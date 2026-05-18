"""
Ingesta DOCX compartida: texto/tablas para RAG + partidas estructuradas + normalización canónica.
"""

from __future__ import annotations

from typing import Any, Dict, List, Tuple

from docx import Document

from app.memory.repository import MemoryRepository
from app.services.tabular_line_item_extract import extract_line_items_from_docx_path
from app.services.economic_normalizer import normalize_line_items, merge_normalized_payload


def _docx_to_text_payload(file_path: str, filename: str) -> Tuple[str, List[Dict[str, Any]]]:
    """Construye contenido textual con párrafos y tablas del DOCX para indexación."""
    doc = Document(file_path)
    blocks: List[str] = [f"### ARCHIVO: {filename} | TIPO: DOCX"]

    paragraph_lines = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
    if paragraph_lines:
        blocks.append("PARRAFOS:\n" + "\n".join(paragraph_lines))

    pages: List[Dict[str, Any]] = []
    for i, table in enumerate(doc.tables):
        rows: List[str] = []
        for row in table.rows:
            cells = [str(c.text or "").strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if not rows:
            continue
        table_name = f"Tabla {i+1}"
        table_text = f"{table_name}:\n" + "\n".join(rows)
        blocks.append(table_text)
        pages.append({"page": table_name, "text": table_text})

    full_text = "\n\n".join(blocks).strip()
    if not pages:
        pages = [{"page": "docx", "text": full_text}]
    return full_text, pages


async def process_docx_document(
    memory: MemoryRepository,
    session_id: str,
    doc_id: str,
    file_path: str,
    filename: str,
) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """
    Construye payload equivalente a OCR para DOCX y persiste partidas + normalizado canónico.
    """
    full_text, pages = _docx_to_text_payload(file_path, filename)

    rows = extract_line_items_from_docx_path(file_path, filename)
    await memory.replace_line_items_for_document(session_id, doc_id, rows)

    try:
        session_state = await memory.get_session(session_id) or {}
        normalized = normalize_line_items(
            session_id=session_id,
            doc_id=doc_id,
            source_filename=filename,
            source_type="docx",
            rows=rows,
            raw_text=full_text,
        )
        updated_state = merge_normalized_payload(session_state, normalized)
        await memory.save_session(session_id, updated_state)
    except Exception:
        pass

    ocr_result: Dict[str, Any] = {
        "extracted_text": full_text,
        "pages": pages,
        "total_pages": len(pages),
        "success": True,
    }
    return ocr_result, rows

