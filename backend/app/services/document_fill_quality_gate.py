from __future__ import annotations

import os
import re
import unicodedata
from dataclasses import dataclass
from typing import Any, Dict, List, Sequence

from docx import Document
from openpyxl import load_workbook

from app.config.settings import settings

_PLACEHOLDER_PATTERNS: Sequence[str] = (
    r"\[\s*[^\]]+\s*\]",
    r"\{\s*[^}]+\s*\}",
    r"\bN\/A\b",
    r"(?<!\S)\.\.\.(?!\S)",
    r"Dato pendiente de confirmar",
)

_DEFERRED_ECONOMIC_PLACEHOLDER_RE = re.compile(
    r"(?i)(tarifa\s+mensual|precio\s+unitario|integraci[oó]n\s+del\s+costo|"
    r"importe\s+mensual|costo\s+mensual|tarifa\s+para\s+horario|tabla\s+de\s+precios|"
    r"an[aá]lisis\s+de\s+precios|precios\s+unitarios)"
)

_DEFERRED_ECONOMIC_FILENAME_RE = re.compile(
    r"(?i)precios[_\s]?unitarios|analisis[_\s]?precios|"
    r"anexo[_\s]?9|cotizaci[oó]n|resumen.*subtotal"
)


def _is_deferred_economic_placeholder(text: str, *, basename: str = "") -> bool:
    """
    Placeholders de precios/tarifas que en etapa ``formats`` aún no tienen fuente
    (la propuesta económica corre después). No deben bloquear el pipeline entero.
    """
    blob = f"{basename} {text or ''}"
    return bool(_DEFERRED_ECONOMIC_PLACEHOLDER_RE.search(blob))


def _is_standalone_ellipsis(text: str) -> bool:
    """Tres puntos como celda vacía (APU), no '750...SAN' en domicilios."""
    compact = re.sub(r"\s+", "", str(text or ""))
    return bool(compact) and bool(re.fullmatch(r"\.{3,}", compact))


def _defer_formats_stage_placeholder(
    *,
    stage: str,
    basename: str,
    text: str,
) -> bool:
    if stage != "formats":
        return False
    if _is_standalone_ellipsis(text):
        return True
    if _DEFERRED_ECONOMIC_FILENAME_RE.search(basename or ""):
        return True
    return _is_deferred_economic_placeholder(text, basename=basename)


_BLANK_SLOT_PATTERNS: Sequence[str] = (
    r"\bQUIEN SUSCRIBE\b[^\n]{0,80}_{4,}",
    r"\bEMPRESA\b[^\n]{0,80}_{4,}",
    r"\bPROCEDIMIENTO\b[^\n]{0,80}_{4,}",
    r"\b(DENOMINACI[OÓ]N SOCIAL|DOMICILIO|AUTORIZACI[OÓ]N DEL GOBIERNO FEDERAL PARA OPERAR|"
    r"NOMBRE O DENOMINACI[OÓ]N SOCIAL|MONTO AFIANZADO|MONEDA|FECHA DE EXPEDICI[OÓ]N|"
    r"OBLIGACI[OÓ]N GARANTIZADA|N[UÚ]MERO ASIGNADO POR \"LA CONTRATANTE\"|OBJETO|"
    r"MONTO DEL CONTRATO|FECHA DE SUSCRIPCI[OÓ]N|TIPO)\s*:\s*_{3,}",
    r"[A-ZÁÉÍÓÚÑ][^\n]{0,120}_{4,}\s*$",
)

_GENERIC_TENDER_MARKERS = {
    "RFC",
    "SAT",
    "IVA",
    "CURP",
    "ISR",
    "IMSS",
    "ISSSTE",
    "MXN",
    "XML",
    "PDF",
    "ZIP",
    "DOCX",
    "XLSX",
}

_ECO_LABELS = ("SUBTOTAL", "IVA", "TOTAL")
_RFC_PATTERN = r"^[A-Z&Ñ]{3,4}\d{6}[A-Z0-9]{3}$"


@dataclass
class FillIssue:
    error_type: str
    severity: str
    document_id: str
    field_key: str
    detected_value: str
    expected_rule: str
    provenance: Dict[str, Any]

    def as_dict(self) -> Dict[str, Any]:
        return {
            "error_type": self.error_type,
            "severity": self.severity,
            "document_id": self.document_id,
            "field_key": self.field_key,
            "detected_value": self.detected_value,
            "expected_rule": self.expected_rule,
            "provenance": self.provenance,
        }


@dataclass(frozen=True)
class FieldPolicy:
    field_key: str
    required: bool
    allow_placeholder: bool
    expected_type: str
    consistency_group: str = ""
    min_confidence: float = 0.0
    pattern: str = ""


@dataclass(frozen=True)
class DocumentPolicy:
    family: str
    policy_id: str
    template_id: str
    tipo: str
    filename_regex: str
    fields: Sequence[FieldPolicy]


class DocumentFieldPolicyRegistry:
    """Registro determinista de políticas de llenado por documento."""

    POLICY_VERSION = "1.4.0"

    _POLICIES: Sequence[DocumentPolicy] = (
        DocumentPolicy(
            family="technical",
            policy_id="tech_carta_presentacion",
            template_id="",
            tipo="tecnico_carta",
            filename_regex=r"CARTA_PRESENTACION",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.8),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.85, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.8),
            ),
        ),
        DocumentPolicy(
            family="formats",
            policy_id="formats_legal_anexos",
            template_id="anexo_7",
            tipo="",
            filename_regex=r"ANEXO_7|ANEXO 7|PERSONALIDAD",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.85),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.9, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.85),
            ),
        ),
        DocumentPolicy(
            family="formats",
            policy_id="formats_legal_conformidad",
            template_id="anexo_11",
            tipo="",
            filename_regex=r"CONFORMIDAD|ANEXO_11|ANEXO 11",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.85),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.9, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.85),
            ),
        ),
        DocumentPolicy(
            family="formats",
            policy_id="formats_legal_manifestacion",
            template_id="anexo_15",
            tipo="",
            filename_regex=r"ANEXO_15|ANEXO 15|MANIFEST",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.85),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.9, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.85),
            ),
        ),
        DocumentPolicy(
            family="technical",
            policy_id="technical_template_mirror",
            template_id="",
            tipo="",
            filename_regex=r"CARTA_PRESENTACION|PROPUESTA_TECNICA|CONTENIDO_NACIONAL|TE-12|PUNTUACION",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.8),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.85, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.8),
            ),
        ),
        DocumentPolicy(
            family="economic",
            policy_id="economic_anexo_ae",
            template_id="anexo_economico",
            tipo="anexo_economico",
            filename_regex=r"ANEXO_AE(?:_|$|\b)|^ANEXO_AE_PROPUESTA_ECONOMICA",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.85),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.9, pattern=_RFC_PATTERN),
                FieldPolicy("subtotal", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
                FieldPolicy("iva", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
                FieldPolicy("total", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
            ),
        ),
        DocumentPolicy(
            family="economic",
            policy_id="economic_tabla_precios",
            template_id="tabla_precios",
            tipo="tabla_precios",
            filename_regex=r"TABLA_PRECIOS_UNITARIOS",
            fields=(
                FieldPolicy("subtotal", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
                FieldPolicy("iva", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
                FieldPolicy("total", True, False, "numeric", consistency_group="economic_totals", min_confidence=0.95),
            ),
        ),
        DocumentPolicy(
            family="economic",
            policy_id="economic_carta_compromiso",
            template_id="carta_compromiso",
            tipo="carta_compromiso",
            filename_regex=r"CARTA_COMPROMISO_PRECIOS|CARTA COMPROMISO",
            fields=(
                FieldPolicy("razon_social", True, False, "text", min_confidence=0.85),
                FieldPolicy("rfc", True, False, "identifier", min_confidence=0.9, pattern=_RFC_PATTERN),
                FieldPolicy("representante_legal", True, False, "text", min_confidence=0.85),
            ),
        ),
    )

    @classmethod
    def resolve_policy(cls, family: str, doc: Dict[str, Any]) -> DocumentPolicy | None:
        template_id = str(doc.get("template_id") or "").strip().lower()
        tipo = str(doc.get("tipo") or "").strip().lower()
        filename = str(os.path.basename(str(doc.get("ruta") or "")) or "").strip().lower()

        for p in cls._POLICIES:
            if p.family != family:
                continue
            if p.template_id and template_id == p.template_id.lower():
                return p
        for p in cls._POLICIES:
            if p.family == family and p.tipo and tipo == p.tipo.lower():
                return p
        for p in cls._POLICIES:
            if p.family == family and p.filename_regex and re.search(p.filename_regex, filename, flags=re.IGNORECASE):
                return p
        return None


class FieldProvenanceResolver:
    """Resuelve procedencia por campo con cascada explícita."""

    ORDER = ("user_override", "normalized_doc", "master_profile", "llm_inference")

    @classmethod
    def resolve(
        cls,
        field_key: str,
        field_provenance: Dict[str, Any],
        fallback_source: str,
        fallback_confidence: float = 0.7,
    ) -> Dict[str, Any]:
        raw = field_provenance.get(field_key)
        if isinstance(raw, dict):
            return {
                "source": str(raw.get("source") or fallback_source),
                "confidence": float(raw.get("confidence", fallback_confidence) or fallback_confidence),
                "anchor": raw.get("anchor"),
            }
        return {"source": fallback_source, "confidence": fallback_confidence, "anchor": None}


def _should_relax_template_gate(basename: str, *, stage: str) -> bool:
    from app.services.document_body_quality import should_relax_fill_quality_gate

    return should_relax_fill_quality_gate(basename, stage=stage)


def _is_placeholder(text: str, *, strict_form_lines: bool = True) -> bool:
    if not text:
        return False
    if _is_separator_only(text):
        return False
    if _is_standalone_ellipsis(text):
        return True
    for pattern in _PLACEHOLDER_PATTERNS:
        if pattern == r"(?<!\S)\.\.\.(?!\S)":
            continue
        if re.search(pattern, text, flags=re.IGNORECASE):
            return True
    if not strict_form_lines:
        return _has_blank_slot_placeholder(text, include_form_line_pattern=False)
    return _has_blank_slot_placeholder(text)


def _normalize_token(value: str) -> str:
    raw = unicodedata.normalize("NFD", str(value or ""))
    raw = "".join(ch for ch in raw if unicodedata.category(ch) != "Mn")
    return raw.upper()


def _is_separator_only(text: str) -> bool:
    compact = re.sub(r"[\s\|]", "", str(text or ""))
    return bool(compact) and bool(re.fullmatch(r"[_\-]+", compact))


def _has_blank_slot_placeholder(text: str, *, include_form_line_pattern: bool = True) -> bool:
    raw = str(text or "").strip()
    if not raw or _is_separator_only(raw):
        return False
    norm = _normalize_token(raw)
    if re.search(r"\bNOMBRE DEL LICITANTE\s*:\s*(?:$|N/A\b|_{3,})", norm, flags=re.IGNORECASE):
        return True
    if "NOMBRE DEL REPRESENTANTE LEGAL" in norm and ("(" in norm or "_" in norm):
        return True
    if "NOMBRE DE LA PERSONA QUE PARTICIPA EN LA PRESENTE LICITACION" in norm:
        return True
    for pattern in _BLANK_SLOT_PATTERNS:
        if not include_form_line_pattern and pattern == r"[A-ZÁÉÍÓÚÑ][^\n]{0,120}_{4,}\s*$":
            continue
        if re.search(pattern, norm, flags=re.IGNORECASE):
            return True
    return False


def _is_placeholder_llm_residual_only(text: str) -> bool:
    """
    En plantillas espejo/catálogo del convocante: solo bloquea restos de LLM
    (corchetes, N/A, «dato pendiente»), **no** campos en blanco del formato
    (NOMBRE DEL LICITANTE: _______, NOMBRE DEL INTERESADO: ___).
    """
    if not text or _is_separator_only(text):
        return False
    if _is_standalone_ellipsis(text):
        return True
    for pattern in _PLACEHOLDER_PATTERNS:
        if pattern == r"(?<!\S)\.\.\.(?!\S)":
            continue
        if re.search(pattern, text, flags=re.IGNORECASE):
            return True
    return False


def _is_placeholder_operational_mirror(text: str) -> bool:
    """Alias retrocompatible: misma política que plantilla relajada."""
    return _is_placeholder_llm_residual_only(text)


def detect_cross_tender_marker(chunks: Sequence[str], session_hint: str) -> str | None:
    """
    Detecta referencias institucionales repetidas que no pertenecen a la sesión.
    """
    hint = _normalize_token(session_hint)
    if not hint:
        return None
    blob = " ".join(str(chunk or "") for chunk in chunks if chunk)
    if not blob:
        return None
    norm_blob = _normalize_token(blob)
    markers = re.findall(r"\(([A-Z][A-Z0-9\-]{3,})\)", norm_blob)
    for marker in sorted(set(markers)):
        if marker in _GENERIC_TENDER_MARKERS:
            continue
        if norm_blob.count(marker) < 2:
            continue
        if marker not in hint:
            return marker
    return None


def _effective_gate_mode() -> str:
    mode = str(getattr(settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "audit") or "audit").strip().lower()
    return "enforce" if mode == "enforce" else "audit"


def _issue_severity(base: str) -> str:
    if _effective_gate_mode() != "enforce":
        return "warn"
    return base


def _profile_required_fields(master_profile: Dict[str, Any]) -> List[FillIssue]:
    issues: List[FillIssue] = []
    required_keys = ("razon_social", "rfc", "representante_legal")
    for key in required_keys:
        value = str((master_profile or {}).get(key) or "").strip()
        if not value or _is_placeholder(value):
            issues.append(
                FillIssue(
                    error_type="required_field_missing",
                    severity=_issue_severity("block"),
                    document_id="profile",
                    field_key=key,
                    detected_value=value,
                    expected_rule="non_empty_and_not_placeholder",
                    provenance={"source": "master_profile", "confidence": 1.0},
                )
            )
    return issues


def _scan_docx(path: str) -> List[str]:
    from app.services.document_contamination_gate import strip_llm_meta_leaks

    out: List[str] = []
    try:
        doc = Document(path)
    except Exception:
        return out
    for p in doc.paragraphs:
        text = strip_llm_meta_leaks((p.text or "").strip())
        if text:
            out.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = strip_llm_meta_leaks((cell.text or "").strip())
                if text:
                    out.append(text)
    return out


def _scan_docx_economic_totals(chunks: Sequence[str]) -> Dict[str, str]:
    out: Dict[str, str] = {}
    rgx = {
        "SUBTOTAL": re.compile(r"(?<![A-Z])SUBTOTAL\s*:\s*\$?\s*([0-9\.,]+)", re.IGNORECASE),
        "IVA": re.compile(
            r"(?:I\.?\s*V\.?\s*A\.?\s*\.?|IVA)(?:\s*\([^)]*\))?\s*:\s*\$?\s*([0-9\.,]+)",
            re.IGNORECASE,
        ),
        "TOTAL": re.compile(
            r"(?<!SUB)TOTAL(?:\s+DE\s+LA\s+PROPUESTA)?\s*:\s*\$?\s*([0-9\.,]+)",
            re.IGNORECASE,
        ),
    }
    for line in chunks:
        for k, p in rgx.items():
            if k in out:
                continue
            m = p.search(line or "")
            if m:
                out[k] = m.group(1)
    return out


def _scan_xlsx_labels(path: str) -> Dict[str, str]:
    out: Dict[str, str] = {}
    try:
        wb = load_workbook(path, data_only=True)
    except Exception:
        return out
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 300), min_col=1, max_col=min(ws.max_column, 12)):
                for idx, cell in enumerate(row):
                    raw = str(cell.value or "").strip().upper()
                    label_key = ""
                    if raw.startswith("SUBTOTAL"):
                        label_key = "SUBTOTAL"
                    elif raw.startswith("IVA") or raw.startswith("I.V.A"):
                        label_key = "IVA"
                    elif raw.startswith("TOTAL"):
                        label_key = "TOTAL"
                    if not label_key or label_key in out:
                        continue
                    if idx + 1 < len(row):
                        right = row[idx + 1].value
                        out[label_key] = "" if right is None else str(right).strip()
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return out


def _scan_xlsx_text(path: str) -> List[str]:
    out: List[str] = []
    try:
        wb = load_workbook(path, data_only=True)
    except Exception:
        return out
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(
                min_row=1,
                max_row=min(ws.max_row, 300),
                min_col=1,
                max_col=min(ws.max_column, 16),
            ):
                for cell in row:
                    text = str(cell.value or "").strip()
                    if text:
                        out.append(text)
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return out


def _as_float(value: Any) -> float | None:
    if value is None:
        return None
    txt = str(value).strip().replace(",", "")
    if not txt:
        return None
    try:
        return float(txt)
    except Exception:
        return None


def _norm_text(value: str) -> str:
    raw = unicodedata.normalize("NFD", str(value or ""))
    raw = "".join(ch for ch in raw if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", raw).strip().lower()


def _contains_profile_value(chunks: Sequence[str], expected_value: str) -> bool:
    needle = _norm_text(expected_value)
    if not needle:
        return False
    return any(needle in _norm_text(chunk) for chunk in chunks if chunk)


def _is_reference_fianza_model(
    *,
    basename: str,
    chunks: Sequence[str],
) -> bool:
    """Detecta formatos oficiales de póliza/fianza que son referencia de afianzadora."""
    blob = _normalize_token(f"{basename} {' '.join(str(chunk or '') for chunk in chunks[:80])}")
    if "FIANZA" not in blob and "AFIANZADORA" not in blob:
        return False
    markers = (
        "AUTORIZACION DEL GOBIERNO FEDERAL PARA OPERAR",
        "NUMERO ASIGNADO POR LA AFIANZADORA",
        "MONTO AFIANZADO",
        "FECHA DE EXPEDICION",
        "MONTO DEL CONTRATO",
        "OBLIGACION GARANTIZADA",
        "LA AFIANZADORA",
    )
    matches = sum(1 for marker in markers if marker in blob)
    return matches >= 3


def _extract_field_value(
    field_key: str,
    *,
    master_profile: Dict[str, Any],
    document_chunks: Sequence[str],
    xlsx_labels: Dict[str, str],
    docx_totals: Dict[str, str],
    provenance_context: Dict[str, Any] | None = None,
    materialization_route: str = "",
) -> str:
    k = field_key.lower()
    if k in ("razon_social", "rfc", "representante_legal"):
        expected = str(master_profile.get(k) or "").strip()
        if expected and _contains_profile_value(document_chunks, expected):
            return expected
        route = str(materialization_route or "").lower()
        if route.startswith("deterministic") or route in ("mirror", "template_locked"):
            from app.core.formats_pilot_slots import is_usable_profile_field_value

            if is_usable_profile_field_value(expected):
                return expected
        return ""
    if k == "subtotal":
        scanned = xlsx_labels.get("SUBTOTAL", "") or docx_totals.get("SUBTOTAL", "")
        if scanned:
            return scanned
    elif k == "iva":
        scanned = xlsx_labels.get("IVA", "") or docx_totals.get("IVA", "")
        if scanned:
            return scanned
    elif k == "total":
        scanned = xlsx_labels.get("TOTAL", "") or docx_totals.get("TOTAL", "")
        if scanned:
            return scanned
    else:
        return ""

    resumen = (
        (provenance_context or {}).get("economic_resumen")
        if isinstance(provenance_context, dict)
        else None
    )
    if isinstance(resumen, dict) and resumen.get(k) is not None:
        val = resumen.get(k)
        if _as_float(val) is not None:
            return str(val)
    return ""


def _requires_economic_totals(
    *,
    stage: str,
    doc: Dict[str, Any],
    policy: DocumentPolicy | None,
) -> bool:
    if policy and any(f.consistency_group == "economic_totals" for f in policy.fields):
        return True
    doc_tipo = str(doc.get("tipo") or "").lower()
    template_id = str(doc.get("template_id") or "").lower()
    blob = f"{os.path.basename(str(doc.get('ruta') or ''))} {doc_tipo} {template_id}".lower()
    return stage == "economic" and bool(
        re.search(r"tabla_precios|anexo_ae", blob)
    )


def _validate_expected_type(value: str, expected_type: str, pattern: str = "") -> tuple[bool, str]:
    if expected_type == "numeric":
        return (_as_float(value) is not None, "numeric_value_required")
    if expected_type == "identifier" and pattern:
        ok = bool(re.match(pattern, str(value or "").strip(), flags=re.IGNORECASE))
        return (ok, f"pattern:{pattern}")
    return (bool(str(value or "").strip()), "non_empty")


def validate_generated_documents_fill(
    *,
    stage: str,
    generated_documents: Sequence[Dict[str, Any]],
    master_profile: Dict[str, Any],
    provenance_context: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    """
    Evalúa calidad de llenado sobre archivos materializados.

    En modo ``audit`` nunca bloquea, pero emite warnings estructurados.
    En modo ``enforce`` convierte reglas críticas a bloqueantes.
    """
    if not bool(getattr(settings, "DOCUMENT_FILL_QUALITY_GATE_ENABLED", True)):
        return {
            "validation_passed": True,
            "blocking_count": 0,
            "warning_count": 0,
            "issues": [],
            "documents_scanned": 0,
            "metrics": {"mode": _effective_gate_mode(), "stage": stage},
        }

    issues: List[FillIssue] = []
    issues.extend(_profile_required_fields(master_profile))
    docs_scanned = 0
    docs_with_policy = 0
    policy_miss_count = 0
    confidence_violations = 0
    field_provenance = (
        provenance_context.get("field_provenance")
        if isinstance(provenance_context, dict) and isinstance(provenance_context.get("field_provenance"), dict)
        else {}
    )
    fallback_source = str((provenance_context or {}).get("source") or "generated_document")
    session_hint = str((provenance_context or {}).get("session_hint") or "").strip()

    for d in generated_documents or []:
        path = str(d.get("ruta") or "").strip()
        if not path or not os.path.exists(path):
            continue
        docs_scanned += 1
        basename = os.path.basename(path)
        ext = os.path.splitext(path)[1].lower()
        policy = DocumentFieldPolicyRegistry.resolve_policy(stage, d)
        if policy is None:
            policy_miss_count += 1
        else:
            docs_with_policy += 1

        chunks: List[str] = []
        xlsx_labels: Dict[str, str] = {}
        docx_totals: Dict[str, str] = {}
        xlsx_chunks: List[str] = []
        if ext == ".docx":
            chunks = _scan_docx(path)
            docx_totals = _scan_docx_economic_totals(chunks)
        elif ext == ".xlsx":
            xlsx_labels = _scan_xlsx_labels(path)
            xlsx_chunks = _scan_xlsx_text(path)
            if str(d.get("fill_status") or "").strip().lower() == "skipped_missing_locator":
                issues.append(
                    FillIssue(
                        error_type="required_field_missing",
                        severity=_issue_severity("block"),
                        document_id=basename,
                        field_key="price_fill",
                        detected_value=str(d.get("valid_locator_count") or 0),
                        expected_rule="economic_excel_requires_valid_locator_mapping",
                        provenance=provenance_context or {"source": "generated_xlsx", "confidence": 1.0},
                    )
                )
            if _requires_economic_totals(stage=stage, doc=d, policy=policy):
                for key in _ECO_LABELS:
                    val = xlsx_labels.get(key, "")
                    if not val:
                        issues.append(
                            FillIssue(
                                error_type="required_field_missing",
                                severity=_issue_severity("block"),
                                document_id=basename,
                                field_key=key.lower(),
                                detected_value=val,
                                expected_rule="label_anchor_has_numeric_value",
                                provenance=provenance_context or {"source": "generated_xlsx", "confidence": 1.0},
                            )
                        )
                        continue
                    if _as_float(val) is None:
                        issues.append(
                            FillIssue(
                                error_type="cross_field_inconsistency",
                                severity=_issue_severity("block"),
                                document_id=basename,
                                field_key=key.lower(),
                                detected_value=val,
                                expected_rule="numeric_value_required",
                                provenance=provenance_context or {"source": "generated_xlsx", "confidence": 1.0},
                            )
                        )

        materialized_chunks = chunks if ext == ".docx" else xlsx_chunks
        relax_template = _should_relax_template_gate(basename, stage=stage)
        if not _is_reference_fianza_model(basename=basename, chunks=materialized_chunks):
            for text in materialized_chunks:
                if relax_template:
                    ph_hit = _is_placeholder_llm_residual_only(text)
                else:
                    ph_hit = _is_placeholder(text, strict_form_lines=True)
                if ph_hit:
                    defer_economic = _defer_formats_stage_placeholder(
                        stage=stage,
                        basename=basename,
                        text=text,
                    )
                    issues.append(
                        FillIssue(
                            error_type="placeholder_detected",
                            severity=_issue_severity("warn" if defer_economic else "block"),
                            document_id=basename,
                            field_key="tarifa_mensual" if defer_economic else "content",
                            detected_value=text[:240],
                            expected_rule=(
                                "deferred_to_economic_stage"
                                if defer_economic
                                else "no_placeholder_tokens"
                            ),
                            provenance=provenance_context
                            or {"source": f"generated_{ext.lstrip('.')}", "confidence": 1.0},
                        )
                    )
                    break
        cross_tender_marker = None if relax_template else detect_cross_tender_marker(
            materialized_chunks, session_hint
        )
        if cross_tender_marker:
            route = str(d.get("materialization_route") or "").lower()
            mirror_mode = str(d.get("mirror_mode") or "").lower()
            is_ingested_mirror = route in ("mirror", "copy_docx") or mirror_mode.startswith(
                ("copy", "mirror", "deterministic")
            )
            issues.append(
                FillIssue(
                    error_type="cross_tender_reference",
                    severity=_issue_severity("warn" if is_ingested_mirror else "block"),
                    document_id=basename,
                    field_key="content",
                    detected_value=cross_tender_marker,
                    expected_rule="matches_current_tender_context",
                    provenance=provenance_context
                    or {"source": f"generated_{ext.lstrip('.')}", "confidence": 1.0},
                )
            )

        _mirror_route = str(d.get("materialization_route") or "").lower() in (
            "mirror",
            "copy_docx",
        ) or str(d.get("mirror_mode") or "").lower().startswith(
            ("copy", "mirror", "deterministic")
        )
        _skip_contamination_block = relax_template and (
            _mirror_route or basename.lower().startswith(("cat_", "cat ", "panel_", "panel "))
        )
        _template_contamination_ok = frozenset(
            {
                "adjudication_language_in_proposal_stage",
                "bases_checklist_in_letter_body",
                "evaluator_perspective_detected",
                "contract_language_in_proposal_stage",
                "document_date_after_submission_deadline",
                "generic_legal_fallback_body",
            }
        )

        if bool(getattr(settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)) and not (
            relax_template and not materialized_chunks
        ):
            from app.services.document_contamination_gate import (
                contamination_hits_to_issues,
                scan_date_after_deadline,
                scan_text_contamination,
            )

            full_text = "\n".join(materialized_chunks)
            if full_text.strip():
                ctx = provenance_context if isinstance(provenance_context, dict) else {}
                hits = scan_text_contamination(
                    full_text, basename=basename, stage=stage
                )
                deadline_iso = str(ctx.get("deadline_dt_iso") or "")
                fecha_es = str(ctx.get("fecha_es") or "")
                if not relax_template:
                    date_hit = scan_date_after_deadline(
                        full_text,
                        deadline_dt_iso=deadline_iso or None,
                        fecha_es=fecha_es,
                    )
                    if date_hit:
                        hits.append(date_hit)
                    from app.services.document_contamination_gate import scan_all_document_dates

                    for extra in scan_all_document_dates(
                        full_text,
                        deadline_dt_iso=deadline_iso or None,
                        canonical_fecha_es=fecha_es,
                    ):
                        if not any(
                            h.error_type == extra.error_type
                            and h.detected_value == extra.detected_value
                            for h in hits
                        ):
                            hits.append(extra)
                if relax_template:
                    hits = [
                        h
                        for h in hits
                        if h.error_type not in _template_contamination_ok
                    ]
                for raw in contamination_hits_to_issues(
                    hits,
                    document_id=basename,
                    provenance=ctx
                    or {"source": f"generated_{ext.lstrip('.')}", "confidence": 1.0},
                ):
                    err_type = str(raw.get("error_type") or "")
                    if _skip_contamination_block and err_type in _template_contamination_ok:
                        continue
                    sev = str(raw.get("severity") or _issue_severity("block"))
                    if _skip_contamination_block and sev == "block":
                        sev = _issue_severity("warn")
                    issues.append(
                        FillIssue(
                            error_type=err_type,
                            severity=sev,
                            document_id=basename,
                            field_key=str(raw.get("field_key") or "content"),
                            detected_value=str(raw.get("detected_value") or ""),
                            expected_rule=str(raw.get("expected_rule") or ""),
                            provenance=raw.get("provenance")
                            if isinstance(raw.get("provenance"), dict)
                            else {},
                        )
                    )

        if ext == ".docx" and stage in ("formats", "technical") and materialized_chunks:
            from app.services.document_body_quality import scan_materialized_doc_text
            from app.services.document_contamination_gate import is_apu_document

            shell_hit = scan_materialized_doc_text(
                "\n".join(materialized_chunks),
                basename=basename,
            )
            defer_shell = relax_template or any(
                _defer_formats_stage_placeholder(
                    stage=stage, basename=basename, text=chunk
                )
                for chunk in materialized_chunks
            ) or (
                stage == "formats"
                and is_apu_document(basename, "\n".join(materialized_chunks[:3]))
            )
            if (
                shell_hit
                and not defer_shell
                and not _is_reference_fianza_model(basename=basename, chunks=materialized_chunks)
            ):
                issues.append(
                    FillIssue(
                        error_type=str(shell_hit["error_type"]),
                        severity=_issue_severity("block"),
                        document_id=basename,
                        field_key=str(shell_hit.get("field_key") or "content"),
                        detected_value=str(shell_hit.get("detected_value") or ""),
                        expected_rule=str(shell_hit.get("expected_rule") or ""),
                        provenance=provenance_context
                        or {"source": f"generated_{ext.lstrip('.')}", "confidence": 1.0},
                    )
                )

        defer_policy_fields = stage == "formats" and any(
            _defer_formats_stage_placeholder(
                stage=stage, basename=basename, text=chunk
            )
            for chunk in materialized_chunks
        )

        if policy is not None and not defer_policy_fields:
            totals: Dict[str, float] = {}
            for fp in policy.fields:
                value = _extract_field_value(
                    fp.field_key,
                    master_profile=master_profile,
                    document_chunks=materialized_chunks,
                    xlsx_labels=xlsx_labels,
                    docx_totals=docx_totals,
                    provenance_context=provenance_context if isinstance(provenance_context, dict) else None,
                    materialization_route=str(d.get("materialization_route") or ""),
                )
                fallback_confidence = float((provenance_context or {}).get("confidence", 0.7))
                prov = FieldProvenanceResolver.resolve(fp.field_key, field_provenance, fallback_source, fallback_confidence)

                if fp.required and not value:
                    issues.append(
                        FillIssue(
                            error_type="required_field_missing",
                            severity=_issue_severity("block"),
                            document_id=basename,
                            field_key=fp.field_key,
                            detected_value=value,
                            expected_rule="required_by_policy",
                            provenance=prov,
                        )
                    )
                    continue
                if value and (not fp.allow_placeholder) and _is_placeholder(value):
                    issues.append(
                        FillIssue(
                            error_type="placeholder_detected",
                            severity=_issue_severity("block"),
                            document_id=basename,
                            field_key=fp.field_key,
                            detected_value=value,
                            expected_rule="no_placeholder_tokens",
                            provenance=prov,
                        )
                    )
                if value:
                    ok_type, rule = _validate_expected_type(value, fp.expected_type, fp.pattern)
                    if not ok_type:
                        issues.append(
                            FillIssue(
                                error_type="cross_field_inconsistency",
                                severity=_issue_severity("block"),
                                document_id=basename,
                                field_key=fp.field_key,
                                detected_value=value,
                                expected_rule=rule,
                                provenance=prov,
                            )
                        )
                    if fp.expected_type == "numeric":
                        n = _as_float(value)
                        if n is not None:
                            totals[fp.field_key.lower()] = n
                    # Si el dato ya está materializado en el archivo, no bloquear por confianza de inferencia.
                    prov = {
                        **prov,
                        "confidence": max(
                            float(prov.get("confidence", 0.0) or 0.0),
                            float(fp.min_confidence or 0.0),
                            0.95,
                        ),
                    }
                if fp.min_confidence > 0 and float(prov.get("confidence", 0.0) or 0.0) < fp.min_confidence:
                    confidence_violations += 1
                    issues.append(
                        FillIssue(
                            error_type="source_confidence_insufficient",
                            severity=_issue_severity("block"),
                            document_id=basename,
                            field_key=fp.field_key,
                            detected_value=value,
                            expected_rule=f"confidence>={fp.min_confidence}",
                            provenance=prov,
                        )
                    )

            if any(f.consistency_group == "economic_totals" for f in policy.fields):
                s = totals.get("subtotal")
                i = totals.get("iva")
                t = totals.get("total")
                if s is not None and i is not None and t is not None:
                    if abs((s + i) - t) > 0.02:
                        issues.append(
                            FillIssue(
                                error_type="cross_field_inconsistency",
                                severity=_issue_severity("block"),
                                document_id=basename,
                                field_key="economic_totals",
                                detected_value=f"subtotal={s},iva={i},total={t}",
                                expected_rule="total == subtotal + iva",
                                provenance={"source": "materialized_document", "confidence": 1.0},
                            )
                        )

    blocking = sum(1 for i in issues if i.severity == "block")
    warnings = sum(1 for i in issues if i.severity == "warn")
    return {
        "validation_passed": blocking == 0,
        "policy_version": DocumentFieldPolicyRegistry.POLICY_VERSION,
        "blocking_count": blocking,
        "warning_count": warnings,
        "issues": [i.as_dict() for i in issues],
        "documents_scanned": docs_scanned,
        "documents_with_policy": docs_with_policy,
        "metrics": {
            "mode": _effective_gate_mode(),
            "stage": stage,
            "issues_total": len(issues),
            "policy_miss_count": policy_miss_count,
            "policy_miss_ratio": (
                round(policy_miss_count / docs_scanned, 3) if docs_scanned else 0.0
            ),
            "confidence_violations": confidence_violations,
        },
    }
