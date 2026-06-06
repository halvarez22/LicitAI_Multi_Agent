"""
Contexto de experiencia / contratos previos desde fuentes de sesión (Fuentes).

100 % agnóstico de licitación y empresa: solo heurísticas estructurales + perfil de evidencia.
"""
from __future__ import annotations

import re
from typing import Any, Dict, Iterator, List, Optional, Tuple

from app.services.evidence_profile_service import (
    build_evidence_profile_from_documents,
)

_EXPERIENCE_FILENAME_TOKENS = (
    "experiencia",
    "curriculum",
    "curriculo",
    "cv ",
    " cv.",
    "contrato",
    "referencias",
    "clientes",
)

_SKIP_FILENAME_TOKENS = (
    "modelo contrato",
    "anexo s",
    "bases",
    "convocatoria",
    "matriz",
)

# Misma familia de patrones que evidence_profile_service (sin acoplar IDs concretos).
_CONTRATO_ID_RE = re.compile(
    r"(?:contrato\s+(?:n[uú]mero|no\.?|n[°º#])\s*|contrato\s+#\s*)"
    r"([A-Z0-9][A-Z0-9\-\/\.]{3,40})",
    re.IGNORECASE,
)

_TEL_RE = re.compile(
    r"(?:Tel\.?|tel[eé]fono|Tel[eé]f)\s*[:\.]?\s*([\d\s\(\)\-+]{7,})",
    re.IGNORECASE,
)

_CP_DOMICILIO_RE = re.compile(
    r"([^\n]{8,160}C\.?\s*P\.?\s*\d{5}[^\n]*)",
    re.IGNORECASE,
)

_CLIENT_NAME_PATTERNS = (
    r"para las unidades de(?:l| la| este| esta)\s+(.{6,160}?)"
    r"(?:,\s*con|\.\s*Con|\s+con una|\s+con un|\s+vigencia|\.)",
    r"correspondiente al\s+(.{6,160}?)"
    r"(?:,\s*con|\.\s*Con|\s+con la empresa|\s+con el|\s+con)",
    r"servicio de\s+(.{6,120}?)\s+para",
    r"prestaci[oó]n de\s+(.{6,120}?)\s+(?:para|en|del)",
    r"(?:organismo|dependencia|instituci[oó]n|cliente|contratante)\s*[:\-]\s*(.{6,120}?)(?:\n|\.|,)",
)

_HEADER_SKIP_TOKENS = (
    "integridad",
    "página",
    "pagina",
    "nitropdf",
    "www.",
    "http",
    "la 12 ",
    "m7f",
)


def req_needs_company_experience(req_id: str, req_nombre: str, req_desc: str) -> bool:
    """True si el requisito pide relación de clientes, currículum o experiencia."""
    blob = f"{req_id} {req_nombre} {req_desc}".upper()
    keys = (
        "TE-03",
        "CLIENTE",
        "CURRICULUM",
        "CURRÍCULUM",
        "EXPERIENCIA",
        "PRINCIPALES CLIENTES",
        "RELACION DE",
        "RELACIÓN DE",
        "TRAYECTORIA",
        "CONTRATOS PREVIOS",
        "CONTRATOS ANTERIORES",
    )
    return any(k in blob for k in keys)


def _is_experience_source_filename(filename: str) -> bool:
    fn = str(filename or "").lower()
    if not fn:
        return False
    if any(t in fn for t in _SKIP_FILENAME_TOKENS):
        return False
    return any(t in fn for t in _EXPERIENCE_FILENAME_TOKENS)


def _clean_label(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "").strip())
    text = re.sub(r"^[\-\|\s\d\.]+", "", text)
    return text.strip(" .,;:-")


def _extract_header_organization(block: str) -> str:
    """Nombre de dependencia/cliente en membrete (antes de Asunto / cuerpo)."""
    head = block[:900]
    if re.search(r"Asunto\s*:", head, re.IGNORECASE):
        head = re.split(r"Asunto\s*:", head, maxsplit=1, flags=re.IGNORECASE)[0]
    candidates: List[str] = []
    for line in head.splitlines():
        line = _clean_label(line)
        if not (15 <= len(line) <= 140):
            continue
        low = line.lower()
        if any(t in low for t in _HEADER_SKIP_TOKENS):
            continue
        if not re.search(r"[A-Za-zÁÉÍÓÚáéíóúÑñ]{8,}", line):
            continue
        if re.fullmatch(r"[\d\s\-LA]+", line, re.IGNORECASE):
            continue
        candidates.append(line)
    if not candidates:
        return ""
    return _clean_label(max(candidates, key=len))


def _extract_client_name_generic(blob: str) -> str:
    for pat in _CLIENT_NAME_PATTERNS:
        m = re.search(pat, blob, re.IGNORECASE | re.DOTALL)
        if m:
            name = _clean_label(m.group(1))
            if len(name) >= 6:
                return name
    header = _extract_header_organization(blob)
    if header:
        return header
    return ""


def _iter_contract_windows(text: str, *, radius: int = 220) -> Iterator[Tuple[str, str]]:
    """(contrato_id, ventana de texto) para cada mención de contrato."""
    if not text:
        return
    seen_spans: set[Tuple[int, int]] = set()
    for m in _CONTRATO_ID_RE.finditer(text):
        span = (m.start(), m.end())
        if span in seen_spans:
            continue
        seen_spans.add(span)
        cid = str(m.group(1) or "").strip().upper()
        if not cid:
            continue
        start = max(0, m.start() - radius)
        end = min(len(text), m.end() + radius)
        yield cid, text[start:end]


def _parse_experience_block(block: str, *, contrato_hint: str = "") -> Optional[Dict[str, str]]:
    blob = block.strip()
    if len(blob) < 40:
        return None

    contrato = str(contrato_hint or "").strip().upper()
    if not contrato:
        m = _CONTRATO_ID_RE.search(blob)
        contrato = m.group(1).strip().upper() if m else ""

    nombre = _extract_client_name_generic(blob)
    if not nombre and contrato:
        nombre = f"Referencia contrato {contrato}"
    if not nombre:
        return None

    tel_m = _TEL_RE.search(blob)
    telefono = _clean_label(tel_m.group(1)) if tel_m else ""

    domicilio = ""
    cp_m = _CP_DOMICILIO_RE.search(blob)
    if cp_m:
        domicilio = _clean_label(cp_m.group(1))

    row: Dict[str, str] = {"nombre": nombre}
    if domicilio:
        row["domicilio"] = domicilio
    if telefono:
        row["telefono"] = telefono
    if contrato:
        row["contrato"] = contrato
    return row


def _parse_client_rows_from_experience_text(text: str) -> List[Dict[str, str]]:
    if not text or len(text) < 80:
        return []
    rows: List[Dict[str, str]] = []
    seen_keys: set[str] = set()

    for cid, window in _iter_contract_windows(text):
        row = _parse_experience_block(window, contrato_hint=cid)
        if row:
            key = str(row.get("contrato") or row.get("nombre") or "").lower()
            if key and key not in seen_keys:
                seen_keys.add(key)
                rows.append(row)

    if not rows:
        blocks = re.split(
            r"(?=Asunto\s*:|A quien corresponda|Por medio de la presente|CONSTANCIA)",
            text,
            flags=re.IGNORECASE,
        )
        for block in blocks:
            row = _parse_experience_block(block)
            if not row:
                continue
            key = str(row.get("contrato") or row.get("nombre") or "").lower()
            if key and key not in seen_keys:
                seen_keys.add(key)
                rows.append(row)

    return rows


def _rows_from_evidence_profile(docs: List[Any]) -> List[Dict[str, str]]:
    """Filas mínimas desde contratos_previos canónicos (sin inventar nombres)."""
    profile = build_evidence_profile_from_documents(docs or [])
    fields = profile.get("fields") if isinstance(profile, dict) else {}
    entry = (fields or {}).get("contratos_previos") if isinstance(fields, dict) else None
    vals = entry.get("value") if isinstance(entry, dict) else None
    if not isinstance(vals, list):
        return []
    rows: List[Dict[str, str]] = []
    for item in vals:
        if not isinstance(item, dict):
            continue
        cid = str(item.get("contrato_id") or "").strip().upper()
        if not cid:
            continue
        src = str(item.get("source_doc") or "").strip()
        rows.append(
            {
                "nombre": f"Referencia contrato {cid}",
                "contrato": cid,
                "source_doc": src,
            }
        )
    return rows


def extract_client_references_from_documents(docs: List[Any]) -> List[Dict[str, str]]:
    """
    Extrae filas nombre/domicilio/teléfono desde PDFs de experiencia en Fuentes.
    Sin nombres de cliente hardcodeados: solo texto de fuentes + perfil de evidencia.
    """
    merged: List[Dict[str, str]] = []
    by_contrato: Dict[str, Dict[str, str]] = {}

    for doc in docs or []:
        content = doc.get("content") if isinstance(doc.get("content"), dict) else {}
        if not content:
            continue
        if str(content.get("status") or "").upper() not in ("ANALYZED", "COMPLETED", "OK"):
            continue
        filename = str(content.get("filename") or content.get("name") or "")
        if not _is_experience_source_filename(filename):
            continue
        text = str(content.get("extracted_text") or doc.get("extracted_text") or "")
        for row in _parse_client_rows_from_experience_text(text):
            row = dict(row)
            row["source_doc"] = filename
            if not row.get("telefono"):
                tel_m = _TEL_RE.search(text)
                if tel_m:
                    row["telefono"] = _clean_label(tel_m.group(1))
            if not row.get("domicilio"):
                cp_m = _CP_DOMICILIO_RE.search(text)
                if cp_m:
                    row["domicilio"] = _clean_label(cp_m.group(1))
            key = str(row.get("contrato") or row.get("nombre") or "").lower()
            if not key:
                continue
            prev = by_contrato.get(key)
            if prev:
                for fld in ("domicilio", "telefono", "nombre"):
                    if not prev.get(fld) and row.get(fld):
                        prev[fld] = row[fld]
            else:
                by_contrato[key] = row

    for row in _rows_from_evidence_profile(docs):
        key = str(row.get("contrato") or "").lower()
        if not key:
            continue
        if key in by_contrato:
            if not by_contrato[key].get("source_doc") and row.get("source_doc"):
                by_contrato[key]["source_doc"] = row["source_doc"]
            continue
        by_contrato[key] = row

    merged = list(by_contrato.values())
    return merged[:8]


async def build_company_experience_context_block(
    memory: Any,
    session_id: str,
    *,
    max_doc_chars: int = 4000,
    max_contracts: int = 10,
) -> str:
    docs = await memory.get_documents(session_id) or []
    if not docs:
        return ""

    session_state: Dict[str, Any] = {}
    if hasattr(memory, "get_session"):
        try:
            session_state = await memory.get_session(session_id) or {}
        except Exception:
            session_state = {}

    parts: List[str] = []

    try:
        from app.services.document_catalog_service import (
            get_entries_by_use_case,
        )
        from app.contracts.document_catalog import DocumentCatalogUseCase

        exp_entries = get_entries_by_use_case(
            session_state, DocumentCatalogUseCase.FILL_TE03_CLIENTS.value
        )
        for entry in exp_entries:
            ents = entry.entities if isinstance(entry.entities, dict) else {}
            refs = ents.get("client_refs") or []
            if not refs:
                continue
            parts.append(
                f"REFERENCIAS DE CLIENTES (catálogo — {entry.filename}):"
            )
            for row in refs[:max_contracts]:
                if not isinstance(row, dict):
                    continue
                name = str(row.get("nombre") or "").strip()
                dom = str(row.get("domicilio") or "").strip()
                tel = str(row.get("telefono") or "").strip()
                cid = str(row.get("contrato") or "").strip()
                line = "- "
                if name:
                    line += name
                elif cid:
                    line += f"Contrato {cid}"
                else:
                    line += "Cliente detectado"
                if dom:
                    line += f" | Domicilio: {dom}"
                if tel:
                    line += f" | Tel: {tel}"
                parts.append(line)
    except Exception:
        pass

    profile = build_evidence_profile_from_documents(docs)
    fields = profile.get("fields") if isinstance(profile, dict) else {}
    contratos_entry = (fields or {}).get("contratos_previos") if isinstance(fields, dict) else None
    if isinstance(contratos_entry, dict):
        vals = contratos_entry.get("value")
        if isinstance(vals, list) and vals:
            parts.append("CONTRATOS / EXPERIENCIA DETECTADOS EN FUENTES DE LA SESIÓN:")
            for item in vals[:max_contracts]:
                if not isinstance(item, dict):
                    continue
                cid = str(item.get("contrato_id") or "").strip()
                src = str(item.get("source_doc") or "").strip()
                elems = item.get("elementos_vigilancia")
                line = f"- Contrato {cid}" if cid else "- Contrato registrado"
                if elems:
                    line += f" ({elems} elementos)"
                if src:
                    line += f" — documento: {src}"
                parts.append(line)

    for doc in docs:
        content = doc.get("content") if isinstance(doc.get("content"), dict) else {}
        if not content:
            continue
        if str(content.get("status") or "").upper() not in ("ANALYZED", "COMPLETED", "OK"):
            continue
        filename = str(content.get("filename") or content.get("name") or "")
        doc_id = str(doc.get("id") or "")
        if session_state:
            try:
                from app.services.document_catalog_service import catalog_from_session_state

                catalog = catalog_from_session_state(session_state)
                if catalog and doc_id in catalog.entries:
                    entry = catalog.entries[doc_id]
                    role = entry.doc_role.value if hasattr(entry.doc_role, "value") else str(entry.doc_role)
                    if role != "company_experience":
                        continue
            except Exception:
                if not _is_experience_source_filename(filename):
                    continue
        elif not _is_experience_source_filename(filename):
            continue
        text = str(content.get("extracted_text") or doc.get("extracted_text") or "").strip()
        if len(text) < 80:
            continue
        snippet = text[:max_doc_chars]
        if len(text) > max_doc_chars:
            snippet += "\n…(texto truncado)"
        parts.append(f"--- TEXTO DE FUENTE: {filename} ---\n{snippet}")

    if not parts:
        return ""

    return (
        "EXPERIENCIA Y CONTRATOS PREVIOS (desde Fuentes / documentos de la empresa en esta sesión):\n"
        + "\n".join(parts)
        + "\n\nUsa estos datos literales para la tabla o relación de clientes. "
        "No inventes filas ni uses corchetes."
    )


def build_experience_sources_ux_summary(
    docs: List[Any],
    session_state: Optional[Dict[str, Any]] = None,
) -> Optional[str]:
    """Resumen corto para chat cuando hay experiencia en Fuentes (sin nombres fijos)."""
    if session_state:
        try:
            from app.services.document_catalog_service import experience_client_refs_from_catalog

            catalog_refs = experience_client_refs_from_catalog(session_state)
            if catalog_refs:
                src = str(catalog_refs[0].get("source_doc") or "").strip()
                src_label = f"**{src}**" if src else "**documentos de experiencia en Fuentes**"
                labels: List[str] = []
                for r in catalog_refs[:3]:
                    name = str(r.get("nombre") or "").strip()
                    cid = str(r.get("contrato") or "").strip()
                    if name:
                        labels.append(name)
                    elif cid:
                        labels.append(f"contrato {cid}")
                if labels:
                    listed = ", ".join(f"**{lbl}**" for lbl in labels)
                    extra = f" (+{len(catalog_refs) - len(labels)} más)" if len(catalog_refs) > len(labels) else ""
                    return (
                        f"Ya tengo {src_label} con referencias de {listed}{extra}. "
                        "Pulsa **Generar** otra vez para usarlas en la tabla de clientes del anexo técnico."
                    )
        except Exception:
            pass

    refs = extract_client_references_from_documents(docs)
    if not refs:
        profile = build_evidence_profile_from_documents(docs or [])
        fields = profile.get("fields") if isinstance(profile, dict) else {}
        entry = (fields or {}).get("contratos_previos") if isinstance(fields, dict) else None
        vals = entry.get("value") if isinstance(entry, dict) else None
        if isinstance(vals, list) and vals:
            ids = [str(v.get("contrato_id") or "").strip() for v in vals if isinstance(v, dict)]
            ids = [x for x in ids if x][:3]
            src = str(vals[0].get("source_doc") or "").strip() if vals else ""
            label = f"**{src}**" if src else "**documentos de experiencia en Fuentes**"
            if ids:
                return (
                    f"Ya tengo {label} con contratos **{', '.join(ids)}**. "
                    "Pulsa **Generar** otra vez para volcarlos al anexo técnico."
                )
        return None

    src = str(refs[0].get("source_doc") or "").strip()
    src_label = f"**{src}**" if src else "**documentos de experiencia en Fuentes**"
    labels: List[str] = []
    for r in refs[:3]:
        name = str(r.get("nombre") or "").strip()
        cid = str(r.get("contrato") or "").strip()
        if name:
            labels.append(name)
        elif cid:
            labels.append(f"contrato {cid}")
    if not labels:
        return (
            f"Ya tengo {src_label} con referencias de experiencia. "
            "Pulsa **Generar** otra vez para usarlas en la tabla de clientes."
        )
    listed = ", ".join(f"**{lbl}**" for lbl in labels)
    extra = f" (+{len(refs) - len(labels)} más)" if len(refs) > len(labels) else ""
    return (
        f"Ya tengo {src_label} con referencias de {listed}{extra}. "
        "Pulsa **Generar** otra vez para usarlas en la tabla de clientes del anexo técnico."
    )


_CLIENT_PLACEHOLDER_CELL_RE = re.compile(
    r"\[(?:Nombre|Domicilio|Tel[eé]fono)\s+del\s+cliente\s+\d+\]",
    re.IGNORECASE,
)


def fill_te03_client_placeholders(path: str, clients: List[Dict[str, str]]) -> bool:
    """Sustituye placeholders genéricos [Domicilio del cliente N] por datos de Fuentes."""
    if not path or not clients:
        return False
    try:
        import docx
    except ImportError:
        return False
    if not str(path).lower().endswith(".docx"):
        return False
    try:
        doc = docx.Document(path)
    except Exception:
        return False

    changed = False
    client_idx = 0

    def _apply_client(text: str, idx: int) -> Tuple[str, bool]:
        if idx >= len(clients):
            return text, False
        c = clients[idx]
        n = idx + 1
        new = text
        subs = (
            (rf"\[Nombre del cliente\s+{n}\]", str(c.get("nombre") or "")),
            (rf"\[Domicilio del cliente\s+{n}\]", str(c.get("domicilio") or "")),
            (rf"\[Tel[eé]fono del cliente\s+{n}\]", str(c.get("telefono") or "")),
        )
        local_changed = False
        for pat, val in subs:
            if val and re.search(pat, new, re.IGNORECASE):
                new = re.sub(pat, val, new, flags=re.IGNORECASE)
                local_changed = True
        return new, local_changed

    def _process_text(text: str) -> str:
        nonlocal changed, client_idx
        if not text or not _CLIENT_PLACEHOLDER_CELL_RE.search(text):
            return text
        new_text, did = _apply_client(text, client_idx)
        if did:
            changed = True
            client_idx += 1
            return new_text
        return text

    for para in doc.paragraphs:
        new_text = _process_text(para.text)
        if new_text != para.text:
            para.text = new_text
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text or "" for cell in row.cells)
            if not _CLIENT_PLACEHOLDER_CELL_RE.search(row_text):
                continue
            if client_idx >= len(clients):
                break
            c = clients[client_idx]
            n = client_idx + 1
            for cell in row.cells:
                new_text = cell.text or ""
                for pat, val in (
                    (rf"\[Nombre del cliente\s+{n}\]", str(c.get("nombre") or "")),
                    (rf"\[Domicilio del cliente\s+{n}\]", str(c.get("domicilio") or "")),
                    (rf"\[Tel[eé]fono del cliente\s+{n}\]", str(c.get("telefono") or "")),
                ):
                    if val and re.search(pat, new_text, re.IGNORECASE):
                        new_text = re.sub(pat, val, new_text, flags=re.IGNORECASE)
                if new_text != cell.text:
                    cell.text = new_text
                    changed = True
            client_idx += 1

    if changed:
        doc.save(path)
    return changed
