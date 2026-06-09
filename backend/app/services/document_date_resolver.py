"""
Resuelve la fecha a imprimir en documentos de propuesta (universal, sin hardcode por licitación).

Cascada: override usuario > cronograma (presentación de proposiciones) > fecha sistema acotada al hito.
"""
from __future__ import annotations

import os
import re
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

_DATE_ES_BODY_REPLACE_RE = re.compile(
    r"\b(\d{1,2})\s+de\s+("
    r"enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre"
    r")\s+de\s+(\d{4})\b",
    re.IGNORECASE,
)

from app.agents.analyst import normalize_cronograma_dict
from app.config.settings import settings
from app.services.cronograma_bases_extract import parse_spanish_date_fragment

_MESES_ES_OUT = (
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
)

_DOCUMENT_DATE_OVERRIDE_SESSION_KEYS = (
    "document_date_user_override",
    "document_date_override",
    "fecha_documento",
    "fecha_documentos",
    "fecha_formatos",
    "fecha_formatos_administrativos",
)

_DOCUMENT_DATE_OVERRIDE_PROFILE_KEYS = (
    "document_date_user_override",
    "document_date_override",
    "fecha_documento",
    "fecha_documentos",
)

_DOCUMENT_DATE_CHAT_PATTERNS = (
    re.compile(
        r"(?is)(?:fecha\s+(?:de\s+)?(?:los\s+)?(?:formatos?|documentos?|anexos?|cartas?|expediente)"
        r"|fecha\s+documental|usar\s+fecha|fecha\s+can[oó]nica)"
        r"[^0-9]{0,40}(\d{1,2}\s+de\s+[a-záéíóúñ]+\s+(?:de|del)\s+(?:año\s+)?20\d{2})"
    ),
    re.compile(
        r"(?is)(?:fecha\s*[:=]\s*)"
        r"(\d{1,2}\s+de\s+[a-záéíóúñ]+\s+(?:de|del)\s+(?:año\s+)?20\d{2})"
    ),
)


def _subtract_business_days(dt: datetime, days: int) -> datetime:
    if days <= 0:
        return dt
    cur = dt
    remaining = days
    while remaining > 0:
        cur -= timedelta(days=1)
        if cur.weekday() < 5:
            remaining -= 1
    return cur


def _format_fecha_es(dt: datetime) -> str:
    return f"{dt.day} de {_MESES_ES_OUT[dt.month - 1]} de {dt.year}"


def _format_fecha_corta(dt: datetime) -> str:
    return dt.strftime("%d/%m/%Y")


def _cronograma_from_session(session_state: Dict[str, Any]) -> Dict[str, str]:
    if not isinstance(session_state, dict):
        return normalize_cronograma_dict({})
    for container_key in ("last_analysis", "analysis_snapshot"):
        raw = session_state.get(container_key)
        if isinstance(raw, dict) and raw.get("cronograma"):
            return normalize_cronograma_dict(raw.get("cronograma"))
    er = session_state.get("execution_results")
    if isinstance(er, dict):
        for step in ("analyst", "compliance", "datagap"):
            block = er.get(step)
            if isinstance(block, dict):
                data = block.get("data") if isinstance(block.get("data"), dict) else block
                if isinstance(data, dict) and data.get("cronograma"):
                    return normalize_cronograma_dict(data.get("cronograma"))
    return normalize_cronograma_dict(session_state.get("cronograma"))


def _hito_date_raw(hito: Dict[str, Any]) -> str:
    """Texto de fecha de hito usable en cronograma (español o ISO → español)."""
    raw = str(hito.get("fecha_texto_raw") or "").strip()
    if raw:
        return raw
    iso_val = str(hito.get("fecha_hora") or "").strip()
    if not iso_val:
        return ""
    if "T" in iso_val:
        try:
            dt = datetime.fromisoformat(iso_val.replace("Z", ""))
            return _format_fecha_es(dt)
        except (TypeError, ValueError):
            pass
    return iso_val


def _cronograma_from_submission_checklist(session_state: Dict[str, Any]) -> Dict[str, str]:
    """Fallback: fechas persistidas en checklist de hitos."""
    checklist = session_state.get("submission_checklist")
    if not isinstance(checklist, dict):
        return {}
    hitos = checklist.get("hitos") or []
    if not isinstance(hitos, list):
        return {}
    out: Dict[str, str] = {}
    for hito in hitos:
        if not isinstance(hito, dict):
            continue
        hid = str(hito.get("id") or "").strip()
        raw = _hito_date_raw(hito)
        if hid and raw:
            out[hid] = raw
    return out


def _is_usable_cronograma_value(val: str) -> bool:
    text = str(val or "").strip()
    if not text:
        return False
    low = text.lower()
    if low in ("no especificado", "n/a", "s/d", "sd", "...", "pendiente"):
        return False
    if parse_spanish_date_fragment(text):
        return True
    if "T" in text:
        try:
            datetime.fromisoformat(text.replace("Z", ""))
            return True
        except (TypeError, ValueError):
            return False
    return False


def _merged_cronograma(session_state: Dict[str, Any]) -> Dict[str, str]:
    cronograma = _cronograma_from_session(session_state)
    checklist_cron = _cronograma_from_submission_checklist(session_state)
    merged = dict(cronograma)
    for key, val in checklist_cron.items():
        if _is_usable_cronograma_value(val):
            merged[key] = val
    merged = normalize_cronograma_dict(merged)
    for key, val in checklist_cron.items():
        if _is_usable_cronograma_value(val):
            merged[key] = val
    bases_blob = str(session_state.get("bases_corpus_hint") or "").strip()
    if bases_blob:
        from app.services.cronograma_bases_extract import merge_cronograma_with_bases

        merged = merge_cronograma_with_bases(merged, bases_blob)
    return merged


def _deadline_from_cronograma(
    cronograma: Dict[str, str],
    hito_key: str,
) -> tuple[str, Optional[datetime]]:
    deadline_raw = str(
        cronograma.get(hito_key) or cronograma.get("recepcion_proposiciones") or ""
    ).strip()
    deadline_dt = parse_spanish_date_fragment(deadline_raw) if deadline_raw else None
    if not deadline_dt and deadline_raw and "T" in deadline_raw:
        try:
            deadline_dt = datetime.fromisoformat(deadline_raw.replace("Z", ""))
        except (TypeError, ValueError):
            deadline_dt = None
    return deadline_raw, deadline_dt


def _doc_date_from_deadline(deadline_dt: datetime, offset_business_days: int) -> datetime:
    doc_dt = _subtract_business_days(deadline_dt, offset_business_days)
    if doc_dt.date() > deadline_dt.date():
        doc_dt = _subtract_business_days(deadline_dt, 1)
    now = datetime.now()
    if now.date() <= deadline_dt.date() and doc_dt.date() > now.date():
        doc_dt = now
    return doc_dt


def extract_document_date_user_override(session_state: Optional[Dict[str, Any]]) -> Optional[str]:
    """
    Localiza override de fecha documental en sesión/perfil (HITL, chat, UI).

    Returns:
        Fragmento de fecha en español parseable, o ``None``.
    """
    if not isinstance(session_state, dict):
        return None

    candidates: List[str] = []
    for key in _DOCUMENT_DATE_OVERRIDE_SESSION_KEYS:
        val = session_state.get(key)
        if isinstance(val, str) and val.strip():
            candidates.append(val.strip())

    profile = session_state.get("master_profile")
    if isinstance(profile, dict):
        for key in _DOCUMENT_DATE_OVERRIDE_PROFILE_KEYS:
            val = profile.get(key)
            if isinstance(val, str) and val.strip():
                candidates.append(val.strip())

    user_inputs = session_state.get("user_inputs")
    if isinstance(user_inputs, dict):
        for key in _DOCUMENT_DATE_OVERRIDE_SESSION_KEYS:
            val = user_inputs.get(key)
            if isinstance(val, str) and val.strip():
                candidates.append(val.strip())

    for raw in candidates:
        parsed = parse_spanish_date_fragment(raw)
        if parsed:
            return _format_fecha_es(parsed)
    return None


def parse_document_date_override_from_chat(text: str) -> Optional[str]:
    """
    Extrae fecha documental explícita del mensaje de chat (sin hardcode por licitación).
    """
    blob = str(text or "").strip()
    if not blob:
        return None
    for pat in _DOCUMENT_DATE_CHAT_PATTERNS:
        m = pat.search(blob)
        if not m:
            continue
        parsed = parse_spanish_date_fragment(m.group(1))
        if parsed:
            return _format_fecha_es(parsed)
    return None


def apply_document_date_override_from_chat(
    session_state: Dict[str, Any],
    user_message: str,
) -> Dict[str, Any]:
    """
    Persiste override de fecha documental si el usuario la indicó en chat.

    Returns:
        ``{"applied": bool, "fecha_es": str|None, "session_patch": dict, "message": str}``
    """
    fecha_es = parse_document_date_override_from_chat(user_message)
    if not fecha_es:
        return {"applied": False, "fecha_es": None, "session_patch": {}, "message": ""}

    cronograma = _merged_cronograma(session_state or {})
    _, deadline_dt = _deadline_from_cronograma(cronograma, "presentacion_proposiciones")
    parsed = parse_spanish_date_fragment(fecha_es)
    clamped = False
    if parsed and deadline_dt and parsed.date() > deadline_dt.date():
        offset = int(getattr(settings, "DOCUMENT_DATE_OFFSET_BUSINESS_DAYS", 2) or 2)
        parsed = _doc_date_from_deadline(deadline_dt, offset)
        fecha_es = _format_fecha_es(parsed)
        clamped = True

    patch = {
        "document_date_user_override": fecha_es,
        "document_date_override_provenance": {
            "source": "chat_user_direct",
            "method": "parse_document_date_override_from_chat",
            "original_message": user_message[:240],
            "clamped_to_deadline": clamped,
        },
    }
    msg = (
        f"Quedó registrada la **fecha documental** **{fecha_es}** para formatos y anexos. "
        "Pulsa **Volver a generar** cuando quieras continuar."
    )
    if clamped:
        msg += (
            " (La fecha indicada era posterior al cierre de proposiciones; "
            f"usé **{fecha_es}**, acotada al calendario de la licitación.)"
        )
    return {
        "applied": True,
        "fecha_es": fecha_es,
        "session_patch": patch,
        "message": msg,
    }


def resolve_document_date(
    session_state: Optional[Dict[str, Any]] = None,
    *,
    user_override: Optional[str] = None,
    hito_key: str = "presentacion_proposiciones",
) -> Dict[str, Any]:
    """
    Devuelve fecha canónica para documentos.

    Returns:
        dict con ``fecha_es``, ``fecha_corta``, ``source``, ``deadline_raw``,
        ``deadline_dt`` (iso o None), ``is_after_deadline``.
    """
    if user_override is None and isinstance(session_state, dict):
        user_override = extract_document_date_user_override(session_state)

    cronograma = _merged_cronograma(session_state or {})
    deadline_raw, deadline_dt = _deadline_from_cronograma(cronograma, hito_key)
    offset = int(getattr(settings, "DOCUMENT_DATE_OFFSET_BUSINESS_DAYS", 2) or 2)

    if user_override and str(user_override).strip():
        parsed = parse_spanish_date_fragment(str(user_override))
        if parsed:
            doc_dt = parsed
            clamped = False
            if deadline_dt and doc_dt.date() > deadline_dt.date():
                doc_dt = _doc_date_from_deadline(deadline_dt, offset)
                clamped = True
            now = datetime.now()
            is_late = bool(deadline_dt and now.date() > deadline_dt.date())
            return {
                "fecha_es": _format_fecha_es(doc_dt),
                "fecha_corta": _format_fecha_corta(doc_dt),
                "source": "user_override_clamped" if clamped else "user_override",
                "deadline_raw": deadline_raw,
                "deadline_dt": deadline_dt.isoformat() if deadline_dt else None,
                "is_after_deadline": is_late,
            }

    if deadline_dt:
        doc_dt = _doc_date_from_deadline(deadline_dt, offset)
        now = datetime.now()
        is_late = now.date() > deadline_dt.date()
        return {
            "fecha_es": _format_fecha_es(doc_dt),
            "fecha_corta": _format_fecha_corta(doc_dt),
            "source": f"cronograma:{hito_key}",
            "deadline_raw": deadline_raw,
            "deadline_dt": deadline_dt.isoformat(),
            "is_after_deadline": is_late,
        }

    now = datetime.now()
    return {
        "fecha_es": _format_fecha_es(now),
        "fecha_corta": _format_fecha_corta(now),
        "source": "system_fallback",
        "deadline_raw": "",
        "deadline_dt": None,
        "is_after_deadline": False,
    }


def normalize_docx_spanish_dates(docx_path: str, canonical_fecha_es: str) -> bool:
    """
    Reemplaza fechas literales en un DOCX materializado por la fecha canónica del expediente.

    Returns:
        True si se modificó el archivo.
    """
    canon = str(canonical_fecha_es or "").strip()
    if not canon or not docx_path or not os.path.exists(docx_path):
        return False
    try:
        from docx import Document
    except ImportError:
        return False

    doc = Document(docx_path)
    changed = False

    def _patch(text: str) -> str:
        nonlocal changed
        new_text = normalize_body_spanish_dates(text, canon)
        if new_text != text:
            changed = True
        return new_text

    for paragraph in doc.paragraphs:
        raw = paragraph.text or ""
        new = _patch(raw)
        if new != raw:
            paragraph.text = new

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    raw = paragraph.text or ""
                    new = _patch(raw)
                    if new != raw:
                        paragraph.text = new

    if changed:
        doc.save(docx_path)
    return changed


def normalize_body_spanish_dates(text: str, canonical_fecha_es: str) -> str:
    """
    Unifica fechas literales en el cuerpo a la fecha canónica del expediente.

    Evita que el LLM inserte la fecha del servidor (p. ej. junio) en TE-01 u otros DOCX.
    """
    canon = str(canonical_fecha_es or "").strip()
    if not canon or not text:
        return text
    canon_key = re.sub(r"\s+", " ", canon.lower())

    def _repl(match: re.Match[str]) -> str:
        frag = re.sub(r"\s+", " ", match.group(0).strip().lower())
        if frag == canon_key:
            return match.group(0)
        return canon

    return _DATE_ES_BODY_REPLACE_RE.sub(_repl, text)


def resolve_addressee_lines(
    session_state: Optional[Dict[str, Any]] = None,
    triage_context: Optional[Dict[str, Any]] = None,
) -> str:
    """Destinatario de cartas: convocante de triage/análisis o genérico."""
    tc = triage_context if isinstance(triage_context, dict) else {}
    dest = str(tc.get("destinatario") or "").strip()
    if dest:
        return dest

    if isinstance(session_state, dict):
        dest = str(session_state.get("destinatario") or "").strip()
        if dest:
            return dest
        for key in ("last_analysis", "analysis_snapshot"):
            block = session_state.get(key)
            if isinstance(block, dict):
                dest = str(block.get("destinatario") or "").strip()
                if dest:
                    return dest

    from app.services.convocante_resolver import extract_convocante_from_text

    blob_parts: list[str] = []
    if isinstance(session_state, dict):
        blob_parts.append(str(session_state.get("bases_corpus_hint") or ""))
        for key in ("last_analysis", "analysis_snapshot"):
            block = session_state.get(key)
            if isinstance(block, dict):
                for k in ("convocante", "autoridad_convocante", "entidad", "dependencia", "alcance_operativo"):
                    blob_parts.append(str(block.get(k) or ""))
    if blob_parts:
        found = extract_convocante_from_text("\n".join(blob_parts))
        if found.get("destinatario"):
            return found["destinatario"]

    tc = triage_context if isinstance(triage_context, dict) else {}
    conv = str(tc.get("convocante") or tc.get("autoridad_convocante") or "").strip()
    if not conv and isinstance(session_state, dict):
        for key in ("last_analysis", "analysis_snapshot"):
            block = session_state.get(key)
            if isinstance(block, dict):
                conv = str(
                    block.get("convocante")
                    or block.get("autoridad_convocante")
                    or block.get("entidad")
                    or ""
                ).strip()
                if conv:
                    break
    if conv:
        return f"{conv.upper()}\nPRESENTE.-"
    return "A QUIEN CORRESPONDA:"
