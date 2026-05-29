"""
Espejo universal de plantillas: copia/rellena archivos Office ingestados para entrega.

No hardcodea anexos por licitación; opera sobre rutas resueltas por sesión.
"""
from __future__ import annotations

import glob
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from app.core.observability import get_logger
from app.services.ingested_file_resolver import IngestedFileRef
from app.services.document_traceability import safe_file_sha256

logger = get_logger(__name__)

_OFFICE_EXTS = frozenset({".doc", ".docx", ".xls", ".xlsx"})


def _ext(path: str) -> str:
    return Path(path).suffix.lower()


def convert_doc_to_docx(source_path: str, out_dir: str) -> Optional[str]:
    """
    Convierte .doc binario a .docx vía LibreOffice (misma estrategia que ingesta).

    Returns:
        Ruta al .docx generado o None si falla.
    """
    os.makedirs(out_dir, exist_ok=True)
    for cmd in (
        [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "docx",
            "--outdir",
            out_dir,
            source_path,
        ],
        [
            "libreoffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "docx",
            "--outdir",
            out_dir,
            source_path,
        ],
    ):
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            if proc.returncode != 0:
                continue
            base = Path(source_path).stem
            candidates = glob.glob(os.path.join(out_dir, f"{base}*.docx"))
            if candidates:
                return candidates[0]
        except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
            logger.warning("template_mirror_lo_convert_failed", error=str(exc))
            continue
    return None


def fill_docx_with_profile(docx_path: str, profile: Dict[str, Any]) -> None:
    """
    Sustitución determinista de marcadores habituales (RFC, razón social, fecha).

    No sustituye líneas de guiones largas (firmas); solo tokens explícitos.
    """
    try:
        from docx import Document
    except ImportError:
        return

    rfc = str(profile.get("rfc") or profile.get("RFC") or "").strip()
    razon = str(
        profile.get("razon_social")
        or profile.get("company_name")
        or profile.get("nombre_empresa")
        or ""
    ).strip()
    rep = str(
        profile.get("representante_legal")
        or profile.get("representante")
        or profile.get("legal_representative")
        or ""
    ).strip()
    domicilio = str(profile.get("domicilio") or profile.get("address") or "").strip()
    fecha = str(profile.get("fecha") or profile.get("fecha_documento") or "").strip()
    zonas = str(profile.get("zonas_ofertadas") or profile.get("zonas") or "").strip()
    numero_ref = str(
        profile.get("numero_referencia")
        or profile.get("licitacion")
        or profile.get("numero_procedimiento")
        or ""
    ).strip()
    tarifa_ref = profile.get("tarifa_mensual_referencia")
    tarifa_ref_txt = ""
    if tarifa_ref not in (None, ""):
        try:
            tarifa_ref_txt = f"${float(tarifa_ref):,.2f} MXN"
        except Exception:
            tarifa_ref_txt = str(tarifa_ref).strip()
    licitacion = str(
        profile.get("licitacion")
        or profile.get("numero_procedimiento")
        or profile.get("tender_name")
        or ""
    ).strip()

    replacements: List[Tuple[str, str]] = []
    if rfc:
        replacements.extend(
            [
                ("{{RFC}}", rfc),
                ("[RFC]", rfc),
                ("RFC: _____", f"RFC: {rfc}"),
                ("R.F.C.", rfc),
            ]
        )
    if razon:
        replacements.extend(
            [
                ("{{RAZON_SOCIAL}}", razon),
                ("[RAZON SOCIAL]", razon),
                ("[NOMBRE DE LA EMPRESA]", razon),
            ]
        )
    if rep:
        replacements.extend(
            [
                ("{{REPRESENTANTE}}", rep),
                ("[REPRESENTANTE LEGAL]", rep),
            ]
        )
    if domicilio:
        replacements.append(("[DOMICILIO]", domicilio))
    if fecha:
        replacements.append(("[FECHA]", fecha))

    if not replacements and not any((tarifa_ref_txt, numero_ref, zonas, licitacion, fecha, domicilio, razon, rep, rfc)):
        return

    doc = Document(docx_path)

    def _apply_paragraph(text: str) -> str:
        out = text
        for old, new in replacements:
            if old in out and new:
                out = out.replace(old, new)
        if razon:
            out = re.sub(r"(?im)^(NOMBRE DEL LICITANTE:\s*)$", rf"\1{razon}", out)
            out = re.sub(
                r"(?i)(NOMBRE DEL LICITANTE:\s*)(?=FECHA:)",
                rf"\1{razon} ",
                out,
            )
            out = re.sub(r"(?i)(denominaci[óo]n\s+social\s*:\s*)_{3,}", rf"\1{razon}", out)
            out = re.sub(r"(?i)(empresa)\s*_{3,}\(\d+\)_{3,}", rf"\1 {razon}", out)
            out = out.replace("(NOMBRE DE LA PERSONA QUE PARTICIPA EN LA PRESENTE LICITACIÓN)", razon)
            out = out.replace("(NOMBRE DE LA PERSONA QUE PARTICIPA EN LA PRESENTE LICITACION)", razon)
        if rep:
            out = re.sub(r"(?i)(quien\s+suscribe)\s*_{3,}", rf"\1 {rep}", out)
            out = re.sub(
                r"(?i)(quien\s+suscribe)\s*_{3,}(\s*,)",
                rf"\1 {rep}\2",
                out,
            )
            out = out.replace("(NOMBRE DEL REPRESENTANTE LEGAL)", rep)
            out = out.replace("NOMBRE DEL REPRESENTANTE LEGAL", rep)
        if fecha:
            out = re.sub(r"(?im)^(FECHA:\s*)$", rf"\1{fecha}", out)
            out = re.sub(r"_{3,}\s*de\s*_{3,}\s*de\s*_{3,}(?:\s*\(\d+\))?", fecha, out)
        if rfc:
            out = re.sub(r"(?im)^(RFC:\s*)$", rf"\1{rfc}", out)
        if rep and out.strip().upper() == "N/A":
            out = rep
        if domicilio:
            out = re.sub(r"(?i)(domicilio\s*:\s*)_{3,}", rf"\1{domicilio}", out)
        if zonas:
            out = re.sub(r"(?i)(zona\s*:\s*)_{3,}", rf"\1{zonas}", out)
        if tarifa_ref_txt:
            out = re.sub(
                r"(?i)(tarifa\s+mensual\s+para\s+horario\s*:\s*)_{3,}",
                rf"\1{tarifa_ref_txt}",
                out,
            )
        if numero_ref:
            out = re.sub(
                r"(?i)(n[uú]mero\s*\(s\)\s*:\s*)_{3,}",
                rf"\1{numero_ref}",
                out,
            )
        if licitacion:
            out = re.sub(r"(?i)(procedimiento)\s*_{3,}", rf"\1 {licitacion}", out)
            out = re.sub(r"(?i)(No\.\s*)_{3,}\(\d+\)_{3,}", rf"\1{licitacion}", out)
            out = re.sub(r"(?i)(procedimiento\s+[A-Z0-9-]+)\(\d+\)_{3,}", rf"\1", out)
        return out

    for para in doc.paragraphs:
        if para.text:
            new_t = _apply_paragraph(para.text)
            if new_t != para.text:
                para.text = new_t

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        new_t = _apply_paragraph(para.text)
                        if new_t != para.text:
                            para.text = new_t

    doc.save(docx_path)


def mirror_template_to_output(
    source_ref: IngestedFileRef,
    output_path: str,
    profile: Optional[Dict[str, Any]] = None,
    *,
    fill_profile: bool = True,
) -> Dict[str, Any]:
    """
    Materializa una plantilla oficial en la carpeta de salida del agente.

    Args:
        source_ref: Archivo ingestado con ruta en disco.
        output_path: Ruta destino (.docx o .xlsx/.xls).
        profile: Datos maestro empresa para relleno ligero.
        fill_profile: Si True, aplica sustituciones en docx.

    Returns:
        Metadatos del artefacto generado (ruta, modo, extensión).

    Raises:
        FileNotFoundError: Si no existe el origen.
        ValueError: Extensión no soportada.
    """
    src = source_ref.file_path
    if not os.path.isfile(src):
        raise FileNotFoundError(src)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    ext = _ext(src)
    out_ext = _ext(output_path)
    profile = profile or {}
    mode = "copy"

    if ext == ".doc":
        if out_ext != ".docx":
            output_path = str(Path(output_path).with_suffix(".docx"))
            out_ext = ".docx"
        tmp = tempfile.mkdtemp(prefix="licitai_mirror_")
        try:
            converted = convert_doc_to_docx(src, tmp)
            if not converted:
                shutil.copy2(src, output_path.replace(".docx", ".doc"))
                mode = "copy_doc_native"
                return {
                    "ruta": output_path.replace(".docx", ".doc"),
                    "mirror_mode": mode,
                    "materialization_route": "mirror",
                    "source_filename": source_ref.filename,
                    "source_doc_id": source_ref.doc_id,
                    "source_path": src,
                    "source_hash": safe_file_sha256(src),
                    "output_hash": safe_file_sha256(output_path.replace(".docx", ".doc")),
                }
            shutil.copy2(converted, output_path)
            mode = "convert_doc_docx"
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    elif ext == ".docx":
        shutil.copy2(src, output_path)
        mode = "copy_docx"
    elif ext in (".xlsx", ".xls"):
        shutil.copy2(src, output_path)
        mode = "copy_excel"
    else:
        raise ValueError(f"Extensión no soportada para espejo: {ext}")

    if fill_profile and out_ext == ".docx" and os.path.isfile(output_path):
        try:
            fill_docx_with_profile(output_path, profile)
            mode = f"{mode}_filled"
        except Exception as exc:
            logger.warning(
                "template_mirror_fill_failed",
                path=output_path,
                error=str(exc),
            )

    return {
        "ruta": output_path,
        "mirror_mode": mode,
        "materialization_route": "mirror",
        "source_filename": source_ref.filename,
        "source_doc_id": source_ref.doc_id,
        "source_path": src,
        "source_hash": safe_file_sha256(src),
        "output_hash": safe_file_sha256(output_path),
    }
