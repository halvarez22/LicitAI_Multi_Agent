"""
Auditoría universal de contenido en entrega CompraNet (sin hardcode por licitación).

Produce matriz JSON por archivo del índice de entrega con hallazgos tipificados.
"""
from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from app.services.cronograma_bases_extract import parse_spanish_date_fragment
from app.services.document_contamination_gate import (
    ContaminationHit,
    contamination_enforce_at_pack,
    contamination_hits_to_issues,
    infer_document_stage,
    scan_all_document_dates,
    scan_conflicting_document_dates,
    scan_text_contamination,
)
from app.services.document_date_resolver import resolve_document_date
from app.services.document_body_quality import (
    TEMPLATE_CONTAMINATION_RELAXED_ERROR_TYPES,
    should_relax_delivery_contamination,
)
from app.services.pliego_formats_enrichment_service import pliego_format_dedupe_key


@dataclass
class AuditFinding:
    error_type: str
    severity: str
    field_key: str
    detected_value: str
    expected_rule: str


@dataclass
class DocumentAuditRow:
    nombre_entrega: str
    sobre: str
    ruta: str
    dedupe_key: str
    findings: List[AuditFinding] = field(default_factory=list)
    fecha_detectadas: List[str] = field(default_factory=list)
    char_count: int = 0
    paragraph_count: int = 0

    @property
    def blocking_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "block")

    @property
    def warn_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "warn")


def _read_docx_text(path: Path) -> tuple[str, int]:
    from docx import Document

    doc = Document(str(path))
    paras = [re.sub(r"\s+", " ", (p.text or "").strip()) for p in doc.paragraphs if (p.text or "").strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = re.sub(r"\s+", " ", (cell.text or "").strip())
                if t:
                    paras.append(t)
    text = "\n".join(paras)
    return text, len(paras)


def _read_xlsx_text(path: Path) -> tuple[str, int]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", 0
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(max_row=min(ws.max_row or 0, 200), values_only=True):
            for v in row:
                if v is not None and str(v).strip():
                    parts.append(str(v).strip())
    text = "\n".join(parts)
    return text, len(parts)


def _load_file_text(path: Path) -> tuple[str, int]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _read_docx_text(path)
    if ext == ".xlsx":
        return _read_xlsx_text(path)
    return "", 0


def audit_delivery_files(
    session_id: str,
    *,
    session_state: Optional[Dict[str, Any]] = None,
    indice_files: Optional[Sequence[Dict[str, Any]]] = None,
    validated_root: Optional[Path] = None,
) -> Dict[str, Any]:
    """
    Audita archivos materializados referenciados en ``INDICE_ENTREGA.json``.

    Args:
        session_id: Identificador de sesión.
        session_state: Estado de sesión para resolver fecha límite.
        indice_files: Entradas del índice; si None, se leen de disco.
        validated_root: Raíz ``_compranet_validated``; default bajo outputs.

    Returns:
        dict serializable con resumen y filas por documento.
    """
    root = validated_root or Path("/data/outputs") / session_id / "_compranet_validated"
    if indice_files is None:
        indice_path = root / "INDICE_ENTREGA.json"
        if indice_path.is_file():
            payload = json.loads(indice_path.read_text(encoding="utf-8"))
            indice_files = list(payload.get("files") or [])
        else:
            indice_files = []

    date_info = resolve_document_date(session_state or {})
    deadline_iso = str(date_info.get("deadline_dt") or "")
    fecha_es = str(date_info.get("fecha_es") or "")

    rows: List[DocumentAuditRow] = []
    for ent in indice_files or []:
        if not isinstance(ent, dict):
            continue
        nombre = str(ent.get("nombre_entrega") or ent.get("dest_name") or "").strip()
        sobre = str(ent.get("sobre") or ent.get("label") or "").strip()
        rel = str(ent.get("ruta_relativa") or ent.get("path") or "").strip()
        if rel:
            fpath = root / rel.replace("\\", "/")
        elif nombre and sobre:
            folder = {
                "SobreComplementaria": "SobreComplementaria",
                "SobreTecnica": "SobreTecnica",
                "SobreEconomica": "SobreEconomica",
            }.get(sobre, sobre)
            fpath = root / folder / nombre
        else:
            continue
        if not fpath.is_file():
            rows.append(
                DocumentAuditRow(
                    nombre_entrega=nombre,
                    sobre=sobre,
                    ruta=str(fpath),
                    dedupe_key=pliego_format_dedupe_key(nombre),
                    findings=[
                        AuditFinding(
                            "deliverable_file_missing",
                            "block",
                            "path",
                            str(fpath),
                            "file_exists_on_disk",
                        )
                    ],
                )
            )
            continue

        text, para_count = _load_file_text(fpath)
        dedupe = pliego_format_dedupe_key(nombre)
        stage = infer_document_stage(sobre=sobre, basename=fpath.name, dedupe_key=dedupe)
        findings: List[AuditFinding] = []

        relax_template = should_relax_delivery_contamination(fpath.name)

        def _append_hit(hit: ContaminationHit, *, default_severity: str = "block") -> None:
            if relax_template and hit.error_type in TEMPLATE_CONTAMINATION_RELAXED_ERROR_TYPES:
                default_severity = "warn"
            issues = contamination_hits_to_issues(
                [hit],
                document_id=fpath.name,
                provenance={"source": "delivery_content_audit", "confidence": 1.0},
            )
            if issues:
                sev = str(issues[0].get("severity") or default_severity)
            else:
                sev = default_severity
            if relax_template and hit.error_type in TEMPLATE_CONTAMINATION_RELAXED_ERROR_TYPES:
                sev = "warn"
            findings.append(
                AuditFinding(
                    hit.error_type,
                    sev,
                    hit.field_key,
                    hit.detected_value,
                    hit.expected_rule,
                )
            )

        hits = scan_text_contamination(
            text, basename=fpath.name, stage=stage, dedupe_key=dedupe
        )
        if relax_template:
            hits = [h for h in hits if h.error_type not in TEMPLATE_CONTAMINATION_RELAXED_ERROR_TYPES]
        for hit in hits:
            _append_hit(hit)
        if not relax_template:
            for date_hit in scan_all_document_dates(
                text,
                deadline_dt_iso=deadline_iso or None,
                canonical_fecha_es=fecha_es,
            ):
                _append_hit(date_hit)
            conflict = scan_conflicting_document_dates(
                text,
                canonical_fecha_es=fecha_es,
                dedupe_key=dedupe,
                basename=fpath.name,
            )
            if conflict:
                _append_hit(conflict)
        if "bajo protesta" not in text.lower() and fpath.suffix.lower() == ".docx":
            if dedupe.startswith("pliego|ANEXO_") or "carta" in nombre.lower():
                findings.append(
                    AuditFinding(
                        "legal_protest_missing",
                        "warn",
                        "content",
                        "",
                        "administrative_letter_requires_protesta",
                    )
                )
        date_frags = re.findall(
            r"\b\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|"
            r"septiembre|octubre|noviembre|diciembre)\s+de\s+\d{4}\b",
            text,
            re.I,
        )
        rows.append(
            DocumentAuditRow(
                nombre_entrega=nombre,
                sobre=sobre,
                ruta=str(fpath),
                dedupe_key=dedupe,
                findings=findings,
                fecha_detectadas=sorted(set(d.lower() for d in date_frags)),
                char_count=len(text),
                paragraph_count=para_count,
            )
        )

    blocking = sum(r.blocking_count for r in rows)
    warns = sum(r.warn_count for r in rows)
    return {
        "session_id": session_id,
        "schema_version": "1.0.0",
        "date_resolution": date_info,
        "summary": {
            "files_audited": len(rows),
            "blocking_findings": blocking,
            "warn_findings": warns,
            "files_with_blocking": sum(1 for r in rows if r.blocking_count),
        },
        "documents": [
            {
                **{k: v for k, v in asdict(r).items() if k != "findings"},
                "findings": [asdict(f) for f in r.findings],
            }
            for r in rows
        ],
    }


def forensic_contamination_blocking(report: Dict[str, Any]) -> bool:
    """True si el informe forense tiene hallazgos bloqueantes."""
    summary = report.get("summary") or {}
    return int(summary.get("blocking_findings") or 0) > 0


def format_forensic_contamination_errors(report: Dict[str, Any], *, limit: int = 12) -> List[str]:
    """Resume hallazgos bloqueantes para logs y errores de empaquetado."""
    errors: List[str] = []
    for doc in report.get("documents") or []:
        nombre = str(doc.get("nombre_entrega") or doc.get("dedupe_key") or "documento")
        for finding in doc.get("findings") or []:
            if str(finding.get("severity") or "").lower() != "block":
                continue
            errors.append(
                f"{nombre}: {finding.get('error_type')} ({str(finding.get('detected_value') or '')[:80]})"
            )
            if len(errors) >= limit:
                return errors
    return errors


def run_forensic_contamination_audit(
    session_id: str,
    *,
    session_state: Optional[Dict[str, Any]] = None,
    validated_root: Optional[Path] = None,
    indice_files: Optional[Sequence[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    """
    Auditoría forense P0 sobre entrega CompraNet (misma matriz que CI).

    Returns:
        Informe JSON con ``summary.blocking_findings`` y filas por archivo.
    """
    report = audit_delivery_files(
        session_id,
        session_state=session_state,
        indice_files=indice_files,
        validated_root=validated_root,
    )
    report["enforce_at_pack"] = contamination_enforce_at_pack()
    report["gate_passed"] = not forensic_contamination_blocking(report)
    return report
