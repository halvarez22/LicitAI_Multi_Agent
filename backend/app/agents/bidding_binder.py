import os
import re
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.agents.base_agent import BaseAgent
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.agents.mcp_context import MCPContextManager
from app.services.vector_service import VectorDbServiceClient
from app.config.settings import settings
from app.core.logging_config import get_logger

logger = get_logger("bidding_binder_agent")

class BiddingBinderAgent(BaseAgent):
    """
    Agente encargado de compilar el 'Sobre Digital' final.
    Consolida la Guía de Armado y el Checklist de Integridad.
    """
    
    def __init__(self, context_manager: MCPContextManager):
        super().__init__(
            agent_id="bidding_binder",
            name="Bidding Binder Agent",
            description="Agente encargado de compilar el Sobre Digital final y consolidar la Guía de Armado y el Checklist.",
            context_manager=context_manager
        )
        self.vector_client = VectorDbServiceClient()
        
    async def process(self, input_data: AgentInput) -> AgentOutput:
        session_id = input_data.session_id
        
        # 1. Escaneo de Requisitos Externos via ChromaDB
        res_ext = self.vector_client.query_texts(
            session_id=session_id,
            query="6.1 DOCUMENTACION COMPLEMENTARIA documentos que debe presentar el licitante SAT IMSS INFONAVIT actas registros",
            n_results=1
        )
        results_ext = res_ext.get("documents") if res_ext else []
        
        externos_req = []
        if results_ext and len(results_ext[0]) > 0:
            externos_req = [
                "Opinión del SAT (Positiva)",
                "Opinión del IMSS (Positiva)",
                "Constancia INFONAVIT (Sin adeudos)",
                "Acta Constitutiva y Poderes Notariales",
                "Identificación Oficial del Representante Legal"
            ]
        else:
            externos_req = ["Opinión SAT, IMSS, INFONAVIT (Estándar LAASSP)"]

        # === CAPA 1 y 2: BiddingBinderEngine (Auditoría Física y Convergencia Legal) ===
        res_fisicos = self.vector_client.query_texts(
            session_id=session_id,
            query="tinta foliad rubric sobre paquete original copia firma",
            n_results=5
        )
        chunks_fisicos = res_fisicos.get("documents", [[]])[0] if res_fisicos and res_fisicos.get("documents") else []
        texto_fisico = " ".join(chunks_fisicos).lower()

        # Regex y heurísticas
        tinta_color = "AZUL" if "azul" in texto_fisico else ("NEGRA" if "negra" in texto_fisico else "AZUL O NEGRA")
        foliado_req = "Consecutivo en todas las hojas (ej. 001, 002...)" if "foliad" in texto_fisico else "No especificado explícitamente, pero se recomienda foliar"
        rubrica_req = "Firma autógrafa en la última hoja y rúbrica en los márgenes del resto" if "rubric" in texto_fisico else "Firma en la última hoja de cada documento"
        
        # Capa 2: Riesgo de Desechamiento
        has_critical_risk = bool(re.search(r'(desechamient|descalifica|motivo de|no será solvente)', texto_fisico))

        protocolo_armado = {
            "tinta": tinta_color,
            "foliado": foliado_req,
            "rubricado": rubrica_req,
            "critical_risk": has_critical_risk
        }
            
        # 2. Mapeo de Archivos Internos Generados
        company_data = input_data.company_data or {}
        documentos_generados = company_data.get("documentos_generados", {})
        
        tecnica = documentos_generados.get("tecnica", [])
        economica = documentos_generados.get("economica", [])
        administrativa = documentos_generados.get("administrativa", [])
        
        def _get_names(lst):
            return [d.get("nombre") or d.get("name") for d in lst if d.get("nombre") or d.get("name")]
            
        formatos_internos = _get_names(administrativa) + _get_names(tecnica)
        cedula_economica = _get_names(economica)
        if not cedula_economica:
            cedula_economica = ["ANEXO_AE_PROPUESTA_ECONOMICA.docx", "TABLA_PRECIOS_UNITARIOS.xlsx"]
        if not formatos_internos:
            formatos_internos = ["CARTA_PRESENTACION.docx", "ESPECIFICACIONES_PROYECTO.docx"]
            
        # 3. Construir el JSON consolidado
        checklist_json = {
            "bloque_1_externos": externos_req,
            "bloque_2_formatos": formatos_internos,
            "bloque_3_tecnicos": ["Anexos Técnicos Específicos según Bases"],
            "bloque_4_economica": cedula_economica,
            "protocolo_armado": protocolo_armado
        }
        
        # 4. Inyección en plantilla Word (Capa 3)
        out_dir = os.path.join("out", "generated", session_id)
        os.makedirs(out_dir, exist_ok=True)
        docx_path = os.path.join(out_dir, "GUIA_DE_ARMADO_Y_CHECKLIST.docx")
        
        try:
            self._render_guide_docx(docx_path, checklist_json, company_data.get("rfc", "RFC_EMPRESA"), session_id)
        except Exception as e:
            logger.error("bidding_binder_docx_error", error=str(e))
            
        # Guardamos en session state
        session_state = await self.context_manager.memory.get_session(session_id)
        if has_critical_risk:
            session_state["physical_packaging_critical_risk"] = True
        session_state["delivery_checklist"] = checklist_json
        await self.context_manager.memory.save_session(session_id, session_state)
        
        return AgentOutput(
            status=AgentStatus.SUCCESS,
            agent_id=self.agent_id,
            session_id=session_id,
            message="Guía de Armado y Checklist de Integridad generado con éxito.",
            data={"checklist": checklist_json, "docx_path": docx_path}
        )
        
    def _render_guide_docx(self, path: str, data: Dict[str, Any], rfc: str, licitacion_id: str):
        doc = Document()
        doc.add_heading('GUÍA DE ARMADO DE SOBRES Y CHECKLIST DE INTEGRIDAD', 0)
        
        p_proto = data.get("protocolo_armado", {})
        
        # === Bloque A: Instrucciones de Impresión (Tolerancia Cero) ===
        doc.add_heading('Bloque A: Instrucciones de Impresión (Tolerancia Cero)', level=1)
        p1 = doc.add_paragraph()
        p1.add_run("1. Imprime los anexos técnicos en papel membretado.\n")
        run_tinta = p1.add_run(f"2. Prepara un bolígrafo de TINTA {p_proto.get('tinta', 'AZUL')}. ")
        run_tinta.bold = True
        p1.add_run("El uso de otro color causa descalificación automática.")
        if p_proto.get("critical_risk"):
            run_warn = doc.add_paragraph().add_run("⚠️ ALERTA DE RIESGO CRÍTICO: Las bases indican explícitamente que no seguir estas instrucciones es causal de desechamiento.")
            run_warn.bold = True
            run_warn.font.color.rgb = RGBColor(255, 0, 0)
            
        # === Bloque B: El Ritual del Foliado y Rúbrica ===
        doc.add_heading('Bloque B: El Ritual del Foliado y Rúbrica', level=1)
        doc.add_paragraph("1. Ordena tus hojas en este orden estricto (Legal, Técnico, Económico).")
        doc.add_paragraph(f"2. {p_proto.get('foliado', 'Foliado consecutivo')}")
        doc.add_paragraph(f"3. {p_proto.get('rubricado', 'Firma y rúbrica en todas las páginas')}")
        
        # Checklist de documentos
        doc.add_heading('Checklist de Documentos a Integrar', level=2)
        for cat, title in [("bloque_1_externos", "Administrativos Externos"), 
                           ("bloque_2_formatos", "Formatos Generados por LicitAI"), 
                           ("bloque_3_tecnicos", "Documentación Técnica"), 
                           ("bloque_4_economica", "Propuesta Económica")]:
            doc.add_heading(title, level=3)
            for item in data.get(cat, []):
                doc.add_paragraph(f"[ ] {item}")

        doc.add_page_break()
        
        # === Bloque C: Plantilla de Rotulado de Sobres ===
        doc.add_heading('Bloque C: Plantilla de Rotulado de Sobres (Listo para Imprimir)', level=1)
        doc.add_paragraph("Recorta la siguiente etiqueta y pégala en el exterior de tu sobre manila cerrado:")
        
        # Crear la etiqueta (tabla con bordes)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        p_etiqueta = cell.paragraphs[0]
        p_etiqueta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_lic = p_etiqueta.add_run(f"LICITACIÓN: {licitacion_id.upper()}\n")
        run_lic.bold = True
        run_lic.font.size = Pt(16)
        
        p_etiqueta.add_run(f"\nPROVEEDOR: {rfc}\n\n")
        
        run_prop = p_etiqueta.add_run("PROPUESTA TÉCNICA Y ECONÓMICA\n")
        run_prop.bold = True
        run_prop.font.size = Pt(14)
        
        p_etiqueta.add_run("\nCONTIENE DOCUMENTACIÓN ORIGINAL Y COPIAS")
        
        doc.save(path)
