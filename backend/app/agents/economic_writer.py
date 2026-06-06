import os
import re
import time
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from app.agents.base_agent import BaseAgent
from app.agents.mcp_context import MCPContextManager
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.services.document_fill_quality_gate import (
    detect_cross_tender_marker,
    validate_generated_documents_fill,
)
from app.services.document_fill_ux_messages import build_fill_validation_event
from app.services.excel_filling_service import ExcelFillingService
from app.services.structured_economic_price_mapper import apply_structured_price_inputs
from app.services.document_traceability import (
    attach_traceability,
    build_materialization_metrics,
    safe_file_sha256,
)
from app.config.settings import settings
from app.services.document_date_resolver import resolve_document_date
from app.services.apu_document_builder import build_apu_markdown
from app.services.economic_document_reapply import build_economic_doc_metadata
from app.utils.doc_formatting import (
    apply_corporate_docx_letterhead,
    apply_corporate_excel_letterhead,
    stamp_corporate_excel_file,
)


def _has_valid_excel_locator(item: Dict[str, Any]) -> bool:
    extra = item.get("extra") if isinstance(item.get("extra"), dict) else {}
    row_idx = item.get("row_index")
    col_idx = extra.get("price_column_index")
    try:
        return int(float(row_idx)) >= 0 and int(float(col_idx)) >= 0
    except (TypeError, ValueError):
        return False


class EconomicWriterAgent(BaseAgent):
    """
    Agente: Generador de Propuesta Económica.
    Genera documentos económicos formales (.xlsx y .docx) a partir de los
    ítems ya calculados por EconomicAgent en Fase 1, sin invocar al LLM.

    POLÍTICA DE TOTALES (decisión de negocio):
    Los documentos oficiales aplican IVA 16 % sobre el subtotal de líneas
    renderizadas. EconomicAgent (Fase 1) puede incluir un margen adicional
    (~15 %) para uso interno de la UI; ese margen NO se traslada al sobre
    físico. Si se desea una única cifra en pantalla y en papel, alinear
    ambos agentes bajo la misma regla fiscal.
    """

    def __init__(self, context_manager: MCPContextManager):
        super().__init__(
            agent_id="economic_writer",
            name="Economic Writer Agent",
            description="Generador automatizado de propuestas económicas y catálogos de precios.",
            context_manager=context_manager
        )
        self.excel_filler = ExcelFillingService()

    async def process(self, agent_input: AgentInput) -> AgentOutput:
        session_id = agent_input.session_id
        correlation_id = agent_input.correlation_id or "no-id"
        started_at = time.perf_counter()
        print(f"[{self.name}] 💰 Iniciando generación de Propuesta Económica para {session_id}...", flush=True)

        context = await self.context_manager.get_global_context(session_id)

        # 1. Recuperar contexto de la empresa
        company_data = agent_input.company_data
        master_profile = company_data.get("master_profile", {})
        
        # 2. Obtener propuesta de Fase 1
        # a) Inyección directa del orquestador via economic_data
        economic_data = agent_input.company_data.get("economic_data")
        
        # b) Estructura estándar results.economic.data
        if not economic_data and "results" in agent_input.company_data:
            res = agent_input.company_data["results"]
            if isinstance(res, dict) and "economic" in res:
                econ = res["economic"]
                economic_data = econ.get("data", econ) if isinstance(econ, dict) else None

        # c) Buscar en el estado de la sesión si venimos en modo generation_only
        session_state = context.get("session_state") or {}
        try:
            from app.services.mini_dictamen_anexos_service import (
                build_and_persist_mini_dictamen,
                build_stage_blocking_questions,
                get_blocking_annex_rows_for_stage,
            )

            await build_and_persist_mini_dictamen(self.context_manager.memory, session_id)
            fresh_state = await self.context_manager.memory.get_session(session_id) or session_state
            blocking_rows = get_blocking_annex_rows_for_stage(fresh_state, "economic_writer")
            if blocking_rows:
                fresh_state["pending_questions"] = build_stage_blocking_questions(
                    "economic_writer", blocking_rows
                ) + list(fresh_state.get("pending_questions") or [])
                fresh_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, fresh_state)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "La generación económica quedó bloqueada por anexos obligatorios "
                        "pendientes de aclaración o sin fuente oficial válida."
                    ),
                    data={"missing": blocking_rows},
                    correlation_id=correlation_id,
                )
        except Exception:
            pass
        if not economic_data:
            tasks = session_state.get("tasks_completed", [])
            for task in reversed(tasks):
                if task.get("task") == "economic_proposal":
                    result_data = task.get("result", {})
                    # extraemos .data del dict o lo asumimos directo
                    economic_data = result_data.get("data", result_data)
                    break

        # d) MPS (master_proposal_state) — fuente tras refresh_economic_validations
        if not economic_data:
            mps = session_state.get("master_proposal_state")
            if isinstance(mps, dict) and mps:
                economic_data = mps

        if not economic_data:
            return AgentOutput(
                status=AgentStatus.ERROR,
                agent_id=self.agent_id,
                session_id=session_id,
                message="No se encontró una propuesta económica calculada en Fase 1.",
                error="No se encontró una propuesta económica calculada en Fase 1.",
                correlation_id=correlation_id,
            )

        items = economic_data.get("items") if isinstance(economic_data.get("items"), list) else []
        if not items:
            user_inputs = session_state.get("economic_user_inputs") or {}
            if bool(user_inputs.get("allow_zero_total_base_ack")):
                economic_data = {**economic_data, "allow_zero_total_base_ack": True}
            else:
                try:
                    line_items = await self.context_manager.memory.get_line_items_for_session(
                        session_id
                    )
                except Exception:
                    line_items = []
                concept_prices = {}
                if isinstance(user_inputs, dict):
                    cp = user_inputs.get("concept_prices")
                    if isinstance(cp, dict):
                        concept_prices = cp
                if line_items and concept_prices:
                    line_items = apply_structured_price_inputs(line_items, concept_prices)
                if line_items:
                    boot: List[Dict[str, Any]] = []
                    for idx, li in enumerate(line_items[:300]):
                        if not isinstance(li, dict):
                            continue
                        boot.append(
                            {
                                "partida": idx + 1,
                                "concepto": li.get("description")
                                or li.get("concepto")
                                or li.get("descripcion")
                                or "Partida",
                                "cantidad": float(li.get("quantity") or li.get("cantidad") or 1),
                                "precio_unitario": float(
                                    li.get("unit_price") or li.get("precio_unitario") or 0
                                ),
                            }
                        )
                    if boot:
                        economic_data = {**economic_data, "items": boot}
                        items = boot
        if not items:
            st_payload = str(economic_data.get("status") or "").strip().lower()
            vres = economic_data.get("validation_result")
            vres = vres if isinstance(vres, dict) else {}
            blocking = bool(vres.get("blocking_issues"))
            missing = economic_data.get("missing")
            has_missing = isinstance(missing, list) and len(missing) > 0
            if st_payload == "waiting_for_data" and (blocking or has_missing):
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "La propuesta económica aún no está lista para generar documentos: "
                        "hay validaciones pendientes o datos incompletos. "
                        "Atiende el asistente o ajusta Excel/cotización y vuelve a ejecutar generar o continuar."
                    ),
                    correlation_id=correlation_id,
                )
            return AgentOutput(
                status=AgentStatus.ERROR,
                agent_id=self.agent_id,
                session_id=session_id,
                message=(
                    "La propuesta económica no tiene partidas cotizables para materializar archivos. "
                    "Escribe `generar propuesta económica` en el chat para recalcular la cotización."
                ),
                error="No se encontró una propuesta económica calculada en Fase 1.",
                correlation_id=correlation_id,
            )
            
        # 3. Normalizar items para el renderizado Excel/Word
        mapeo_items = []
        for idx, item in enumerate(economic_data.get("items", [])):
            cantidad = float(item.get("cantidad", 1))
            precio = float(item.get("precio_unitario", 0.0))
            importe = item.get("subtotal", cantidad * precio)
            mapeo_items.append({
                "partida": item.get("partida", idx + 1),
                "descripcion": item.get("concepto", item.get("descripcion", "")),
                "unidad": item.get("unidad", "Servicio"),
                "cantidad": cantidad,
                "precio_unitario": precio,
                "importe": importe
            })
            
        # Fallback de sumatoria si el motor no proporcionó totales
        calc_subtotal = sum(i["importe"] for i in mapeo_items)
        
        # 1. Obtener totales maestros desde el motor económico (Prioridad absoluta)
        subtotal = float(economic_data.get("total_base") or calc_subtotal)
        total = float(economic_data.get("grand_total") or (subtotal * 1.16))
        
        # 2. Re-calcular IVA como la diferencia si no viene explícito (Maneja IVA 0%, 8%, 16%, etc.)
        iva = round(total - subtotal, 2)
        
        subtotal = round(subtotal, 2)
        total = round(total, 2)
        
        allow_zero = bool(economic_data.get("allow_zero_total_base_ack"))
        if subtotal < 0.01 and not allow_zero:
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=(
                    "La propuesta económica no tiene importe cotizable antes de IVA (subtotal ~0). "
                    "Captura precios unitarios en la fase económica.\n\n"
                    "Si esta licitación no requiere importe base, escribe en el chat: "
                    "**'Esta licitación no requiere importe base'** para confirmar y continuar."
                ),
                correlation_id=correlation_id,
            )
        validation_result = (
            economic_data.get("validation_result")
            if isinstance(economic_data.get("validation_result"), dict)
            else {}
        )
        perfil_usado = str(validation_result.get("perfil_usado") or "generic")
        _date_info = resolve_document_date(session_state)
        resumen = {
            "subtotal": round(subtotal, 2),
            "iva": iva,
            "total": total,
            "moneda": economic_data.get("currency", "MXN"),
            "fecha": _date_info.get("fecha_corta") or datetime.now().strftime("%d/%m/%Y"),
            "fecha_es": _date_info.get("fecha_es") or "",
            "perfil_usado": perfil_usado,
        }
        
        # 4. Generación de Archivos (misma raíz que TechnicalWriter/FormatsAgent)
        output_base_dir = os.path.join("/data", "outputs", session_id, "2.propuesta_economica")
        os.makedirs(output_base_dir, exist_ok=True)
        billing_spec = self._resolve_proportional_billing_spec(economic_data, mapeo_items)
        excel_lineage: Dict[str, Any] = {}
        econ_doc_meta = build_economic_doc_metadata(
            session_id=session_id,
            session_state=session_state,
            master_profile=master_profile,
            resumen=resumen,
            company_data=company_data,
        )

        mirrored_documents: List[Dict[str, Any]] = []
        if bool(getattr(settings, "TEMPLATE_MIRROR_ENABLED", True)):
            mirrored_documents = await self._mirror_economic_templates(
                session_id=session_id,
                session_state=session_state,
                output_base_dir=output_base_dir,
                master_profile=master_profile,
                economic_items=items,
                mapeo_items=mapeo_items,
            )

        # 4.1 Generar Excel de Precios (Modo Espejo vs Genérico) — resumen si no hubo plantillas
        excel_path = os.path.join(output_base_dir, "TABLA_PRECIOS_UNITARIOS.xlsx")
        
        # Verificar si podemos usar el protocolo de llenado sobre el original
        # Tomamos el primer item que tenga coordenadas de Excel
        first_excel_item = next((i for i in items if _has_valid_excel_locator(i)), None)
        
        if not mirrored_documents:
            if first_excel_item:
                source_file = first_excel_item["extra"].get("source_filename")
                excel_lineage = {
                    "source_filename": source_file,
                    "source_doc_id": first_excel_item.get("document_id"),
                    "source_path": first_excel_item["extra"].get("source_path"),
                    "materialization_route": "fill_excel",
                }
                items_to_fill = []
                for it in items:
                    ext = it.get("extra") or {}
                    if _has_valid_excel_locator(it):
                        items_to_fill.append({
                            "sheet_name": it.get("sheet_name"),
                            "row_index": it.get("row_index"),
                            "price_column_index": ext.get("price_column_index"),
                            "final_price": it.get("precio_unitario"),
                        })

                try:
                    print(
                        f"[{self.name}] Activando llenado espejo Excel: {source_file}",
                        flush=True,
                    )
                    ref_path = first_excel_item["extra"].get("source_path")
                    excel_path = self.excel_filler.fill_proposal_excel(
                        session_id=session_id,
                        source_filename=source_file,
                        items_to_fill=items_to_fill,
                        output_filename="CATALOGO_LLENADO_OFICIAL.xlsx",
                        source_path=ref_path,
                        output_dir=output_base_dir,
                    )
                    excel_lineage["source_path"] = ref_path or excel_lineage.get("source_path")
                except Exception as e:
                    print(
                        f"[{self.name}] Falló llenado espejo, usando genérico: {e}",
                        flush=True,
                    )
                    self._generate_price_excel(
                        excel_path,
                        mapeo_items,
                        master_profile,
                        resumen,
                        billing_spec=billing_spec,
                        doc_metadata=econ_doc_meta,
                    )
                    excel_lineage = {"materialization_route": "deterministic"}
            else:
                self._generate_price_excel(
                    excel_path,
                    mapeo_items,
                    master_profile,
                    resumen,
                    billing_spec=billing_spec,
                    doc_metadata=econ_doc_meta,
                )
                excel_lineage = {"materialization_route": "deterministic"}

            if os.path.isfile(excel_path):
                stamp_corporate_excel_file(excel_path, econ_doc_meta)

            word_path = os.path.join(output_base_dir, "ANEXO_AE_PROPUESTA_ECONOMICA.docx")
            self._generate_anexo_ae(
                word_path,
                mapeo_items,
                resumen,
                master_profile,
                billing_spec=billing_spec,
                doc_metadata=econ_doc_meta,
            )

            carta_path = os.path.join(output_base_dir, "CARTA_COMPROMISO_PRECIOS.docx")
            self._generate_carta_compromiso(
                carta_path, resumen, master_profile, doc_metadata=econ_doc_meta
            )

            apu_path = os.path.join(output_base_dir, "ANALISIS_PRECIOS_UNITARIOS.docx")
            self._generate_apu_document(
                apu_path,
                resumen=resumen,
                master_profile=master_profile,
                mapeo_items=mapeo_items,
                session_id=session_id,
                session_state=session_state,
            )

            generated_documents = [
                attach_traceability(
                    {
                        "nombre": "Tabla de Precios Unitarios",
                        "ruta": excel_path,
                        "tipo": "tabla_precios",
                        "template_id": "tabla_precios",
                    },
                    source_doc_id=str(excel_lineage.get("source_doc_id") or "") or None,
                    source_filename=str(excel_lineage.get("source_filename") or "") or None,
                    source_path=str(excel_lineage.get("source_path") or "") or None,
                    source_hash=safe_file_sha256(str(excel_lineage.get("source_path") or "")),
                    template_id="tabla_precios",
                    materialization_route=str(excel_lineage.get("materialization_route") or "deterministic"),
                    output_hash=safe_file_sha256(excel_path),
                ),
                attach_traceability(
                    {
                        "nombre": "Anexo AE - Propuesta Económica",
                        "ruta": word_path,
                        "tipo": "anexo_economico",
                        "template_id": "anexo_economico",
                    },
                    template_id="anexo_economico",
                    materialization_route="deterministic",
                    output_hash=safe_file_sha256(word_path),
                ),
                attach_traceability(
                    {
                        "nombre": "Carta Compromiso de Precios",
                        "ruta": carta_path,
                        "tipo": "carta_compromiso",
                        "template_id": "carta_compromiso",
                    },
                    template_id="carta_compromiso",
                    materialization_route="deterministic",
                    output_hash=safe_file_sha256(carta_path),
                ),
                attach_traceability(
                    {
                        "nombre": "Análisis de Precios Unitarios",
                        "ruta": apu_path,
                        "tipo": "analisis_precios_unitarios",
                        "template_id": "apu",
                    },
                    template_id="apu",
                    materialization_route="deterministic_apu",
                    output_hash=safe_file_sha256(apu_path),
                    provenance_ui={
                        "source": "deterministic_economic",
                        "confidence": 0.98,
                    },
                ),
            ]
        else:
            generated_documents = list(mirrored_documents)
            if not any("carta" in str(d.get("nombre", "")).lower() for d in generated_documents):
                carta_path = os.path.join(output_base_dir, "CARTA_COMPROMISO_PRECIOS.docx")
                self._generate_carta_compromiso(carta_path, resumen, master_profile)
                generated_documents.append(
                    attach_traceability(
                        {
                            "nombre": "Carta Compromiso de Precios",
                            "ruta": carta_path,
                            "tipo": "carta_compromiso",
                            "template_id": "carta_compromiso",
                        },
                        template_id="carta_compromiso",
                        materialization_route="deterministic",
                        output_hash=safe_file_sha256(carta_path),
                    )
                )

        print(f"[{self.name}] ✅ Propuesta económica generada con éxito.", flush=True)
        fill_gate = validate_generated_documents_fill(
            stage="economic",
            generated_documents=generated_documents,
            master_profile=master_profile,
            provenance_context={
                "source": "economic_writer",
                "confidence": 0.95,
                "session_hint": f"{session_id} {session_state.get('name', '')}",
                "fecha_es": resumen.get("fecha_es"),
                "deadline_dt_iso": _date_info.get("deadline_dt"),
                "economic_resumen": {
                    "subtotal": resumen.get("subtotal"),
                    "iva": resumen.get("iva"),
                    "total": resumen.get("total"),
                },
            },
        )
        validation_events = [
            build_fill_validation_event(it, stage="economic")
            for it in (fill_gate.get("issues") or [])
            if isinstance(it, dict)
        ]
        if not bool(fill_gate.get("validation_passed", True)):
            from app.services.document_fill_ux_messages import build_fill_blocking_question

            company_name = str(master_profile.get("razon_social") or "").strip()
            human_question = build_fill_blocking_question(
                "economic",
                fill_gate.get("issues") or [],
                company_name=company_name,
            )
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=human_question,
                data={
                    "documentos": generated_documents,
                    "resumen_economico": resumen,
                    "document_fill_quality_gate": fill_gate,
                    "validation_events": validation_events,
                    "materialization_metrics": build_materialization_metrics(
                        stage="economic",
                        documents=generated_documents,
                        elapsed_ms=int((time.perf_counter() - started_at) * 1000),
                    ),
                    "missing": [
                        {
                            "field": "document_fill_quality_gate",
                            "label": "Completar datos de la cotización",
                            "question": human_question,
                            "type": "document_fill_quality_gate_blocking",
                            "blocking_items": fill_gate.get("issues") or [],
                        }
                    ],
                    "stage": "economic",
                },
                correlation_id=correlation_id,
            )

        return AgentOutput(
            status=AgentStatus.SUCCESS,
            agent_id=self.agent_id,
            session_id=session_id,
            data={
                "folder": output_base_dir,
                "documentos": generated_documents,
                "resumen_economico": resumen,
                "document_fill_quality_gate": fill_gate,
                "validation_events": validation_events,
                "materialization_metrics": build_materialization_metrics(
                    stage="economic",
                    documents=generated_documents,
                    elapsed_ms=int((time.perf_counter() - started_at) * 1000),
                ),
            },
            correlation_id=correlation_id
        )



    async def _mirror_economic_templates(
        self,
        *,
        session_id: str,
        session_state: Dict[str, Any],
        output_base_dir: str,
        master_profile: Dict[str, Any],
        economic_items: List[Dict[str, Any]],
        mapeo_items: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        """
        Espeja plantillas económicas ingestadas (.xlsx/.xls/.docx) al sobre económico.

        Prioriza catálogo de sesión; rellena Excel si hay coordenadas en partidas.
        """
        from app.services.ingested_file_resolver import (
            build_ingested_file_index,
            resolve_ingested_file,
        )
        from app.services.session_template_catalog import (
            build_session_template_catalog,
            normalize_filename_key,
        )
        from app.services.template_mirror_service import mirror_template_to_output

        catalog = session_state.get("session_template_catalog")
        documents = await self.context_manager.memory.get_documents(session_id)
        if not catalog:
            catalog = build_session_template_catalog(session_id, documents)
        index = build_ingested_file_index(documents)

        templates = [
            it
            for it in (catalog.get("items") or [])
            if isinstance(it, dict)
            and it.get("sobre_inferido") == "economico"
            and it.get("accion_recomendada") == "generar"
            and it.get("document_class") == "plantilla_oferta"
        ]
        max_n = int(getattr(settings, "TEMPLATE_MIRROR_MAX_ECONOMIC", 20) or 20)
        if max_n > 0:
            templates = templates[:max_n]

        line_items: List[Dict[str, Any]] = []
        try:
            line_items = await self.context_manager.memory.get_line_items_for_session(
                session_id
            )
        except Exception:
            line_items = []
        concept_prices = {}
        user_inputs = session_state.get("economic_user_inputs") or {}
        if isinstance(user_inputs, dict):
            cp = user_inputs.get("concept_prices")
            if isinstance(cp, dict):
                concept_prices = cp
        if line_items and concept_prices:
            line_items = apply_structured_price_inputs(line_items, concept_prices)

        profile_fill = {
            "rfc": master_profile.get("rfc"),
            "razon_social": master_profile.get("razon_social"),
            "representante_legal": master_profile.get("representante_legal"),
            "domicilio": master_profile.get("domicilio_fiscal") or master_profile.get("domicilio"),
            "licitacion": session_id,
        }

        out_docs: List[Dict[str, Any]] = []
        session_hint = f"{session_id} {session_state.get('name', '')}"
        for idx, tpl in enumerate(templates, start=1):
            fn = str(tpl.get("source_filename") or "")
            ref = resolve_ingested_file(
                fn,
                index,
                doc_id=tpl.get("doc_id"),
                source_path=tpl.get("source_path"),
            )
            if not ref:
                continue
            cross_tender_marker = detect_cross_tender_marker(
                [str(getattr(ref, "extracted_text", "") or "")],
                session_hint,
            )
            if cross_tender_marker:
                print(
                    f"[{self.name}] Plantilla económica omitida por contaminación cross-tender: "
                    f"{fn} ({cross_tender_marker})",
                    flush=True,
                )
                continue
            ext = ref.file_path.rsplit(".", 1)[-1].lower()
            safe = re.sub(r"[^\w\s-]", "", normalize_filename_key(fn).replace(" ", "_"))[:48]
            out_name = f"ECON_{idx:02d}_{safe}.{ext if ext != 'doc' else 'docx'}"
            out_path = os.path.join(output_base_dir, out_name)
            try:
                if ext in ("xlsx", "xls") and line_items:
                    doc_line_items = []
                    if str(ref.doc_id or ""):
                        doc_line_items = [
                            li
                            for li in line_items
                            if str(li.get("document_id") or "") == str(ref.doc_id or "")
                        ]
                    items_to_fill = []
                    for li in doc_line_items:
                        extra = li.get("extra") if isinstance(li.get("extra"), dict) else {}
                        if not _has_valid_excel_locator(li):
                            continue
                        qty = li.get("cantidad")
                        try:
                            qty_f = float(qty) if qty is not None else None
                        except (TypeError, ValueError):
                            qty_f = None
                        items_to_fill.append(
                            {
                                "sheet_name": li.get("sheet_name"),
                                "row_index": li.get("row_index"),
                                "price_column_index": extra.get("price_column_index"),
                                "final_price": float(li.get("precio_unitario") or 0),
                                "quantity": qty_f,
                                "amount_column_index": extra.get("subtotal_column_index"),
                                "total_column_index": extra.get("total_column_index"),
                                "quantity_column_index": extra.get("quantity_column_index"),
                            }
                        )
                    if items_to_fill:
                        out_path = self.excel_filler.fill_proposal_excel(
                            session_id=session_id,
                            source_filename=ref.filename,
                            items_to_fill=items_to_fill,
                            output_filename=out_name,
                            source_path=ref.file_path,
                            output_dir=output_base_dir,
                        )
                        out_docs.append(
                            attach_traceability(
                                {
                                    "nombre": fn,
                                    "ruta": out_path,
                                    "tipo": "plantilla_economica_espejo",
                                },
                                source_doc_id=str(ref.doc_id or tpl.get("doc_id") or "") or None,
                                source_filename=ref.filename,
                                source_path=ref.file_path,
                                source_hash=safe_file_sha256(ref.file_path),
                                mirror_mode="excel_fill_line_items",
                                materialization_route="fill_excel",
                                output_hash=safe_file_sha256(out_path),
                                provenance_ui=tpl.get("provenance_ui") if isinstance(tpl.get("provenance_ui"), dict) else None,
                            )
                        )
                        continue

                    if doc_line_items:
                        meta = mirror_template_to_output(ref, out_path, profile_fill, fill_profile=False)
                        out_docs.append(
                            attach_traceability(
                                {
                                    "nombre": fn,
                                    "ruta": meta["ruta"],
                                    "tipo": "plantilla_economica_espejo",
                                    "document_class": tpl.get("document_class"),
                                    "expected_fill_mode": "line_items",
                                    "fill_status": "skipped_missing_locator",
                                    "line_items_detected": len(doc_line_items),
                                    "valid_locator_count": 0,
                                },
                                source_doc_id=str(ref.doc_id or tpl.get("doc_id") or "") or None,
                                source_filename=ref.filename,
                                source_path=meta.get("source_path") or ref.file_path,
                                source_hash=meta.get("source_hash"),
                                mirror_mode="copy_excel_missing_locator",
                                materialization_route=meta.get("materialization_route") or "mirror",
                                output_hash=meta.get("output_hash") or safe_file_sha256(meta.get("ruta")),
                                provenance_ui=tpl.get("provenance_ui") if isinstance(tpl.get("provenance_ui"), dict) else None,
                            )
                        )
                        continue

                meta = mirror_template_to_output(
                    ref, out_path, profile_fill, fill_profile=(ext == "docx")
                )
                out_docs.append(
                    attach_traceability(
                        {
                            "nombre": fn,
                            "ruta": meta["ruta"],
                            "tipo": "plantilla_economica_espejo",
                        },
                        source_doc_id=str(ref.doc_id or tpl.get("doc_id") or "") or None,
                        source_filename=ref.filename,
                        source_path=meta.get("source_path") or ref.file_path,
                        source_hash=meta.get("source_hash"),
                        mirror_mode=meta.get("mirror_mode"),
                        materialization_route=meta.get("materialization_route") or "mirror",
                        output_hash=meta.get("output_hash") or safe_file_sha256(meta.get("ruta")),
                        provenance_ui=tpl.get("provenance_ui") if isinstance(tpl.get("provenance_ui"), dict) else None,
                    )
                )
            except Exception as exc:
                print(
                    f"[{self.name}] Espejo económico omitido {fn}: {exc}",
                    flush=True,
                )
        return out_docs

    @staticmethod
    def _resolve_proportional_billing_spec(
        economic_data: Dict[str, Any], items: List[Dict]
    ) -> Optional[Dict[str, Any]]:
        """
        Especificación universal de facturación proporcional Partida 1 (sin session_id fijo).
        Prioridad: payload del motor económico → variables de entorno → ítems mensuales P1.
        """
        for key in ("billing_proportional", "formula_incomplete_month"):
            raw = economic_data.get(key)
            if not isinstance(raw, dict):
                continue
            divisor = int(raw.get("days_divisor") or raw.get("divisor_dias") or 0)
            months = int(raw.get("months") or raw.get("meses") or 9)
            if divisor > 0:
                tarifa = float(
                    raw.get("tarifa_mensual")
                    or raw.get("monthly_rate")
                    or 0.0
                )
                if tarifa < 0.01:
                    tarifa = EconomicWriterAgent._first_partida1_monthly_rate(items)
                dias = int(raw.get("dias_transcurridos") or raw.get("days_elapsed") or 15)
                return {
                    "months": months,
                    "days_divisor": divisor,
                    "tarifa_mensual": tarifa,
                    "dias_transcurridos": max(1, dias),
                }

        env_div = int(os.environ.get("LICITAI_PROP_BILLING_DAYS_DIVISOR", "0") or "0")
        if env_div > 0:
            env_months = int(os.environ.get("LICITAI_PROP_BILLING_MONTHS", "9") or "9")
            tarifa = EconomicWriterAgent._first_partida1_monthly_rate(items)
            if tarifa > 0:
                return {
                    "months": env_months,
                    "days_divisor": env_div,
                    "tarifa_mensual": tarifa,
                    "dias_transcurridos": int(
                        os.environ.get("LICITAI_PROP_BILLING_DAYS_ELAPSED", "15") or "15"
                    ),
                }
        return None

    @staticmethod
    def _first_partida1_monthly_rate(items: List[Dict]) -> float:
        for it in items:
            partida = it.get("partida")
            unidad = str(it.get("unidad") or "").lower()
            if str(partida) in ("1", "1.0") and "mensual" in unidad:
                return float(it.get("precio_unitario") or 0.0)
        for it in items:
            if str(it.get("partida")) in ("1", "1.0"):
                return float(it.get("precio_unitario") or 0.0)
        return 0.0

    @staticmethod
    def _append_proportional_billing_block(
        ws, start_row: int, spec: Dict[str, Any], tarifa_cell: str
    ) -> None:
        """Bloque visible + fórmula Excel ((Tarifa*meses)/divisor)*días."""
        months = int(spec.get("months") or 9)
        divisor = int(spec.get("days_divisor") or 275)
        dias = int(spec.get("dias_transcurridos") or 15)
        r0 = start_row
        ws.cell(row=r0, column=1, value="FACTURACIÓN PROPORCIONAL — PARTIDA 1").font = Font(
            bold=True, size=11
        )
        ws.cell(
            row=r0 + 1,
            column=1,
            value=(
                "Si el servicio inicia después del 1.° día del mes, el pliego establece:"
            ),
        )
        ws.cell(
            row=r0 + 2,
            column=1,
            value=(
                f"((Tarifa mensual × {months} meses) / {divisor}) × días naturales transcurridos "
                f"(desde ingreso hasta fin de mes)"
            ),
        )
        ws.cell(row=r0 + 3, column=1, value="Tarifa mensual de referencia (P. Unitario P1):")
        ws.cell(row=r0 + 3, column=2, value=f"={tarifa_cell}")
        ws.cell(row=r0 + 4, column=1, value="Días naturales transcurridos (ejemplo editable):")
        dias_cell = f"B{r0 + 4}"
        ws.cell(row=r0 + 4, column=2, value=dias)
        formula = f"=(({tarifa_cell}*{months})/{divisor})*{dias_cell}"
        ws.cell(row=r0 + 5, column=1, value="Importe proporcional calculado:")
        ws.cell(row=r0 + 5, column=2, value=formula)
        ws.cell(row=r0 + 6, column=1, value="Fórmula literal del pliego:")
        ws.cell(
            row=r0 + 6,
            column=2,
            value=(
                f"(({months}×Tarifa)/{divisor})×Días — constantes {months} meses y {divisor} días"
            ),
        )

    def _generate_price_excel(
        self,
        path: str,
        items: List[Dict],
        profile: Dict,
        resumen: Dict,
        billing_spec: Optional[Dict[str, Any]] = None,
        doc_metadata: Optional[Dict[str, Any]] = None,
    ):
        """Crea un Excel profesional con fórmulas y formato."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Propuesta Económica"
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        center_align = Alignment(horizontal="center", vertical="center")
        border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        header_row = (
            apply_corporate_excel_letterhead(ws, doc_metadata)
            if doc_metadata
            else 3
        )
        if not doc_metadata:
            ws.merge_cells("A1:F1")
            ws["A1"] = profile.get("razon_social", "EMPRESA LICITANTE").upper()
            ws["A1"].font = Font(bold=True, size=14)
            ws["A1"].alignment = center_align
        
        # Títulos de Columnas
        headers = ["Partida", "Descripción", "Unidad", "Cantidad", "P. Unitario", "Importe"]
        for col, text in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col, value=text)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
            
        # Datos
        current_row = header_row + 1
        first_p1_row: Optional[int] = None
        for item in items:
            ws.cell(row=current_row, column=1, value=item.get("partida")).border = border
            ws.cell(row=current_row, column=2, value=item.get("descripcion")).border = border
            ws.cell(row=current_row, column=3, value=item.get("unidad")).border = border
            ws.cell(row=current_row, column=4, value=item.get("cantidad")).border = border
            ws.cell(row=current_row, column=5, value=item.get("precio_unitario")).border = border
            ws.cell(row=current_row, column=6, value=item.get("importe")).border = border
            if first_p1_row is None and str(item.get("partida")) in ("1", "1.0"):
                first_p1_row = current_row
            current_row += 1
            
        # Totales desde resumen (calculados en Fase 1, no se recalculan)
        ws.cell(row=current_row + 1, column=5, value="SUBTOTAL:").font = Font(bold=True)
        ws.cell(row=current_row + 1, column=6, value=resumen["subtotal"]).font = Font(bold=True)
        iva_pct = (resumen['iva'] / resumen['subtotal'] * 100) if resumen['subtotal'] > 0 else 16.0
        ws.cell(row=current_row + 2, column=5, value=f"IVA ({iva_pct:g}%):").font = Font(bold=True)
        ws.cell(row=current_row + 2, column=6, value=resumen["iva"]).font = Font(bold=True)
        ws.cell(row=current_row + 3, column=5, value="TOTAL:").font = Font(bold=True)
        ws.cell(row=current_row + 3, column=6, value=resumen["total"]).font = Font(bold=True)

        if billing_spec and billing_spec.get("days_divisor"):
            tarifa_ref = (
                f"E{first_p1_row}"
                if first_p1_row
                else str(round(float(billing_spec.get("tarifa_mensual") or 0), 2))
            )
            self._append_proportional_billing_block(
                ws, current_row + 5, billing_spec, tarifa_ref
            )
        
        # Ajustar anchos
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
        
        wb.save(path)

    def _generate_anexo_ae(
        self,
        path: str,
        items: List[Dict],
        resumen: Dict,
        profile: Dict,
        billing_spec: Optional[Dict[str, Any]] = None,
        doc_metadata: Optional[Dict[str, Any]] = None,
    ):
        """Genera el Word del Anexo AE (Propuesta Económica Detallada)."""
        doc = Document()
        if doc_metadata:
            apply_corporate_docx_letterhead(doc, doc_metadata)

        doc.add_heading('ANEXO AE: PROPUESTA ECONÓMICA', 0)
        
        p = doc.add_paragraph()
        run = p.add_run(f"LICITANTE: {profile.get('razon_social', '...')}\n")
        run.bold = True
        p.add_run(f"RFC: {profile.get('rfc', '...')}\n")
        p.add_run(f"REPRESENTANTE: {profile.get('representante_legal', '...')}\n")
        p.add_run(f"FECHA: {resumen['fecha']}")

        doc.add_paragraph("\nPor medio de la presente, sometemos a su consideración nuestra propuesta económica detallada:")

        # Tabla Word
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Partida'
        hdr_cells[1].text = 'Concepto'
        hdr_cells[2].text = 'Cant.'
        hdr_cells[3].text = 'Importe'

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('partida'))
            row_cells[1].text = item.get('descripcion')
            row_cells[2].text = str(item.get('cantidad'))
            row_cells[3].text = f"${item.get('importe'):,.2f}"

        if billing_spec and billing_spec.get("days_divisor"):
            months = int(billing_spec.get("months") or 9)
            divisor = int(billing_spec.get("days_divisor") or 275)
            tarifa = float(billing_spec.get("tarifa_mensual") or 0.0)
            dias = int(billing_spec.get("dias_transcurridos") or 15)
            monto = (tarifa * months) / divisor * dias if divisor else 0.0
            doc.add_heading("Facturación proporcional — Partida 1", level=2)
            doc.add_paragraph(
                "Si el servicio de limpieza inicia posterior al primer día del mes, "
                "aplica la fórmula del pliego:"
            )
            doc.add_paragraph(
                f"((Tarifa mensual × {months} meses) / {divisor}) × días naturales transcurridos "
                f"(desde la fecha de ingreso hasta el último día del mes)."
            )
            doc.add_paragraph(
                f"Tarifa mensual de referencia: ${tarifa:,.2f} MXN | "
                f"Días de ejemplo: {dias} | Importe proporcional: ${monto:,.2f} MXN"
            )

        doc.add_paragraph(f"\nSUBTOTAL: ${resumen['subtotal']:,.2f}")
        # Calculamos el porcentaje real para mostrarlo en el documento
        iva_pct = (resumen['iva'] / resumen['subtotal'] * 100) if resumen['subtotal'] > 0 else 16.0
        doc.add_paragraph(f"I.V.A. ({iva_pct:g}%): ${resumen['iva']:,.2f}")
        para_total = doc.add_paragraph(f"TOTAL DE LA PROPUESTA: ${resumen['total']:,.2f}")
        para_total.runs[0].bold = True

        doc.add_paragraph("\nVIGENCIA DE LA PROPUESTA: 30 DÍAS NATURALES.")
        
        doc.add_paragraph("\n\n__________________________________")
        doc.add_paragraph(f"{profile.get('representante_legal', 'Representante Legal')}\nfirma")

        doc.save(path)

    def _generate_apu_document(
        self,
        path: str,
        *,
        resumen: Dict[str, Any],
        master_profile: Dict[str, Any],
        mapeo_items: List[Dict[str, Any]],
        session_id: str,
        session_state: Dict[str, Any],
    ) -> None:
        """APU en perspectiva concursante con montos del motor económico."""
        dom = str(
            master_profile.get("domicilio_fiscal")
            or master_profile.get("domicilio")
            or ""
        ).strip()
        ciudad = dom.split(",")[0].strip() if dom else "México"
        content = build_apu_markdown(
            razon_social=str(master_profile.get("razon_social") or ""),
            rfc=str(master_profile.get("rfc") or ""),
            representante=str(master_profile.get("representante_legal") or ""),
            domicilio=dom,
            fecha_es=str(resumen.get("fecha_es") or resumen.get("fecha") or ""),
            procedimiento=session_id.replace("_", " "),
            subtotal=float(resumen.get("subtotal") or 0),
            iva=float(resumen.get("iva") or 0),
            total=float(resumen.get("total") or 0),
            line_items=mapeo_items,
            ciudad=ciudad,
        )
        from app.agents.formats import _save_docx

        metadata = build_economic_doc_metadata(
            session_id=session_id,
            session_state=session_state,
            master_profile=master_profile,
            resumen=resumen,
        )
        _save_docx("ANÁLISIS DE PRECIOS UNITARIOS", content, path, metadata)

    def _generate_carta_compromiso(
        self,
        path: str,
        resumen: Dict,
        profile: Dict,
        doc_metadata: Optional[Dict[str, Any]] = None,
    ):
        """Genera la carta formal de compromiso de precios."""
        doc = Document()
        if doc_metadata:
            apply_corporate_docx_letterhead(doc, doc_metadata)
        doc.add_heading('CARTA COMPROMISO DE PRECIOS', 1)
        
        p = doc.add_paragraph(f"\nMéxico, a {resumen['fecha']}\n")
        p.alignment = 2 # Derecha
        
        doc.add_paragraph("A QUIEN CORRESPONDA:")
        
        body = f"""
        Quien suscribe, C. {profile.get('representante_legal', '...')}, en mi carácter de Representante Legal 
        de la empresa {profile.get('razon_social', '...')}, con RFC {profile.get('rfc', '...')}, manifiesto bajo protesta de decir verdad que:
        
        Los precios presentados en nuestra propuesta económica de fecha {resumen['fecha']} por un total de 
        ${resumen['total']:,.2f} ({resumen['moneda']}), permanecerán firmes y vigentes durante la totalidad 
        del proceso de adjudicación y, en caso de resultar adjudicado, durante la vigencia del contrato respectivo.
        
        Asimismo, garantizamos que los precios no están sujetos a variaciones por fluctuaciones de mercado o 
        costos de insumos durante el periodo mencionado.
        
        Atentamente,
        """
        doc.add_paragraph(body)
        
        doc.add_paragraph("\n\n__________________________________")
        doc.add_paragraph(f"{profile.get('representante_legal', '...')}\n{profile.get('razon_social', '...')}")
        
        doc.save(path)
