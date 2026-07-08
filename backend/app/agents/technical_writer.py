import os
import re
import unicodedata
import json
import time
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Any, Dict, List, Optional, Set
from app.agents.base_agent import BaseAgent
from app.agents.mcp_context import MCPContextManager
from app.services.vector_service import VectorDbServiceClient
from app.services.resilient_llm import ResilientLLMClient
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.contracts.document_inventory import (
    DocumentEnvelope,
    DocumentInventory,
    InventoryItemStatus,
)
from app.core.formats_pilot_slots import build_formats_pilot_missing_entries
from app.core.observability import get_logger
from app.utils.doc_formatting import (
    ANTI_PLACEHOLDER_PROMPT_RULE,
    repair_docx_file_placeholders,
    strip_bracket_placeholders_for_docx,
    strip_markdown_for_docx,
)
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.services.company_experience_context import (
    build_company_experience_context_block,
    build_experience_sources_ux_summary,
    extract_client_references_from_documents,
    fill_te03_client_placeholders,
    req_needs_company_experience,
)
from app.services.document_fill_quality_gate import (
    detect_cross_tender_marker,
    validate_generated_documents_fill,
)
from app.services.document_fill_ux_messages import build_fill_validation_event
from app.services.document_traceability import (
    attach_traceability,
    build_materialization_metrics,
    safe_file_sha256,
)
from app.config.settings import settings as app_settings

logger = get_logger(__name__)

# Prefijos de ID que identifican un requisito como redactable técnicamente.
# Provienen del esquema de IDs que asigna ComplianceAgent (zona "tecnico").
TECH_ID_PREFIXES = ("2.",)


def _clean_profile_text(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    text = re.sub(r"\s*\|\s*", ", ", text)
    text = re.sub(r"\bOTRA NO ESPECIFICADA EN EL\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bN/?A\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*,\s*,+", ", ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip(" ,;-")


def _build_carta_presentacion_text(
    *,
    razon_social: str,
    rfc: str,
    representante: str,
    domicilio: str,
    tender_name: str,
    fecha_es: str,
    destinatario: str = "",
) -> str:
    domicilio_line = _clean_profile_text(domicilio) or "No disponible"
    dest_block = (destinatario or "A QUIEN CORRESPONDA:").strip()
    return (
        "LUGAR Y FECHA\n"
        f"{fecha_es}\n\n"
        f"{dest_block}\n\n"
        f"Por medio de la presente, {razon_social}, con RFC {rfc}, "
        f"por conducto de su representante legal {representante}, con domicilio en {domicilio_line}, "
        f"presenta su propuesta técnica para {tender_name}.\n\n"
        "Manifestamos bajo protesta de decir verdad que la documentación técnica adjunta "
        "se integra conforme a los requisitos, anexos y especificaciones solicitados en las bases, "
        "y que contamos con capacidad legal, técnica, operativa y material para cumplir con las obligaciones "
        "que deriven del procedimiento y, en su caso, del contrato correspondiente.\n\n"
        "Asimismo, nos comprometemos a sostener el contenido de la propuesta técnica presentada "
        "y a atender los requerimientos de aclaración, evaluación y formalización que emita la convocante "
        "conforme a la normatividad aplicable.\n\n"
        "ATENTAMENTE\n\n"
        f"{representante}\n"
        f"Representante Legal de {razon_social}"
    )


def _extract_te12_threshold(req_context: str) -> str:
    ctx = str(req_context or "")
    m = re.search(
        r"(cuando\s+menos\s+\d+\s+de\s+(?:los|un)\s+\d+\s+m[aá]ximos?\s+que\s+se\s+pueden\s+obtener)",
        ctx,
        flags=re.IGNORECASE,
    )
    if m:
        return m.group(1).strip().rstrip(".") + "."
    m = re.search(
        r"(\d+\s+de\s+\d+\s+m[aá]ximos?)",
        ctx,
        flags=re.IGNORECASE,
    )
    if m:
        return f"La puntuación mínima requerida será de {m.group(1).strip()}."
    return "La puntuación mínima requerida será la prevista expresamente en las bases del procedimiento."


def _build_te12_text(
    *,
    razon_social: str,
    rfc: str,
    representante: str,
    domicilio: str,
    tender_name: str,
    fecha_es: str,
    req_context: str,
    destinatario: str = "",
) -> str:
    domicilio_line = _clean_profile_text(domicilio) or "No disponible"
    threshold_line = _extract_te12_threshold(req_context)
    dest_block = (destinatario or "A QUIEN CORRESPONDA:").strip()
    return (
        "PUNTUACIÓN O UNIDADES PORCENTUALES A OBTENER EN LA PROPUESTA TÉCNICA\n\n"
        "LUGAR Y FECHA\n"
        f"{fecha_es}\n\n"
        f"{dest_block}\n\n"
        f"Quien suscribe, {representante}, en representación de {razon_social}, RFC {rfc}, "
        f"con domicilio en {domicilio_line}, comparece para manifestar que conoce y acepta "
        f"el criterio de evaluación técnica aplicable a {tender_name}.\n\n"
        f"{threshold_line}\n\n"
        "Bajo protesta de decir verdad, manifestamos que la propuesta técnica presentada se integra "
        "conforme a los requisitos y criterios de evaluación establecidos por la convocante, "
        "por lo que solicitamos sea considerada para su revisión y dictaminación.\n\n"
        "ATENTAMENTE\n\n"
        f"{representante}\n"
        f"Representante Legal de {razon_social}"
    )


def _normalized_blob(value: Any) -> str:
    raw = unicodedata.normalize("NFD", str(value or ""))
    raw = "".join(ch for ch in raw if unicodedata.category(ch) != "Mn")
    return raw.lower()


def _mirror_source_has_cross_tender_marker(ref: Any, session_hint: str) -> bool:
    text = str(getattr(ref, "extracted_text", "") or "")
    return bool(detect_cross_tender_marker([text], session_hint))


def _infer_participation_zones(documents: List[Dict[str, Any]]) -> List[str]:
    zones: Set[str] = set()
    for doc in documents or []:
        if not isinstance(doc, dict):
            continue
        content = doc.get("content") if isinstance(doc.get("content"), dict) else {}
        meta = doc.get("metadata") if isinstance(doc.get("metadata"), dict) else {}
        filename = str(content.get("filename") or meta.get("filename") or "")
        blob = _normalized_blob(filename)
        if "zona a" in blob or " za " in f" {blob} " or blob.endswith(" za.xlsx"):
            zones.add("A")
        if "zona b" in blob or " zb " in f" {blob} " or blob.endswith(" zb.xlsx"):
            zones.add("B")
        if "zona c" in blob or " zc " in f" {blob} " or blob.endswith(" zc.xlsx"):
            zones.add("C")
        if "zona d" in blob or " zd " in f" {blob} " or blob.endswith(" zd.xlsx"):
            zones.add("D")
    return sorted(zones)


def _infer_contenido_nacional_pct(vector_db: VectorDbServiceClient, session_id: str) -> str:
    try:
        res = vector_db.query_texts(session_id, "contenido nacional 65% anexo iii-k punto 46", n_results=4)
        for doc in res.get("documents") or []:
            text = str(doc or "")
            if "contenido nacional" not in text.lower():
                continue
            m = re.search(r"(\d{1,3})\s*%", text)
            if m:
                return m.group(1)
    except Exception as exc:
        logger.warning("technical_writer_contenido_nacional_pct_failed", session_id=session_id, error=str(exc))
    return ""


def _build_contenido_nacional_text(
    *,
    razon_social: str,
    rfc: str,
    representante: str,
    session_id: str,
    tender_name: str,
    fecha_es: str,
    zonas: List[str],
    porcentaje: str,
    destinatario: str = "",
) -> str:
    zonas_txt = ", ".join(zonas) if zonas else "correspondientes a la propuesta presentada"
    if len(zonas) == 1:
        zona_clause = f"la zona {zonas_txt}"
    elif len(zonas) > 1:
        zona_clause = f"las zonas {zonas_txt}"
    else:
        zona_clause = "la zona en que participo"
    porcentaje_txt = porcentaje or "65"
    proc_ref = tender_name.strip() or session_id.replace("_", " ").upper()
    dest_block = (destinatario or "A QUIEN CORRESPONDA:").strip()
    return (
        "MANIFESTACIÓN DE NACIONALIDAD MEXICANA, PRODUCCIÓN EN MÉXICO Y GRADO DE CONTENIDO NACIONAL\n\n"
        f"{fecha_es}\n\n"
        f"{dest_block}\n\n"
        f"Me refiero al procedimiento de contratación identificado como {proc_ref}, "
        f"en el que mi representada, la empresa {razon_social}, RFC {rfc}, participa a través de la propuesta contenida en el presente sobre.\n\n"
        "Sobre el particular y en términos de lo previsto por el Acuerdo por el que se establecen las reglas para la determinación "
        "del grado de contenido nacional, tratándose de procedimientos de contratación de carácter nacional, el suscrito manifiesta "
        "bajo protesta de decir verdad que los bienes ofertados para la partida 2 correspondiente a "
        f"{zona_clause}, serán producidos en México y contendrán un grado de contenido nacional de cuando menos el {porcentaje_txt} por ciento, "
        "en el supuesto de que sea adjudicado el pedido respectivo. Asimismo, manifiesto bajo protesta de decir verdad ser de nacionalidad mexicana.\n\n"
        "De igual forma, manifiesto tener conocimiento de lo previsto en el artículo 57 de la Ley de Adquisiciones, "
        "Arrendamientos y Servicios del Sector Público y me comprometo, en caso de ser requerido, a aceptar una verificación "
        "del cumplimiento de los requisitos sobre el contenido nacional de los bienes ofertados mediante la exhibición de la "
        "información documental correspondiente y/o a través de una inspección física de la planta industrial en la que se producen.\n\n"
        "ATENTAMENTE\n\n"
        f"{razon_social}\n\n"
        f"{representante}\n"
        "Representante Legal"
    )


def _is_technical_writable(req: Dict[str, Any]) -> bool:
    """Incluye requisitos técnicos aunque el id no sea 2.x (p. ej. Forma AT- en cap. 7)."""
    r_id = str(req.get("id", ""))
    nombre_raw = str(req.get("nombre", "") or "")
    desc_raw = str(req.get("descripcion", "") or "")
    snippet_raw = str(req.get("snippet", "") or "")
    nombre = nombre_raw.upper()
    descripcion = desc_raw.upper()
    blob = f"{r_id} {nombre} {descripcion}".upper()

    from app.services.document_deliverable_filter import is_company_credential_present_only

    if is_company_credential_present_only(nombre_raw, desc_raw, snippet_raw):
        return False

    # Si el compliance ya clasificó el tipo_accion, usarlo directamente
    tipo_accion = str(req.get("tipo_accion", "")).lower()
    if tipo_accion == "generar":
        return True
    if tipo_accion in ("informativo", "presentar_fisico", "requiere_datos_licitante"):
        return False

    # Excluir requisitos informativos de las bases que NO son documentos a redactar
    _INFORMATIVE_BASES_PATTERNS = (
        "FECHA Y HORA DE LOS EVENTOS",
        "FECHA Y HORA DEL ACTO",
        "VISITA AL SITIO",
        "JUNTA DE ACLARACIONES",
        "LUGAR DE ENTREGA",
        "FORMATO DE PRESENTACION",
        "FORMATO DE PROPUESTA ECONOMICA",
        "APERTURA DE PROPUESTAS",
        "ACTO DE PRESENTACION",
        "PRESENTACION DE LAS PROPUESTAS",
        "LECTURA DEL DICTAMEN",
        "EL COMITE ANALIZARA",
        "LA PROPUESTA TECNICA DEBERA DESCRIBIR",
    )
    if any(p in blob for p in _INFORMATIVE_BASES_PATTERNS):
        return False

    if req.get("tipo") == "tecnico":
        return True
    if req.get("inventory_synthetic") is True:
        return True
    if any(r_id.startswith(p) for p in TECH_ID_PREFIXES):
        return True
    if re.search(r"\bAT[-_]?\d", blob):
        return True
    keys = (
        "PROGRAMA",
        "RELACIÓN",
        "RELACION",
        "MAQUINARIA",
        "EXPERIENCIA",
        "SUPERINTENDENTE",
        "SUPERVISOR",
        "CURRÍCULUM",
        "CURRICULUM",
        "CAPACIDAD FINANCIERA",
        "CONTRATOS EN VIGOR",
        "MODELO DE CONTRATO",
        "PROPUESTA TÉCNICA",
        "PROPUESTA TECNICA",
        "INSTALACIÓN",
        "INSTALACION",
        "CALENDARIZ",
    )
    return any(k in blob for k in keys)


def _inventory_doc_already_on_disk(tech_dir: str, canonical_id: str) -> bool:
    """Evita regenerar si ya existe un .docx cuyo nombre incluye el canonical_id."""
    if not canonical_id or not os.path.isdir(tech_dir):
        return False
    token = re.sub(r"[^\w\-]+", "_", canonical_id).strip("_").lower()
    if len(token) < 3:
        return False
    try:
        for fn in os.listdir(tech_dir):
            if not fn.endswith(".docx") or fn.startswith("~$"):
                continue
            if token in fn.lower():
                return True
    except OSError:
        return False
    return False


def _should_block_by_quality_gate(
    *,
    total_items: int,
    generar_count: int,
    unknown_count: int,
    evidence_match_ratio: float,
    presentar_fisico_count: int = 0,
    triage_context: Optional[Dict[str, Any]] = None,
    queued_requirements_count: int = 0,
    actionable_action_count: int = 0,
) -> Dict[str, Any]:
    """
    Gate duro: evita generación cuando la lista documental está degradada.

    Excepción OBRA: en licitaciones de obra pública (LOPSRM / categoría OBRA),
    los requisitos técnicos son formas predefinidas (AT-10, AT-13, AE-02) que el
    licitante llena y presenta físicamente. El ComplianceAgent los clasifica
    correctamente como presentar_fisico. Si generar_count == 0 pero hay ítems
    presentar_fisico, no es un error de clasificación — es el comportamiento
    esperado para este tipo de licitación.
    """
    if not bool(app_settings.DOCUMENT_QUALITY_HARD_GATE_ENABLED):
        return {"block": False, "reason": "", "metrics": {}}

    if queued_requirements_count > 0:
        return {
            "block": False,
            "reason": "queued_requirements_present",
            "metrics": {
                "queued_requirements_count": queued_requirements_count,
                "total_items": total_items,
                "generar_count": generar_count,
                "actionable_action_count": actionable_action_count,
            },
        }

    # ── Excepción por categoría de licitación ──────────────────────────────
    tender_category = ""
    if isinstance(triage_context, dict):
        tender_category = str(triage_context.get("tender_category") or "").upper()

    if tender_category == "OBRA":
        if generar_count == 0 and presentar_fisico_count > 0:
            return {
                "block": False,
                "reason": "obra_category_no_generate_items_expected",
                "metrics": {
                    "total_items": total_items,
                    "generar_count": generar_count,
                    "presentar_fisico_count": presentar_fisico_count,
                    "tender_category": tender_category,
                    "evidence_match_ratio": evidence_match_ratio,
                },
            }
        if total_items == 0:
            return {"block": False, "reason": "", "metrics": {"tender_category": tender_category}}

    min_items = max(1, int(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MIN_ITEMS", 3) or 3))
    max_unknown = float(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MAX_UNKNOWN_RATIO", 0.6) or 0.6)
    min_evidence = float(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MIN_EVIDENCE_MATCH_RATIO", 0.5) or 0.5)
    if total_items < min_items:
        return {"block": False, "reason": "", "metrics": {}}
    unknown_ratio = (unknown_count / total_items) if total_items else 0.0
    if generar_count == 0 and actionable_action_count == 0:
        return {
            "block": True,
            "reason": "no_actionable_generate_items",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "actionable_action_count": actionable_action_count,
                "unknown_ratio": unknown_ratio,
                "evidence_match_ratio": evidence_match_ratio,
                "tender_category": tender_category or None,
            },
        }
    if unknown_ratio > max_unknown:
        return {
            "block": True,
            "reason": "unknown_ratio_above_threshold",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "unknown_ratio": unknown_ratio,
                "threshold_unknown_ratio": max_unknown,
                "evidence_match_ratio": evidence_match_ratio,
                "tender_category": tender_category or None,
            },
        }
    if evidence_match_ratio < min_evidence:
        return {
            "block": True,
            "reason": "evidence_match_ratio_below_threshold",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "unknown_ratio": unknown_ratio,
                "evidence_match_ratio": evidence_match_ratio,
                "threshold_evidence_match_ratio": min_evidence,
                "tender_category": tender_category or None,
            },
        }
    return {"block": False, "reason": "", "metrics": {}}


def _merge_document_inventory_technical(
    company_data: Dict[str, Any],
    tech_dir: str,
    tech_requirements: List[Dict[str, Any]],
    seen_ids: Set[str],
) -> List[Dict[str, Any]]:
    """
    Añade tareas desde ``document_inventory`` (Modo Fábrica): categoría técnica y pendiente.

    Los dicts resultantes son compatibles con el bucle de generación existente.
    """
    raw = company_data.get("document_inventory")
    if not isinstance(raw, dict) or not raw.get("items"):
        return tech_requirements
    try:
        inv = DocumentInventory.model_validate(raw)
    except Exception as e:
        logger.warning("technical_writer_inventory_parse_failed", error=str(e))
        return tech_requirements

    extra: List[Dict[str, Any]] = []
    for it in inv.items:
        if it.category != DocumentEnvelope.TECHNICAL:
            continue
        if it.status != InventoryItemStatus.PENDING:
            continue
        cid = (it.canonical_id or "").strip()
        if not cid:
            continue
        key = cid.lower()
        if key in seen_ids:
            continue
        if _inventory_doc_already_on_disk(tech_dir, cid):
            continue
        seen_ids.add(key)
        extra.append(
            {
                "id": cid,
                "nombre": it.display_name,
                "descripcion": (it.description or "").strip(),
                "tipo": "tecnico",
                "from_document_inventory": True,
                "generator_hint": it.generator_hint,
            }
        )
    return tech_requirements + extra


class TechnicalWriterAgent(BaseAgent):
    """
    Agente 3: Redactor Técnico
    Genera UN documento Word por cada requisito técnico detectado por Compliance.
    SINCRONIZACIÓN TOTAL: Usa la lista maestra del Auditor para cobertura del 100%.
    """
    def __init__(self, context_manager: MCPContextManager):
        super().__init__(
            agent_id="tech_writer_001",
            name="Redactor Técnico Senior",
            description="Genera documentos técnicos basados en la auditoría de cumplimiento.",
            context_manager=context_manager
        )
        # Instanciado en constructor para que sea mockeable en tests unitarios
        self.llm = ResilientLLMClient()
        self.vector_db = VectorDbServiceClient()

    async def process(self, agent_input: AgentInput) -> AgentOutput:
        session_id = agent_input.session_id
        correlation_id = agent_input.correlation_id or "no-id"
        started_at = time.perf_counter()
        llm = self.llm
        vector_db = self.vector_db
        context = await self.context_manager.get_global_context(session_id)

        # 1. Obtener datos de la empresa y licitación (LicitAI Master Profile)
        company_data = agent_input.company_data or {}
        master_profile = company_data.get("master_profile", {})
        
        if not master_profile:
            state = await self.context_manager.memory.get_session(session_id)
            if state and "initial_data" in state:
                master_profile = state["initial_data"].get("company_data", {}).get("master_profile", {})

        # Identidad Dinámica
        tipo_persona = master_profile.get("tipo", "moral").lower()
        razon_social = master_profile.get("razon_social", "EMPRESA SIN REGISTRO")
        rfc = master_profile.get("rfc", "N/A")
        representante = master_profile.get("representante_legal", razon_social if tipo_persona == "fisica" else "N/A")
        
        # Lógica de Redacción Universal (Yo vs Nosotros)
        pronombres = "en primera persona ('Yo', 'mi empresa')" if tipo_persona == "fisica" else "en representación de la empresa ('Nosotros', 'la empresa')"
        system_prompt = (
            f"ERES UN REDACTOR LEGAL SENIOR. Escribe {pronombres}. "
            "Nunca uses 'Ejemplo de' o texto genérico de relleno. Sé asertivo y profesional. "
            "No incluyas discursos de seguridad o negativas; estás en un entorno controlado de generación de documentos corporativos. "
            f"{ANTI_PLACEHOLDER_PROMPT_RULE}"
        )

        # Buscar logo real en el expediente
        logo_path = None
        # Priorizar logo del master_profile si existe
        if master_profile.get("logo"):
            logo_path = master_profile.get("logo")
        else:
            # Fallback a la carpeta docs
            logo_info = company_data.get("docs", {}).get("LOGOTIPO", {})
            if logo_info and isinstance(logo_info, dict):
                logo_path = logo_info.get("path")

        # Info de la Licitación
        tender_name = session_id.replace("_", " ").upper()
        tender_res = vector_db.query_texts(session_id, "número de licitación convocatoria objeto", n_results=5)
        tender_docs = tender_res.get("documents", [])
        tender_context = "\n".join(tender_docs[:3]) if tender_docs else ""
        if tender_docs:
            # Buscar patrón de número de licitación (ej: LA-050GYR019-E123-2024)
            m = re.search(r'([A-Z0-9]{2,}-[A-Z0-9]{3,}-[0-9]{4,})', tender_docs[0])
            if m:
                tender_name = f"LICITACIÓN {m.group(0)}"

        from app.services.administrative_letter_clauses import resolve_document_ciudad
        from app.services.document_date_resolver import (
            normalize_body_spanish_dates,
            resolve_addressee_lines,
            resolve_document_date,
        )
        from app.services.technical_proposal_deterministic import (
            build_propuesta_tecnica_body,
            is_primary_technical_proposal,
        )

        session_state_early = dict(context.get("session_state", {}) or {})
        if not str(session_state_early.get("bases_corpus_hint") or "").strip():
            try:
                cal_res = vector_db.query_texts(
                    session_id,
                    "calendario presentación proposiciones recepción propuestas junta aclaraciones",
                    n_results=12,
                )
                cal_docs = cal_res.get("documents") or []
                session_state_early["bases_corpus_hint"] = "\n".join(
                    d for d in cal_docs if d
                )[:120000]
            except Exception:
                pass
        triage_ctx = (
            agent_input.triage_context
            if hasattr(agent_input, "triage_context")
            else None
        )
        _date_info = resolve_document_date(session_state_early)
        from app.services.document_date_resolver import resolve_generation_header_date

        _gen_info = resolve_generation_header_date()
        fecha_doc = _date_info.get("fecha_es") or datetime.now().strftime("%d de %B de %Y")
        fecha_hdr = _gen_info.get("fecha_es") or fecha_doc
        fecha_es = fecha_doc  # cuerpo / fecha canónica expediente
        destinatario = resolve_addressee_lines(session_state_early, triage_ctx)

        # Metadata para Word
        _dom_fiscal = master_profile.get("domicilio_fiscal") or master_profile.get("domicilio") or ""
        doc_metadata = {
            "logo_path": logo_path,
            "tender_name": tender_name,
            "fecha": fecha_doc,
            "fecha_documental": fecha_doc,
            "fecha_encabezado": fecha_hdr,
            "fecha_generacion": fecha_hdr,
            "fecha_corta": _date_info.get("fecha_corta") or "",
            "fecha_documental_source": _date_info.get("source", ""),
            "fecha_encabezado_source": _gen_info.get("source", "generation_timestamp"),
            "generated_at_iso": _gen_info.get("generated_at_iso", ""),
            "empresa": razon_social,
            "rfc": rfc,
            "representante": representante,
            "tipo_persona": tipo_persona,
            "domicilio": _dom_fiscal,
            "ciudad": resolve_document_ciudad(master_profile, str(_dom_fiscal)),
            "footer_text": f"{razon_social} | RFC: {rfc} | Domicilio: {_dom_fiscal or 'S/D'}",
            "destinatario": destinatario,
            "formal_closing": True,
        }

        # 1. Crear estructura de carpetas
        base_output_dir = os.path.join("/data", "outputs", session_id)
        tech_dir = os.path.join(base_output_dir, "1.propuesta tecnica")
        os.makedirs(tech_dir, exist_ok=True)

        # 1.1 Inyección de Overrides Económicos (Hito 3.1)
        session_state = context.get("session_state", {})
        from app.services.technical_canonical_v1 import gate_technical_generation_chat_first

        _tech_chat_gate = gate_technical_generation_chat_first(session_state)
        if _tech_chat_gate:
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=str(_tech_chat_gate.get("message") or ""),
                data={
                    "technical_capture_v1": _tech_chat_gate.get("technical_capture_v1"),
                    "missing_labels": _tech_chat_gate.get("missing_labels") or [],
                },
                correlation_id=correlation_id,
            )
        user_inputs = session_state.get("economic_user_inputs") or {}
        econ_parts = []
        for k, v in user_inputs.items():
            if k == "concept_prices" and isinstance(v, dict):
                for c, p in v.items():
                    econ_parts.append(f"Precio Unitario de {c.upper()}: {p}")
            elif v is not None:
                # Normalizar nombres de llaves para el LLM
                lbl = k.replace("_", " ").upper()
                econ_parts.append(f"{lbl}: {v}")
        
        economic_block = ""
        if econ_parts:
            economic_block = "\nDATOS ECONÓMICOS CONFIRMADOS (USAR ESTOS VALORES SI EL DOCUMENTO LO REQUIERE):\n" + "\n".join(econ_parts)

        try:
            from app.services.mini_dictamen_anexos_service import (
                build_and_persist_mini_dictamen,
                build_stage_blocking_questions,
                get_blocking_annex_rows_for_stage,
            )

            await build_and_persist_mini_dictamen(self.context_manager.memory, session_id)
            fresh_state = await self.context_manager.memory.get_session(session_id) or session_state
            blocking_rows = get_blocking_annex_rows_for_stage(fresh_state, "technical")
            if blocking_rows:
                fresh_state["pending_questions"] = build_stage_blocking_questions(
                    "technical", blocking_rows
                ) + list(fresh_state.get("pending_questions") or [])
                fresh_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, fresh_state)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "La generación técnica quedó bloqueada por anexos obligatorios con "
                        "fuente inválida, referencial o pendiente de aclaración."
                    ),
                    data={"missing": blocking_rows},
                    correlation_id=correlation_id,
                )
        except Exception as exc:
            logger.warning(
                "technical_writer_mini_dictamen_guard_failed",
                session_id=session_id,
                error=str(exc),
            )

        # 2. SELECCIÓN DE REQUISITOS (Sincronización Total con Auditor + inventario canónico)
        tasks = session_state.get("tasks_completed", [])
        compliance_data: Dict[str, Any] = {}

        # a) Inyección directa del orquestador via compliance_master_list
        if agent_input.company_data and "compliance_master_list" in agent_input.company_data:
            compliance_data = agent_input.company_data["compliance_master_list"]

        # b) Resultados de Fase 1 via results.compliance.data
        if not compliance_data and agent_input.company_data and "results" in agent_input.company_data:
            results = agent_input.company_data["results"]
            if isinstance(results, dict) and "compliance" in results:
                compliance_data = results["compliance"].get("data", {})

        # c) Tarea persistida master_compliance_list en tasks_completed
        if not compliance_data:
            for task in reversed(tasks):
                if task.get("task") == "stage_completed:compliance":
                    res = task.get("result") or {}
                    if isinstance(res, dict) and res.get("data"):
                        compliance_data = res["data"]
                        break
        if not compliance_data:
            for task in reversed(tasks):
                if task.get("task") == "master_compliance_list":
                    res = task.get("result") or {}
                    compliance_data = res.get("data", res) if isinstance(res, dict) else {}
                    break

        tech_requirements = []
        action_counts: Dict[str, int] = {"generar": 0, "presentar_fisico": 0, "informativo": 0, "unknown": 0}
        # Sólo la zona "tecnico" tiene ítems redactables técnicamente.
        # Las zonas administrativo/formatos se gestionan por FormatsAgent.
        all_candidates = compliance_data.get("tecnico", [])

        from app.config.settings import settings
        from app.services.document_deliverable_filter import (
            count_actionable_generation_actions,
            is_generable_tipo_accion,
            is_technical_writer_queue_eligible,
            normalize_deliverable_key,
            should_show_deliverable_in_ui,
        )

        seen_ids: set[str] = set()
        seen_sigs: set[str] = set()
        for req in all_candidates:
            action = str(req.get("tipo_accion", "unknown") or "unknown").lower()
            if is_technical_writer_queue_eligible(req) and not is_generable_tipo_accion(action):
                action = "generar"
            if action not in action_counts:
                action = "unknown"
            action_counts[action] = action_counts.get(action, 0) + 1
            if not is_technical_writer_queue_eligible(req):
                continue
            nombre_u = str(req.get("nombre") or "")
            desc_u = str(req.get("descripcion") or "")
            if not should_show_deliverable_in_ui(
                nombre_u,
                desc_u,
                str(req.get("snippet") or ""),
                action,
            ):
                continue
            r_id = str(req.get("id", ""))
            if not r_id and req.get("nombre"):
                r_id = str(req.get("nombre", ""))[:80]
            sig = normalize_deliverable_key(nombre_u, "tecnico")
            if not _is_technical_writable(req) or r_id in seen_ids or sig in seen_sigs:
                continue
            tech_requirements.append(req)
            seen_ids.add(r_id)
            seen_sigs.add(sig)

        tech_requirements.sort(
            key=lambda r: (
                0 if str(r.get("tipo_accion") or "").lower() == "generar" else 1,
                0 if r.get("evidence_match") else 1,
                -len(str(r.get("snippet") or "")),
            )
        )

        from app.services.session_template_catalog import build_catalog_mirror_reqs

        catalog_tech = build_catalog_mirror_reqs(
            session_state,
            seen_ids,
            exclude_sobre=("administrativo", "economico"),
        )
        tech_requirements = catalog_tech + tech_requirements
        # Inventario canónico (Modo Fábrica) antes del gate: si hay plantillas PENDING,
        # deben contar como trabajo generable aunque compliance marque todo como unknown.
        tech_requirements = _merge_document_inventory_technical(
            company_data, tech_dir, tech_requirements, seen_ids
        )

        total_candidates = len(all_candidates)
        evidence_true = sum(1 for r in all_candidates if bool(r.get("evidence_match")))
        evidence_ratio = (evidence_true / total_candidates) if total_candidates else 1.0
        actionable_action_count = count_actionable_generation_actions(action_counts)
        gate = _should_block_by_quality_gate(
            total_items=total_candidates,
            generar_count=action_counts.get("generar", 0),
            unknown_count=action_counts.get("unknown", 0),
            evidence_match_ratio=evidence_ratio,
            presentar_fisico_count=action_counts.get("presentar_fisico", 0),
            triage_context=agent_input.triage_context,
            queued_requirements_count=len(tech_requirements),
            actionable_action_count=actionable_action_count,
        )
        if gate.get("block"):
            # Solo bloqueamos si no hay cola real ni acciones materializables.
            has_anything_to_do = (
                len(tech_requirements) > 0
                or actionable_action_count > 0
                or action_counts.get("presentar_fisico", 0) > 0
            )
            logger.warning(
                "technical_writer_quality_gate_triggered",
                session_id=session_id,
                reason=str(gate.get("reason")),
                metrics=gate.get("metrics"),
                continuing=has_anything_to_do,
            )
            if not has_anything_to_do:
                from app.services.document_quality_ux import (
                    build_document_quality_agent_pause_message,
                    build_document_quality_pending_question,
                )

                missing = [
                    build_document_quality_pending_question(
                        gate=gate,
                        session_state=company_data,
                        stage="technical",
                    )
                ]
                await self._save_pending_questions(session_id, missing)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=build_document_quality_agent_pause_message(stage="technical"),
                    data={"missing": missing, "document_quality_gate": gate},
                    correlation_id=correlation_id,
                )

        from app.services.ingested_file_resolver import (
            build_ingested_file_index,
            resolve_ingested_file,
        )
        from app.services.template_mirror_service import mirror_template_to_output

        session_documents: List[Dict[str, Any]] = []
        try:
            session_documents = await self.context_manager.memory.get_documents(session_id)
        except Exception as doc_exc:
            logger.warning("tech_get_documents_failed", session_id=session_id, error=str(doc_exc))

        file_index = build_ingested_file_index(session_documents)
        participation_zones = _infer_participation_zones(session_documents)
        contenido_nacional_pct = _infer_contenido_nacional_pct(vector_db, session_id)
        mirror_enabled = bool(getattr(settings, "TEMPLATE_MIRROR_ENABLED", True))
        mirror_max = int(getattr(settings, "TEMPLATE_MIRROR_MAX_ADMIN", 40) or 40)
        tech_mirror_queue: List[tuple] = []
        tech_llm_queue: List[Dict[str, Any]] = []
        unresolved_catalog_mirrors: List[str] = []
        session_hint = f"{session_id} {session_state.get('name', '')}"
        for req in tech_requirements:
            q = str(req.get("archivo_fuente") or req.get("nombre") or "")
            ref = resolve_ingested_file(
                q,
                file_index,
                doc_id=req.get("source_doc_id"),
                source_path=req.get("source_path"),
            )
            ext = (ref.file_path.rsplit(".", 1)[-1].lower() if ref and ref.file_path else "")
            if bool(req.get("from_session_catalog")) and not ref:
                unresolved_catalog_mirrors.append(str(req.get("nombre") or q or "sin_nombre"))
                logger.warning(
                    "technical_catalog_source_unresolved_skip_llm",
                    session_id=session_id,
                    requested=q,
                    source_doc_id=str(req.get("source_doc_id") or ""),
                    source_path=str(req.get("source_path") or ""),
                )
                continue
            if mirror_enabled and ref and ext in ("doc", "docx", "xls", "xlsx"):
                if _mirror_source_has_cross_tender_marker(ref, session_hint):
                    logger.warning(
                        "technical_mirror_cross_tender_fallback",
                        session_id=session_id,
                        requested=q,
                        source_filename=ref.filename,
                        from_session_catalog=bool(req.get("from_session_catalog")),
                    )
                    if bool(req.get("from_session_catalog")):
                        # La plantilla catalogada se descarta para espejado y dejamos que
                        # el inventario/documento técnico canónico tome el relevo.
                        continue
                    fallback_req = dict(req)
                    fallback_req["archivo_fuente"] = ""
                    fallback_req["cross_tender_mirror_skipped"] = True
                    fallback_req["mirror_conflict_source"] = ref.filename
                    tech_llm_queue.append(fallback_req)
                    continue
                tech_mirror_queue.append((req, ref))
            else:
                tech_llm_queue.append(req)

        if unresolved_catalog_mirrors:
            logger.warning(
                "technical_catalog_mirror_sources_unresolved",
                session_id=session_id,
                count=len(unresolved_catalog_mirrors),
                names=unresolved_catalog_mirrors[:10],
            )

        if mirror_max > 0 and len(tech_mirror_queue) > mirror_max:
            tech_mirror_queue = tech_mirror_queue[:mirror_max]

        max_tech = int(getattr(settings, "TECH_WRITER_MAX_GENERABLE_DOCS", 12) or 12)
        if max_tech > 0 and len(tech_llm_queue) > max_tech:
            tech_llm_queue = tech_llm_queue[:max_tech]

        tech_requirements = tech_llm_queue
        pre_mirror_files: List[Dict[str, Any]] = []
        profile_fill = {
            "rfc": rfc,
            "razon_social": razon_social,
            "representante_legal": representante,
            "domicilio": master_profile.get("domicilio_fiscal") or master_profile.get("domicilio"),
            "fecha": fecha_es,
            "licitacion": tender_name,
        }
        for idx_m, (req, ref) in enumerate(tech_mirror_queue, start=1):
            raw_name = req.get("nombre", "Documento")
            safe = re.sub(r"[^\w\s-]", "", str(raw_name).replace(" ", "_"))[:50].strip("_")
            src_ext = ref.file_path.rsplit(".", 1)[-1].lower()
            out_ext = ".docx" if src_ext == "doc" else f".{src_ext}"
            if out_ext not in (".docx", ".xlsx", ".xls"):
                out_ext = ".docx"
            filepath = os.path.join(tech_dir, f"mirror_{idx_m:02d}_{safe}{out_ext}")
            try:
                normalized_name = _normalized_blob(raw_name or ref.filename)
                if "contenido nacional" in normalized_name:
                    contenido_text = _build_contenido_nacional_text(
                        razon_social=razon_social,
                        rfc=rfc,
                        representante=representante,
                        session_id=session_id,
                        tender_name=tender_name,
                        fecha_es=fecha_es,
                        zonas=participation_zones,
                        porcentaje=contenido_nacional_pct,
                        destinatario=destinatario,
                    )
                    _save_docx(raw_name, contenido_text, filepath, doc_metadata)
                    meta = {
                        "ruta": filepath,
                        "mirror_mode": "deterministic_contenido_nacional",
                        "source_filename": ref.filename,
                    }
                else:
                    meta = mirror_template_to_output(ref, filepath, profile_fill, fill_profile=True)
                if str(meta.get("ruta") or filepath).lower().endswith(".docx"):
                    repair_docx_file_placeholders(str(meta.get("ruta") or filepath))
                pre_mirror_files.append(
                    attach_traceability(
                        {
                            "nombre": raw_name,
                            "ruta": meta["ruta"],
                            "status": "OK",
                            "tipo": "tecnico_mirror",
                        },
                        source_doc_id=str(ref.doc_id or req.get("source_doc_id") or "") or None,
                        source_filename=ref.filename,
                        source_path=meta.get("source_path") or ref.file_path,
                        source_hash=meta.get("source_hash"),
                        mirror_mode=meta.get("mirror_mode"),
                        materialization_route=meta.get("materialization_route") or "mirror",
                        output_hash=meta.get("output_hash") or safe_file_sha256(meta.get("ruta")),
                        provenance_ui=req.get("provenance_ui") if isinstance(req.get("provenance_ui"), dict) else None,
                    )
                )
            except Exception as mir_exc:
                logger.warning(
                    "technical_mirror_failed",
                    session_id=session_id,
                    source=ref.filename,
                    error=str(mir_exc),
                )
                tech_requirements.append(req)

        if not tech_requirements and not pre_mirror_files:
            tender_cat = str((agent_input.triage_context or {}).get("tender_category") or "").upper()
            if tender_cat == "OBRA":
                logger.info(
                    "technical_writer_obra_skip",
                    session_id=session_id,
                    reason="all_technical_items_are_presentar_fisico",
                    total_candidates=total_candidates,
                    presentar_fisico_count=action_counts.get("presentar_fisico", 0),
                )
                return AgentOutput(
                    status=AgentStatus.SUCCESS,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "Licitación de obra pública: los requisitos técnicos son formas predefinidas "
                        "(AT/AE) que se presentan físicamente. No hay documentos técnicos que redactar."
                    ),
                    data={"documentos": [], "obra_category_skip": True},
                    correlation_id=correlation_id,
                )
            return AgentOutput(
                status=AgentStatus.SUCCESS,
                agent_id=self.agent_id,
                session_id=session_id,
                message="No hay requisitos técnicos por redactar.",
                correlation_id=correlation_id
            )

        missing_slots = build_formats_pilot_missing_entries(
            master_profile,
            blocking_job_id=agent_input.job_id,
        )
        if missing_slots:
            field_keys = [m["field"] for m in missing_slots]
            logger.info(
                "technical_writer_pilot_blocked",
                agent_id=self.agent_id,
                session_id=session_id,
                correlation_id=correlation_id,
                blocking_job_id=agent_input.job_id,
                missing_fields=field_keys,
                missing_count=len(missing_slots),
            )
            await self._save_pending_questions(session_id, missing_slots)
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=f"Para redactar la propuesta técnica necesito: {', '.join([m['label'] for m in missing_slots])}",
                data={"missing": missing_slots},
                correlation_id=correlation_id,
            )

        generated_files = list(pre_mirror_files)
        descriptions_map = {}

        # 3. Generar CARTA DE PRESENTACIÓN (PRODUCCIÓN)
        print(f"[TechWriter] Redactando Carta de Presentación Real para {razon_social}...")
        carta_text = _build_carta_presentacion_text(
            razon_social=razon_social,
            rfc=rfc,
            representante=representante,
            domicilio=master_profile.get("domicilio_fiscal", "S/D"),
            tender_name=tender_name,
            fecha_es=fecha_es,
            destinatario=destinatario,
        )
        
        carta_path = os.path.join(tech_dir, "01_CARTA_PRESENTACION_PROPUESTA_TECNICA.docx")
        _save_docx("CARTA DE PRESENTACIÓN DE PROPUESTA TÉCNICA", carta_text, carta_path, doc_metadata)
        _carta_display = "Carta de Presentación de Propuesta Técnica.docx"
        generated_files.append(
            attach_traceability(
                {
                    "nombre": "Carta de Presentación",
                    "source_filename": _carta_display,
                    "ruta": carta_path,
                    "status": "OK",
                    "tipo": "tecnico_carta",
                    "template_id": "",
                },
                source_filename=_carta_display,
                template_id="",
                materialization_route="deterministic",
                output_hash=safe_file_sha256(carta_path),
            )
        )

        session_state = context.get("session_state") or {}
        session_hint = f"{session_id} {session_state.get('name', '')}"
        company_experience_block = await build_company_experience_context_block(
            self.context_manager.memory,
            session_id,
        )

        # 4. Generar documentos del Auditor
        for i, req in enumerate(tech_requirements, start=2):
            req_id = req.get("id", f"2.{i-1}")
            req_nombre = req.get("nombre", f"Requisito Técnico {i}")
            req_desc = req.get("descripcion", "")
            hint = req.get("generator_hint")
            hint_block = ""
            if hint:
                hint_block = f"\nGUÍA_DE_PLANTILLA_O_CLAVE: {hint}"

            print(f"[TechWriter] Generando documento final: {req_id} - {req_nombre}")

            # Buscar contexto específico del requisito en las bases
            # Construir query combinando ID de la forma + nombre + descripción para máxima precisión
            _rag_id_clean = req_id.replace("_", "-").replace(".", "-")
            req_context_res = vector_db.query_texts(
                session_id,
                f"Forma {_rag_id_clean} {req_nombre} {req_desc}",
                n_results=4,
            )
            req_context = ""
            if req_context_res.get("documents"):
                # Priorizar fragmentos que mencionan el ID de la forma explícitamente
                all_docs = req_context_res.get("documents", [])
                id_upper = _rag_id_clean.upper()
                # Primero los que contienen el ID de la forma
                priority = [d for d in all_docs if id_upper in (d or "").upper()]
                rest = [d for d in all_docs if id_upper not in (d or "").upper()]
                ordered = priority + rest
                req_context = "\n".join(d for d in ordered[:2] if d and d.strip())

            doc_prompt = f"""Redacta el documento oficial ESPECÍFICO para el siguiente requisito de licitación:

REQUISITO: {req_id} - {req_nombre}
DESCRIPCIÓN: {req_desc}

EMPRESA: {razon_social}
RFC: {rfc}
REPRESENTANTE LEGAL: {representante}
DOMICILIO: {master_profile.get('domicilio_fiscal', 'S/D')}
LICITACIÓN: {tender_name}
FECHA CANÓNICA DEL EXPEDIENTE (usar únicamente esta en el documento): {fecha_es}

CONTEXTO DE LAS BASES (fragmento literal donde se describe este requisito — úsalo como fuente de verdad para la estructura y contenido del documento):
{req_context[:1500] if req_context else "Ver bases de licitación"}

INSTRUCCIONES CRÍTICAS:
- El CONTEXTO DE LAS BASES define QUÉ debe contener este documento. Léelo primero.
- Si el contexto pide currículum empresarial o relación de clientes: NO inventes filas con corchetes ([Domicilio], [Teléfono], etc.). Escribe un párrafo indicando que se anexan contratos y referencias en el expediente de la empresa, o lista solo clientes reales si constan en el contexto.
- Si el contexto describe una DECLARACIÓN o MANIFESTACIÓN bajo protesta, genera esa declaración.
- Si el contexto describe un PROGRAMA o CALENDARIO, genera esa estructura.
- NO generes texto sobre "solicitudes de aclaración", "junta de aclaraciones" ni "responsabilidad del licitante" a menos que el contexto lo indique explícitamente para ESTE documento.
- NO uses placeholders como [Nombre], [Fecha], [Descripción], [COMPLETAR] — usa los datos reales de la empresa o omite la fila.
- El documento debe ser directamente presentable en la licitación.
- Bajo protesta de decir verdad cuando aplique.
- Firma: {representante}, {razon_social}{hint_block}"""

            # Inyectar datos económicos solo si el documento los necesita explícitamente
            # (documentos de costos, análisis de precios, propuesta económica)
            _req_blob_upper = f"{req_id} {req_nombre} {req_desc}".upper()
            _needs_economic = any(k in _req_blob_upper for k in (
                "PRECIO", "COSTO", "IMPORTE", "ECONÓMIC", "ECONOMIC", "FASAR",
                "SALARIO", "IMSS", "INFONAVIT", "AGUINALDO", "PRESUPUEST",
                "ANÁLISIS DE PRECIO", "ANALISIS DE PRECIO",
            ))
            if _needs_economic and economic_block:
                doc_prompt += f"\n\nDATOS ECONÓMICOS CONFIRMADOS (usar en este documento):\n{economic_block}"

            if req_needs_company_experience(req_id, req_nombre, req_desc) and company_experience_block:
                doc_prompt += f"\n\n{company_experience_block}"

            req_id_upper = str(req_id).upper()
            req_nombre_upper = str(req_nombre).upper()
            if "TE-12" in req_id_upper or "TE-12" in req_nombre_upper:
                doc_text = _build_te12_text(
                    razon_social=razon_social,
                    rfc=rfc,
                    representante=representante,
                    domicilio=master_profile.get("domicilio_fiscal", "S/D"),
                    tender_name=tender_name,
                    fecha_es=fecha_es,
                    req_context=req_context,
                    destinatario=destinatario,
                )
            elif is_primary_technical_proposal(req_id, req_nombre, req_desc):
                exp_block = ""
                if req_needs_company_experience(req_id, req_nombre, req_desc) and company_experience_block:
                    exp_block = company_experience_block
                doc_text = build_propuesta_tecnica_body(
                    razon_social=razon_social,
                    rfc=rfc,
                    representante=representante,
                    domicilio=str(_dom_fiscal),
                    tender_name=tender_name,
                    req_nombre=req_nombre,
                    req_desc=req_desc,
                    req_context=req_context,
                    experience_block=exp_block,
                )
            else:
                resp = await llm.generate(prompt=doc_prompt, system_prompt=system_prompt, correlation_id=correlation_id)
                doc_text = resp.response if resp.success else f"Contenido para {req_nombre}"

            doc_text = normalize_body_spanish_dates(doc_text or "", fecha_es)
            from app.services.document_contamination_gate import strip_llm_meta_leaks

            doc_text = strip_llm_meta_leaks(doc_text)

            # Transliterar acentos y caracteres especiales antes de sanitizar.
            # unicodedata.normalize('NFD') descompone 'á' → 'a' + combining accent,
            # luego el encode/decode ASCII elimina solo los diacríticos.
            _nombre_nfd = unicodedata.normalize('NFD', req_nombre)
            _nombre_ascii = _nombre_nfd.encode('ascii', 'ignore').decode('ascii')
            safe_nombre = re.sub(r'[^a-zA-Z0-9\s]', '', _nombre_ascii)[:50].strip().replace(" ", "_")
            file_path = os.path.join(tech_dir, f"{i:02d}_{req_id.replace('.','_')}_{safe_nombre}.docx")

            route = "deterministic_propuesta_tecnica" if is_primary_technical_proposal(
                req_id, req_nombre, req_desc
            ) else ("deterministic_te12" if "TE-12" in req_id_upper or "TE-12" in req_nombre_upper else "llm")
            _save_docx(req_nombre, doc_text, file_path, doc_metadata)
            _te_display = f"{req_id}: {req_nombre}.docx" if req_id else f"{req_nombre}.docx"
            generated_files.append(
                attach_traceability(
                    {
                        "nombre": f"{req_id}: {req_nombre}",
                        "source_filename": _te_display,
                        "ruta": file_path,
                        "status": "OK",
                        "tipo": "tecnico",
                        "template_id": "",
                    },
                    source_doc_id=str(req.get("source_doc_id") or "") or None,
                    source_filename=str(req.get("archivo_fuente") or _te_display).strip() or _te_display,
                    source_path=str(req.get("source_path") or "") or None,
                    source_hash=safe_file_sha256(str(req.get("source_path") or "")),
                    template_id="",
                    materialization_route=route,
                    output_hash=safe_file_sha256(file_path),
                    provenance_ui=req.get("provenance_ui") if isinstance(req.get("provenance_ui"), dict) else None,
                )
            )
            descriptions_map[os.path.basename(file_path)] = req_desc

        for gf in generated_files:
            gpath = str(gf.get("ruta") or "")
            if gpath.lower().endswith(".docx") and os.path.exists(gpath):
                repair_docx_file_placeholders(gpath)

        session_docs = await self.context_manager.memory.get_documents(session_id) or []
        client_refs = extract_client_references_from_documents(session_docs)
        try:
            from app.services.document_catalog_service import experience_client_refs_from_catalog

            catalog_refs = experience_client_refs_from_catalog(session_state)
            if catalog_refs:
                client_refs = catalog_refs
        except Exception:
            pass
        experience_summary = build_experience_sources_ux_summary(session_docs, session_state)
        if client_refs:
            for gf in generated_files:
                gpath = str(gf.get("ruta") or "")
                if gpath.lower().endswith(".docx") and os.path.exists(gpath):
                    if fill_te03_client_placeholders(gpath, client_refs):
                        repair_docx_file_placeholders(gpath)

        result_data = {
            "titulo": "Propuesta Técnica Completa",
            "folder": tech_dir,
            "documentos": generated_files,
            "descriptions": descriptions_map,
            "action_type_stats": action_counts,
            "materialization_metrics": build_materialization_metrics(
                stage="technical",
                documents=generated_files,
                elapsed_ms=int((time.perf_counter() - started_at) * 1000),
            ),
        }
        if experience_summary:
            result_data["experience_sources_ux"] = experience_summary

        fill_gate = validate_generated_documents_fill(
            stage="technical",
            generated_documents=generated_files,
            master_profile=master_profile,
            provenance_context={"source": "technical_writer", "confidence": 0.9, "session_hint": session_hint},
        )
        result_data["document_fill_quality_gate"] = fill_gate
        result_data["validation_events"] = [
            build_fill_validation_event(it, stage="technical")
            for it in (fill_gate.get("issues") or [])
            if isinstance(it, dict)
        ]
        if not bool(fill_gate.get("validation_passed", True)):
            from app.services.document_fill_ux_messages import build_fill_blocking_question

            company_name = str(master_profile.get("razon_social") or "").strip()
            human_question = build_fill_blocking_question(
                "technical",
                fill_gate.get("issues") or [],
                company_name=company_name,
                experience_summary=experience_summary,
            )
            missing = [
                {
                    "field": "document_fill_quality_gate",
                    "label": "Completar datos de la empresa",
                    "question": human_question,
                    "document_hint": f"Gate llenado: {fill_gate.get('metrics')}",
                    "type": "document_fill_quality_gate_blocking",
                    "blocking_items": fill_gate.get("issues") or [],
                }
            ]
            await self._save_pending_questions(session_id, missing)
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=human_question,
                data={**result_data, "missing": missing, "stage": "technical"},
                correlation_id=correlation_id,
            )

        # Persistir metadatos de descripción
        meta_path = os.path.join(base_output_dir, "descriptions.json")
        try:
            existing_meta = {}
            if os.path.exists(meta_path):
                with open(meta_path, 'r', encoding='utf-8') as f:
                    existing_meta = json.load(f)
            existing_meta.update(descriptions_map)
            with open(meta_path, 'w', encoding='utf-8') as f:
                json.dump(existing_meta, f, indent=4, ensure_ascii=False)
        except Exception as e:
            logger.warning("metadata_persist_failed", session_id=session_id, error=str(e))

        await self.context_manager.record_task_completion(session_id, "technical_writing_COMPLETED", result_data)
        return AgentOutput(
            status=AgentStatus.SUCCESS,
            agent_id=self.agent_id,
            session_id=session_id,
            data=result_data,
            correlation_id=correlation_id
        )

    async def _save_pending_questions(self, session_id: str, missing_fields: List[Dict[str, Any]]) -> None:
        """Persiste preguntas HITL para el chatbot (mismo contrato que FormatsAgent)."""
        try:
            session_state = await self.context_manager.memory.get_session(session_id)
            if session_state:
                session_state["pending_questions"] = missing_fields
                session_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, session_state)
        except Exception as e:
            logger.error(
                "technical_writer_save_questions_failed",
                session_id=session_id,
                error=str(e),
            )


def _save_docx(title: str, content: str, file_path: str, metadata: dict = None):
    doc = docx.Document()
    section = doc.sections[0]
    
    # Márgenes estándar
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Header: Logo y Datos de Licitación
    header = section.header
    htable = header.add_table(1, 2, Inches(6.5))
    htable.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Celda 1: Logo
    if metadata and metadata.get("logo_path"):
        from app.utils.doc_formatting import add_logo_picture_to_run

        if not add_logo_picture_to_run(
            htable.cell(0, 0).paragraphs[0].add_run(),
            str(metadata["logo_path"]),
            width_inches=1.5,
        ):
            logger.warning("logo_insert_failed", path=metadata.get("logo_path"))
            
    # Celda 2: Datos (Derecha)
    p_info = htable.cell(0, 1).paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if metadata:
        run = p_info.add_run(f"{metadata.get('tender_name', 'LICITACIÓN').upper()}\n")
        run.bold = True
        run.font.size = Pt(9)
        fecha_hdr = str(
            metadata.get("fecha_encabezado")
            or metadata.get("fecha_generacion")
            or metadata.get("fecha")
            or ""
        ).strip()
        if not fecha_hdr:
            from app.services.document_date_resolver import resolve_generation_header_date

            fecha_hdr = resolve_generation_header_date()["fecha_es"]
        run_date = p_info.add_run(f"Fecha: {fecha_hdr}")
        run_date.font.size = Pt(8)

    # Pie de Página
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if metadata:
        run_foot = p_foot.add_run(f"{metadata.get('footer_text', '')}")
        run_foot.font.size = Pt(7)
        run_foot.italic = True

    # Cuerpo del Documento
    doc.add_heading(title.upper(), 0)

    body = strip_bracket_placeholders_for_docx(strip_markdown_for_docx(content or ""))
    body_low = body.lower()[:500]
    has_header = "lugar y fecha" in body_low or "presente" in body_low
    if not has_header:
        footer_text = metadata.get("footer_text", "") if metadata else ""
        lugar = (
            footer_text.split("Domicilio:")[-1].split(",")[0].strip()
            if "Domicilio:" in footer_text
            else "México"
        )
        fecha_hdr = str(
            (metadata or {}).get("fecha_encabezado")
            or (metadata or {}).get("fecha_generacion")
            or (metadata or {}).get("fecha")
            or ""
        ).strip()
        if not fecha_hdr:
            from app.services.document_date_resolver import resolve_generation_header_date

            fecha_hdr = resolve_generation_header_date()["fecha_es"]
        p_fecha = doc.add_paragraph(f"LUGAR Y FECHA: {lugar} a {fecha_hdr}")
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dest = (metadata or {}).get("destinatario") or "A QUIEN CORRESPONDA:"
        p_dest = doc.add_paragraph(f"\n{dest}")
        p_dest.bold = True
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para in body.split("\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    # Firma Final
    doc.add_paragraph("\n\n")
    p_atentamente = doc.add_paragraph("ATENTAMENTE\n")
    p_atentamente.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_firma = doc.add_paragraph("___________________________\n")
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_firma = p_firma.add_run(f"{metadata.get('representante', '').upper()}")
    run_firma.bold = True
    
    p_cargo = doc.add_paragraph("REPRESENTANTE LEGAL")
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(file_path)
