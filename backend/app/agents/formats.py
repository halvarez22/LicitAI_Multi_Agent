import os
import re
import json
import time
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Any, Dict, List, Optional, Set

from app.agents.base_agent import BaseAgent
from app.agents.mcp_context import MCPContextManager
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.contracts.document_inventory import (
    DocumentEnvelope,
    DocumentInventory,
    InventoryItemStatus,
)
from app.services.document_fill_quality_gate import (
    detect_cross_tender_marker,
    validate_generated_documents_fill,
)
from app.services.document_fill_ux_messages import build_fill_validation_event
from app.core.formats_pilot_slots import (
    build_formats_pilot_missing_entries,
    is_usable_profile_field_value,
)
from app.core.observability import get_logger
from app.services.document_traceability import (
    attach_traceability,
    build_materialization_metrics,
    safe_file_sha256,
)
from app.utils.doc_formatting import (
    ANTI_PLACEHOLDER_PROMPT_RULE,
    CONCURSANTE_LEXICON_SYSTEM_RULE,
    LEGAL_AUTHORIZATION_SYSTEM_RULE,
    strip_markdown_for_docx,
    is_markdown_table_line,
    parse_markdown_table,
)
from app.services.document_date_resolver import (
    normalize_body_spanish_dates,
    normalize_docx_spanish_dates,
    resolve_addressee_lines,
    resolve_document_date,
)
from app.services.document_contamination_gate import is_apu_document, strip_llm_meta_leaks
from app.core.template_engine import LegalTemplateEngine, TemplateIntegrityError
from app.services.resilient_llm import ResilientLLMClient
from app.services.vector_service import VectorDbServiceClient
from app.config.settings import settings as app_settings

logger = get_logger(__name__)


def _formats_inventory_doc_path(output_dir: str, canonical_id: str) -> str:
    """Ruta del .docx ya materializado para este canonical_id, si existe."""
    if not canonical_id or not os.path.isdir(output_dir):
        return ""
    token = re.sub(r"[^\w\-]+", "_", canonical_id).strip("_").lower()
    if len(token) < 3:
        return ""
    try:
        for fn in os.listdir(output_dir):
            if not fn.endswith(".docx") or fn.startswith("~$"):
                continue
            if token in fn.lower():
                return os.path.join(output_dir, fn)
    except OSError:
        return ""
    return ""


def _formats_inventory_doc_exists(output_dir: str, canonical_id: str) -> bool:
    """Idempotencia: no regenerar si ya existe .docx que incluye el canonical_id."""
    return bool(_formats_inventory_doc_path(output_dir, canonical_id))


def _merge_document_inventory_legal(
    company_data: Dict[str, Any],
    output_dir: str,
    reqs_to_process: List[Dict[str, Any]],
    seen_ids: Set[str],
) -> None:
    """
    Modo fábrica: añade ítems ``legal_administrative`` pendientes del inventario canónico.

    No duplica ``id`` ya visto en compliance ni archivos ya materializados en disco.
    Solo agrega ítems con evidencia real (Tier A) o explícitamente añadidos por el usuario
    (Tier C). Los ítems Tier B (inferidos) sin anchors se descartan para evitar alucinaciones.
    """
    raw = company_data.get("document_inventory")
    if not isinstance(raw, dict) or not raw.get("items"):
        return
    try:
        inv = DocumentInventory.model_validate(raw)
    except Exception as e:
        logger.warning("formats_inventory_parse_failed", error=str(e))
        return

    for it in inv.items:
        if it.category != DocumentEnvelope.LEGAL:
            continue
        if it.status != InventoryItemStatus.PENDING:
            continue
        cid = (it.canonical_id or "").strip().replace(".", "_")
        if not cid:
            continue
        key = cid.lower()
        if key in seen_ids:
            continue
        if _formats_inventory_doc_exists(output_dir, it.canonical_id):
            continue

        # ── FILTRO ANTI-ALUCINACIÓN (document_inventory) ────────────────────
        # Ítems Tier B (inferidos) sin anchors reales son candidatos a alucinación.
        # Solo aceptamos:
        #   - Tier A (anchored): tienen evidencia literal en las bases
        #   - Tier C (user_added): el usuario los agregó explícitamente
        #   - Tier B con al menos un anchor con snippet >= 20 chars
        tier_val = str(getattr(it, "tier", "") or "").lower()
        if tier_val == "inferred":
            anchors = list(getattr(it, "anchors", []) or [])
            has_real_anchor = any(
                len(str(getattr(a, "snippet", "") or "").strip()) >= 20
                for a in anchors
            )
            if not has_real_anchor:
                logger.info(
                    "formats_inventory_inferred_no_anchor_discarded",
                    canonical_id=cid,
                    display_name=str(it.display_name or "")[:80],
                )
                seen_ids.add(key)
                continue

        seen_ids.add(key)
        reqs_to_process.append(
            {
                "id": cid,
                "nombre": it.display_name,
                "descripcion": (it.description or "").strip(),
                "tipo": "formato",
                "from_document_inventory": True,
                "generator_hint": it.generator_hint,
            }
        )


def _resolve_metadata_domicilio(metadata: Dict[str, Any]) -> str:
    """Domicilio usable desde metadata o pie de página (sin S/D ni vacíos)."""
    raw = str(metadata.get("domicilio") or "").strip()
    if is_usable_profile_field_value(raw):
        return raw
    footer_text = str(metadata.get("footer_text") or "")
    if "Domicilio:" in footer_text:
        raw = footer_text.split("Domicilio:", 1)[1].strip()
    return raw if is_usable_profile_field_value(raw) else ""


def _sanitize_legal_content(
    content: str,
    *,
    session_id: str,
    metadata: Dict[str, Any],
) -> str:
    """
    Sanea placeholders frecuentes de plantillas legales antes de guardar DOCX.

    Regla: ningún documento final debe contener marcadores entre corchetes ni
    tokens genéricos tipo "N/A" como dato principal. No inyecta el token
    «Dato pendiente…» (el gate de llenado lo trataría como bloqueo).
    """
    text = (content or "").strip()
    if not text:
        return text

    razon_social = str(metadata.get("empresa") or "la empresa").strip()
    representante = str(metadata.get("representante") or "representante legal").strip()
    rfc = str(metadata.get("rfc") or "").strip()
    fecha = str(metadata.get("fecha") or "").strip()
    hora = str(metadata.get("hora") or datetime.now().strftime("%H:%M")).strip()
    domicilio = _resolve_metadata_domicilio(metadata)
    ciudad = domicilio.split(",")[0].strip() if domicilio else "México"
    pending_slot = "________________________"

    replacements = {
        "[Dirección de la empresa]": domicilio or pending_slot,
        "[Ciudad, Estado, Código Postal]": domicilio or pending_slot,
        "[Fecha actual]": fecha or pending_slot,
        "[Fecha]": fecha or pending_slot,
        "[Hora]": hora,
        "[Nombre del Representante Legal/Apoderado]": representante,
        "[Nombre del Representante Legal o Destinatario]": representante,
        "[Nombre del Destinatario]": str(
            metadata.get("destinatario") or "A QUIEN CORRESPONDA:"
        ).split("\n")[0].strip()
        or "A QUIEN CORRESPONDA:",
        "[Nombre completo del concursante]": razon_social,
        "[Número de Licitación o Nombre del Proceso]": session_id,
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)

    # Paginación y plazos típicos de plantillas (evita sustituir por texto bloqueante).
    text = re.sub(
        r"\[?\s*p[aá]gina\s+[^\]\n]{0,80}\]?",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\[[^\]]{0,120}\bmeses\b[^\]]*\]",
        "12 (doce) meses",
        text,
        flags=re.IGNORECASE,
    )

    # Reemplazos semánticos de fallback comunes.
    text = re.sub(r"\bRFC:\s*N/A\b", f"RFC: {rfc or pending_slot}", text, flags=re.I)
    text = re.sub(r"\bRepresentante\s+Legal:\s*N/A\b", f"Representante Legal: {representante}", text, flags=re.I)
    text = re.sub(r"\bLugar y fecha:\s*N/A\b", f"Lugar y fecha: {ciudad}, {fecha or pending_slot}", text, flags=re.I)
    if domicilio:
        text = text.replace(
            "domicilio en Dato pendiente de confirmar por el representante legal.",
            f"domicilio en {domicilio}.",
        )
    if representante:
        text = re.sub(
            r"(Firma del Representante Legal:?\s*)Dato pendiente de confirmar por el representante legal\.",
            rf"\1{representante}",
            text,
            flags=re.I,
        )

    # Corchetes/llaves residuales → ranura visible (HITL), no token que dispare gate.
    text = re.sub(r"\[[^\]]+\]|\{[^}]+\}", pending_slot, text)
    text = re.sub(r"\s{2,}", " ", text)
    return text


def _normalize_formats_blob(value: Any) -> str:
    return str(value or "").strip().lower()


def _infer_formats_contenido_nacional_pct(
    vector_db: VectorDbServiceClient,
    session_id: str,
) -> str:
    try:
        res = vector_db.query_texts(session_id, "contenido nacional 65% anexo iii-k punto 46", n_results=4)
        for doc in res.get("documents") or []:
            text = str(doc or "")
            if "contenido nacional" not in text.lower():
                continue
            match = re.search(r"(\d{1,3})\s*%", text)
            if match:
                return match.group(1)
    except Exception as exc:
        logger.warning("formats_contenido_nacional_pct_failed", session_id=session_id, error=str(exc))
    return ""


def _build_formats_contenido_nacional_text(
    *,
    razon_social: str,
    rfc: str,
    representante: str,
    session_id: str,
    tender_name: str,
    fecha_es: str,
    zonas: List[str],
    porcentaje: str,
    destinatario: str = "",
) -> str:
    zonas_txt = ", ".join(zonas) if zonas else "correspondientes a la propuesta presentada"
    if len(zonas) == 1:
        zona_clause = f"la zona {zonas_txt}"
    elif len(zonas) > 1:
        zona_clause = f"las zonas {zonas_txt}"
    else:
        zona_clause = "la zona en que participo"
    porcentaje_txt = porcentaje or "65"
    proc_ref = tender_name.strip() or session_id.replace("_", " ").upper()
    dest_block = (destinatario or "A QUIEN CORRESPONDA:").strip()
    return (
        "MANIFESTACIÓN DE NACIONALIDAD MEXICANA, PRODUCCIÓN EN MÉXICO Y GRADO DE CONTENIDO NACIONAL\n\n"
        f"{fecha_es}\n\n"
        f"{dest_block}\n\n"
        f"Me refiero al procedimiento de contratación identificado como {proc_ref}, "
        f"en el que mi representada, la empresa {razon_social}, RFC {rfc}, participa a través de la propuesta contenida en el presente sobre.\n\n"
        "Sobre el particular y en términos de lo previsto por el Acuerdo por el que se establecen las reglas para la determinación "
        "del grado de contenido nacional, tratándose de procedimientos de contratación de carácter nacional, el suscrito manifiesta "
        "bajo protesta de decir verdad que los bienes ofertados para la partida 2 correspondiente a "
        f"{zona_clause}, serán producidos en México y contendrán un grado de contenido nacional de cuando menos el {porcentaje_txt} por ciento, "
        "en el supuesto de que sea adjudicado el pedido respectivo. Asimismo, manifiesto bajo protesta de decir verdad ser de nacionalidad mexicana.\n\n"
        "De igual forma, manifiesto tener conocimiento de lo previsto en el artículo 57 de la Ley de Adquisiciones, "
        "Arrendamientos y Servicios del Sector Público y me comprometo, en caso de ser requerido, a aceptar una verificación "
        "del cumplimiento de los requisitos sobre el contenido nacional de los bienes ofertados mediante la exhibición de la "
        "información documental correspondiente y/o a través de una inspección física de la planta industrial en la que se producen.\n\n"
        "ATENTAMENTE\n\n"
        f"{razon_social}\n\n"
        f"{representante}\n"
        "Representante Legal"
    )


def _pick_reference_service_price(user_inputs: Dict[str, Any]) -> Optional[float]:
    candidates: List[tuple[str, float]] = []

    def _try_float(value: Any) -> Optional[float]:
        if value in (None, ""):
            return None
        try:
            parsed = float(value)
            return parsed if parsed > 0 else None
        except (TypeError, ValueError):
            return None

    def _collect_from_mapping(mapping: Any) -> None:
        if not isinstance(mapping, dict):
            return
        for key, value in mapping.items():
            k = str(key or "")
            if k.startswith("price_struct_service_"):
                parsed = _try_float(value)
                if parsed is not None:
                    candidates.append((k, parsed))
                continue
            if re.search(r"(tarifa|precio|unit_price|price)", k, flags=re.IGNORECASE):
                parsed = _try_float(value)
                if parsed is not None:
                    candidates.append((k, parsed))

    for direct_key in ("tarifa_mensual", "tarifa_mensual_referencia", "reference_unit_price"):
        parsed = _try_float(user_inputs.get(direct_key))
        if parsed is not None:
            return parsed

    _collect_from_mapping(user_inputs.get("concept_prices"))
    _collect_from_mapping(user_inputs)
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0])
    return candidates[0][1]


async def _pick_reference_service_price_from_line_items(
    memory: Any,
    session_id: str,
) -> Optional[float]:
    """Fallback: primera partida tabular con precio unitario > 0."""
    try:
        rows = await memory.get_line_items_for_session(session_id) or []
    except Exception:
        return None
    for row in rows:
        if not isinstance(row, dict):
            continue
        for key in ("unit_price", "precio_unitario", "price", "tarifa_mensual", "total_price"):
            parsed = None
            try:
                val = row.get(key)
                if val not in (None, ""):
                    parsed = float(val)
            except (TypeError, ValueError):
                parsed = None
            if parsed is not None and parsed > 0:
                return parsed
    return None


def _zones_from_economic_user_inputs(user_inputs: Dict[str, Any]) -> List[str]:
    """Extrae zonas A/B/C… desde claves ``price_struct_service_{zona}_``."""
    zones: List[str] = []
    sources: List[Any] = [user_inputs]
    cp = user_inputs.get("concept_prices")
    if isinstance(cp, dict):
        sources.append(cp)
    for mapping in sources:
        if not isinstance(mapping, dict):
            continue
        for key in mapping:
            m = re.search(r"price_struct_service_([a-z])_", str(key or "").lower())
            if not m:
                continue
            z = m.group(1).upper()
            if z not in zones:
                zones.append(z)
    return zones


def _mirror_source_has_cross_tender_marker(ref: Any, session_hint: str) -> bool:
    text = str(getattr(ref, "extracted_text", "") or "")
    return bool(detect_cross_tender_marker([text], session_hint))


def _build_panel_authoritative_reqs(panel_payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Cola «generar» alineada 1:1 con el panel UI (sin dedupe contra compliance).

    Restaura cantidad y nombres de anexos cuando la lista de compliance quedó
    demasiado filtrada.
    """
    from app.services.pliego_formats_enrichment_service import pliego_format_dedupe_key
    from app.services.document_deliverable_filter import is_economic_writer_domain

    panel_reqs: List[Dict[str, Any]] = []
    panel_seen: Set[str] = set()
    for bucket in (
        "sobre_1_tecnico",
        "sobre_2_economico",
        "requisitos_legales",
        "otros_requisitos_criticos",
    ):
        for row in panel_payload.get(bucket) or []:
            if not isinstance(row, dict):
                continue
            if str(row.get("tipo_accion_final") or row.get("tipo") or "") != "generar":
                continue
            nombre_panel = str(
                row.get("nombre_canonico") or row.get("nombre") or ""
            ).strip()
            if not nombre_panel:
                continue
            nombre_low = nombre_panel.lower()
            if re.search(
                r"modelo.*propuesta\s+t[eé]cnica|propuesta\s+t[eé]cnica.*modelo",
                nombre_low,
            ):
                continue
            desc_panel = str(row.get("descripcion") or "")
            snippet_panel = str(
                row.get("snippet_representativo") or row.get("snippet") or ""
            )
            if is_economic_writer_domain(nombre_panel, desc_panel, snippet_panel):
                continue
            if is_apu_document(nombre_panel, desc_panel, ""):
                continue
            dedupe_key = pliego_format_dedupe_key(nombre_panel)
            if dedupe_key in panel_seen:
                continue
            panel_seen.add(dedupe_key)
            safe_id = re.sub(r"[^\w-]", "_", dedupe_key)[:48].strip("_") or "anexo"
            panel_reqs.append(
                {
                    "id": f"panel_{safe_id}",
                    "nombre": nombre_panel,
                    "descripcion": desc_panel,
                    "snippet": snippet_panel,
                    "tipo_accion": "generar",
                    "from_document_inventory": True,
                    "inventory_synthetic": True,
                    "evidence_match": True,
                    "from_formats_panel": True,
                    "panel_dedupe_key": dedupe_key,
                }
            )
    return panel_reqs


def _should_block_by_quality_gate(
    *,
    total_items: int,
    generar_count: int,
    unknown_count: int,
    evidence_match_ratio: float,
    presentar_fisico_count: int = 0,
    triage_context: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Gate duro para frenar sobre-generación administrativa por clasificación débil.

    Excepción OBRA: en licitaciones de obra pública (LOPSRM / categoría OBRA),
    los formatos son predefinidos (AT/AE) que el licitante llena y presenta
    físicamente. Si generar_count == 0 pero hay ítems presentar_fisico, no es
    un error de clasificación — es el comportamiento esperado.
    """
    if not bool(app_settings.DOCUMENT_QUALITY_HARD_GATE_ENABLED):
        return {"block": False, "reason": "", "metrics": {}}

    # ── Excepción por categoría de licitación ──────────────────────────────
    tender_category = ""
    if isinstance(triage_context, dict):
        tender_category = str(triage_context.get("tender_category") or "").upper()

    if tender_category == "OBRA":
        if generar_count == 0 and presentar_fisico_count > 0:
            return {
                "block": False,
                "reason": "obra_category_no_generate_items_expected",
                "metrics": {
                    "total_items": total_items,
                    "generar_count": generar_count,
                    "presentar_fisico_count": presentar_fisico_count,
                    "tender_category": tender_category,
                    "evidence_match_ratio": evidence_match_ratio,
                },
            }
        if total_items == 0:
            return {"block": False, "reason": "", "metrics": {"tender_category": tender_category}}

    min_items = max(1, int(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MIN_ITEMS", 3) or 3))
    max_unknown = float(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MAX_UNKNOWN_RATIO", 0.6) or 0.6)
    min_evidence = float(getattr(app_settings, "DOCUMENT_QUALITY_GATE_MIN_EVIDENCE_MATCH_RATIO", 0.5) or 0.5)
    if total_items < min_items:
        return {"block": False, "reason": "", "metrics": {}}
    unknown_ratio = (unknown_count / total_items) if total_items else 0.0
    if generar_count == 0:
        return {
            "block": True,
            "reason": "no_actionable_generate_items",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "unknown_ratio": unknown_ratio,
                "evidence_match_ratio": evidence_match_ratio,
                "tender_category": tender_category or None,
            },
        }
    if unknown_ratio > max_unknown:
        return {
            "block": True,
            "reason": "unknown_ratio_above_threshold",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "unknown_ratio": unknown_ratio,
                "threshold_unknown_ratio": max_unknown,
                "evidence_match_ratio": evidence_match_ratio,
                "tender_category": tender_category or None,
            },
        }
    if evidence_match_ratio < min_evidence:
        return {
            "block": True,
            "reason": "evidence_match_ratio_below_threshold",
            "metrics": {
                "total_items": total_items,
                "generar_count": generar_count,
                "unknown_ratio": unknown_ratio,
                "evidence_match_ratio": evidence_match_ratio,
                "threshold_evidence_match_ratio": min_evidence,
                "tender_category": tender_category or None,
            },
        }
    return {"block": False, "reason": "", "metrics": {}}


class FormatsAgent(BaseAgent):
    """
    Agente 5: Generador de Formatos Finales (Deduplicación de Coronación).
    Lógica blindada para extraer el 100% de la lista del ComplianceAgent.
    """
    def __init__(self, context_manager: MCPContextManager):
        super().__init__(
            agent_id="formats_001",
            name="Generador de Formatos",
            description="Generador oficial de documentos administrativos (1.x).",
            context_manager=context_manager
        )
        # Instanciado en constructor para que sea mockeable en tests unitarios
        self.llm = ResilientLLMClient()
        self.template_engine = LegalTemplateEngine()
        self.vector_db = VectorDbServiceClient()

    @staticmethod
    def _template_id_for_requirement(req: Dict[str, Any]) -> str | None:
        """Mapea un requisito de formato a template legal bloqueado."""
        rid = str(req.get("id", "")).strip().lower()
        name = str(req.get("nombre", "")).strip().lower()
        desc = str(req.get("descripcion", "")).strip().lower()
        label_tax = str(req.get("label_taxonomica", "")).strip().lower()
        text = f"{rid} {name} {desc} {label_tax}"
        if "anexo 7" in text or "personalidad" in text:
            return "anexo_7"
        if "anexo 11" in text or "conformidad" in text:
            return "anexo_11"
        if "anexo 15" in text or "50" in text or "60" in text:
            return "anexo_15"
        if "decl_integridad" in text or "no colusion" in text or "colusión" in text:
            return "decl_integridad_no_colusion"
        if "decl_mipyme" in text or "mipyme" in text or "estratificacion" in text or "estratificación" in text:
            return "decl_mipyme"
        return None

    def _template_data(self, session_id: str, master_profile: Dict[str, Any], metadata: Dict[str, Any], economic_overrides: Dict[str, Any] = None) -> Dict[str, Any]:
        """Construye datos dinámicos para render de templates legales."""
        domicilio = master_profile.get("domicilio_fiscal") or master_profile.get("domicilio") or ""
        lugar = (
            master_profile.get("ciudad")
            or (str(domicilio).split("|", 1)[0].strip() if domicilio else "")
            or "México"
        )
        actividad_principal = (
            master_profile.get("actividad_principal")
            or master_profile.get("giro")
            or "prestación de servicios"
        )
        data = {
            "razon_social": master_profile.get("razon_social", "N/A"),
            "rfc": master_profile.get("rfc", "N/A"),
            "numero_licitacion": session_id,
            "servicio": master_profile.get("giro", "servicio licitado"),
            "nombre_representante": master_profile.get("representante_legal", "N/A"),
            "domicilio": domicilio,
            "lugar": lugar,
            "fecha": metadata.get("fecha", ""),
            "tipo_licitacion": "Licitacion Publica",
            "autoridad_convocante": "Convocante",
            "actividad_principal": actividad_principal,
        }
        # Hito 3.2: Inyectar overrides económicos en el diccionario de la plantilla
        if economic_overrides:
            for k, v in economic_overrides.items():
                if k != "concept_prices":
                    data[f"econ_{k}"] = v
        return data

    async def process(self, agent_input: AgentInput) -> AgentOutput:
        session_id = agent_input.session_id
        correlation_id = agent_input.correlation_id or "no-id"
        started_at = time.perf_counter()
        llm = self.llm
        context = await self.context_manager.get_global_context(session_id)

        # 1. RECUPERAR DATOS DE IDENTIDAD (PRODUCCIÓN)
        company_data = agent_input.company_data or {}
        master_profile = company_data.get("master_profile", {})
        
        if not master_profile:
            state = await self.context_manager.memory.get_session(session_id)
            if state and "initial_data" in state:
                master_profile = state["initial_data"].get("company_data", {}).get("master_profile", {})
        if not master_profile and agent_input.company_id:
            try:
                company_db = await self.context_manager.memory.get_company(str(agent_input.company_id))
                if isinstance(company_db, dict):
                    master_profile = company_db.get("master_profile") or company_db.get("catalog") or {}
            except Exception:
                pass

        tipo_persona = master_profile.get("tipo", "moral").lower()
        razon_social = master_profile.get("razon_social", "EMPRESA SIN REGISTRO")
        rfc = master_profile.get("rfc")
        representante = master_profile.get("representante_legal")

        # --- Hito 4: piloto bloqueo por slots (sin escribir bajo /data/outputs) ---
        missing_slots = build_formats_pilot_missing_entries(
            master_profile,
            blocking_job_id=agent_input.job_id,
        )
        if missing_slots:
            field_keys = [m["field"] for m in missing_slots]
            logger.info(
                "formats_pilot_blocked",
                agent_id=self.agent_id,
                session_id=session_id,
                correlation_id=correlation_id,
                blocking_job_id=agent_input.job_id,
                pilot="formats_administrativo",
                missing_fields=field_keys,
                missing_count=len(missing_slots),
            )
            await self._save_pending_questions(session_id, missing_slots)

            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=f"Para generar tus documentos necesito: {', '.join([m['label'] for m in missing_slots])}",
                data={"missing": missing_slots},
                correlation_id=correlation_id,
            )

        # Valores seguros después de validación
        rfc = rfc or "N/A"
        representante = representante or (razon_social if tipo_persona == "fisica" else "N/A")
        
        # Lógica de Redacción Universal (Yo vs Nosotros)
        pronombres = "en primera persona ('Yo', 'mi empresa')" if tipo_persona == "fisica" else "en representación de la empresa ('Nosotros', 'la empresa')"
        system_prompt = (
            f"ERES UN REDACTOR LEGAL EXPERTO EN LICITACIONES MEXICANAS. Escribe {pronombres}. "
            "REGLA DE ORO (ESPEJO ESTRICTO): Si se te proporciona una 'PLANTILLA OFICIAL O FORMATO DE LAS BASES', "
            "tu tarea es realizar una TRANSCRIPCIÓN FIEL. Mantén la estructura, todas las columnas de las tablas, "
            "el orden de los párrafos y el lenguaje legal exacto. "
            "SOLO debes rellenar los datos de la empresa, fechas y placeholders. "
            "ESTÁ PROHIBIDO omitir secciones, simplificar tablas o cambiar el formato original. "
            "Si el contexto muestra una tabla con columnas específicas, genera esa tabla EXACTAMENTE IGUAL. "
            "Nunca escribas 'Dato pendiente de confirmar', nunca dejes blancos reales y nunca cites una institución "
            "o convocante distinta a la del procedimiento actual. Si no tienes un dato variable, omite la cláusula "
            "o devuelve solo el texto sustentado por contexto real. "
            f"{ANTI_PLACEHOLDER_PROMPT_RULE} "
            f"{LEGAL_AUTHORIZATION_SYSTEM_RULE} "
            f"{CONCURSANTE_LEXICON_SYSTEM_RULE}"
        )

        # Buscar logo real
        logo_path = master_profile.get("logo")
        if not logo_path:
            logo_info = company_data.get("docs", {}).get("LOGOTIPO", {})
            if logo_info and isinstance(logo_info, dict):
                logo_path = logo_info.get("path")

        # Metadata para Word
        _MESES_F = {
            "January": "enero", "February": "febrero", "March": "marzo",
            "April": "abril", "May": "mayo", "June": "junio",
            "July": "julio", "August": "agosto", "September": "septiembre",
            "October": "octubre", "November": "noviembre", "December": "diciembre"
        }
        session_state = dict(context.get("session_state", {}) or {})
        if not str(session_state.get("bases_corpus_hint") or "").strip():
            try:
                cal_res = self.vector_db.query_texts(
                    session_id,
                    "recepción de propuestas calendario evento fecha presentación apertura proposiciones",
                    n_results=20,
                )
                cal_docs = cal_res.get("documents") or []
                session_state["bases_corpus_hint"] = "\n".join(
                    d for d in cal_docs if d
                )[:120000]
            except Exception as exc:
                logger.warning(
                    "formats_bases_corpus_hint_failed",
                    session_id=session_id,
                    error=str(exc)[:120],
                )
        try:
            from app.services.obra_economic_annex_clauses import (
                fetch_obra_e1_format_corpus_from_index,
            )

            e1_blob = fetch_obra_e1_format_corpus_from_index(session_id)
            hint = str(session_state.get("bases_corpus_hint") or "")
            if e1_blob and "carta compromiso de proposici" not in hint.lower():
                session_state["bases_corpus_hint"] = f"{hint}\n\n{e1_blob}"[:160000]
        except Exception as exc:
            logger.warning(
                "formats_e1_corpus_hint_failed",
                session_id=session_id,
                error=str(exc)[:120],
            )
        try:
            from app.services.convocante_resolver import (
                extract_convocante_from_text,
                merge_convocante_into_session_patch,
            )

            conv_hint = str(session_state.get("bases_corpus_hint") or "")
            if len(conv_hint) < 800:
                conv_res = self.vector_db.query_texts(
                    session_id,
                    "H. Ayuntamiento Dirección General convocante licitación pública número comité",
                    n_results=20,
                )
                conv_hint = "\n".join(
                    d for d in (conv_res.get("documents") or []) if d
                )[:120000]
            patch = merge_convocante_into_session_patch(session_state, conv_hint)
            if not patch.get("convocante"):
                patch = extract_convocante_from_text(conv_hint)
            if patch.get("convocante"):
                la = dict(session_state.get("last_analysis") or {})
                for k, v in patch.items():
                    if v and not str(la.get(k) or "").strip():
                        la[k] = v
                session_state["last_analysis"] = la
                if patch.get("destinatario"):
                    session_state["destinatario"] = patch["destinatario"]
                session_state["convocante"] = patch.get("convocante") or session_state.get("convocante")
        except Exception as conv_exc:
            logger.warning(
                "formats_convocante_enrich_failed",
                session_id=session_id,
                error=str(conv_exc)[:120],
            )
        _date_info = resolve_document_date(session_state)
        _fecha_f = _date_info.get("fecha_es") or datetime.now().strftime("%d de %B de %Y")

        _dom_fiscal = master_profile.get("domicilio_fiscal") or master_profile.get("domicilio") or ""
        _dom_footer = _dom_fiscal if is_usable_profile_field_value(_dom_fiscal) else "S/D"
        from app.services.administrative_letter_clauses import (
            is_invalid_letter_lugar,
            resolve_document_ciudad,
            resolve_letter_session_metadata,
        )

        _letter_meta = resolve_letter_session_metadata(session_state)
        _ciudad = resolve_document_ciudad(
            master_profile, str(_dom_fiscal), letter_meta=_letter_meta
        )
        doc_metadata = {
            "logo_path": logo_path,
            "tender_name": session_id.replace("_", " ").upper(),
            "fecha": _fecha_f,
            "fecha_corta": _date_info.get("fecha_corta", ""),
            "deadline_dt_iso": _date_info.get("deadline_dt"),
            "date_source": _date_info.get("source", ""),
            "hora": datetime.now().strftime("%H:%M"),
            "empresa": razon_social,
            "rfc": rfc,
            "representante": representante,
            "domicilio": _dom_fiscal if is_usable_profile_field_value(_dom_fiscal) else "",
            "footer_text": f"{razon_social} | RFC: {rfc} | Domicilio: {_dom_footer}",
            "destinatario": _letter_meta.get("destinatario")
            or resolve_addressee_lines(
                session_state, agent_input.triage_context if hasattr(agent_input, "triage_context") else None
            ),
            "convocante": _letter_meta.get("convocante", ""),
            "entidad": _letter_meta.get("entidad", ""),
            "dependencia": _letter_meta.get("dependencia", ""),
            "concurso_label": _letter_meta.get("concurso_label", ""),
            "ciudad": _ciudad,
            "formal_closing": True,
            "bases_corpus_hint": session_state.get("bases_corpus_hint", ""),
            "materialization_provenance": "llm_controlled",
        }

        # --- Hito 3.2: Inyectar Overrides Económicos ---
        user_inputs = session_state.get("economic_user_inputs") or {}
        econ_parts = []
        for k, v in user_inputs.items():
            if k == "concept_prices" and isinstance(v, dict):
                for c, p in v.items():
                    econ_parts.append(f"Precio Unitario de {c.upper()}: {p}")
            elif v is not None:
                lbl = k.replace("_", " ").upper()
                econ_parts.append(f"{lbl}: {v}")
        
        economic_block = ""
        if econ_parts:
            economic_block = "\nDATOS ECONÓMICOS CONFIRMADOS (USAR ESTOS VALORES SI EL DOCUMENTO LO REQUIERE):\n" + "\n".join(econ_parts)

        try:
            from app.services.mini_dictamen_anexos_service import (
                build_and_persist_mini_dictamen,
                build_stage_blocking_questions,
                get_blocking_annex_rows_for_stage,
            )

            await build_and_persist_mini_dictamen(self.context_manager.memory, session_id)
            fresh_state = await self.context_manager.memory.get_session(session_id) or session_state
            blocking_rows = get_blocking_annex_rows_for_stage(fresh_state, "formats")
            if blocking_rows:
                fresh_state["pending_questions"] = build_stage_blocking_questions(
                    "formats", blocking_rows
                ) + list(fresh_state.get("pending_questions") or [])
                fresh_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, fresh_state)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "La generación administrativa quedó bloqueada por anexos obligatorios "
                        "con fuente inválida, referencial o pendiente de aclaración."
                    ),
                    data={"missing": blocking_rows},
                    correlation_id=correlation_id,
                )
        except Exception as exc:
            logger.warning(
                "formats_mini_dictamen_guard_failed",
                session_id=session_id,
                error=str(exc),
            )

        # 2. RECUPERAR LISTA MAESTRA (SINCRONIZACIÓN CORONACIÓN)
        # Orden de prioridad:
        # a) Inyección directa del orquestador via compliance_master_list
        # b) Resultados de Fase 1 via results.compliance.data
        # c) Tarea persistida master_compliance_list en tasks_completed
        session_state = context.get("session_state", {})
        tasks = session_state.get("tasks_completed", [])
        compliance_data: Dict[str, Any] = {}
        injected = company_data.get("compliance_master_list")
        if isinstance(injected, dict) and injected:
            compliance_data = injected
        if not compliance_data:
            for task in reversed(tasks):
                if task.get("task") == "stage_completed:compliance":
                    res = task.get("result") or {}
                    if isinstance(res, dict) and res.get("data"):
                        compliance_data = res["data"]
                        break
        if not compliance_data:
            for task in reversed(tasks):
                if task.get("task") == "master_compliance_list":
                    res = task.get("result") or {}
                    compliance_data = res.get("data", res) if isinstance(res, dict) else {}
                    break
        reqs_to_process: List[Dict[str, Any]] = []
        seen_ids: Set[str] = set()
        raw_list = compliance_data.get("administrativo", []) + compliance_data.get("formatos", [])
        action_counts: Dict[str, int] = {"generar": 0, "presentar_fisico": 0, "informativo": 0, "unknown": 0}
        total_candidates = len(raw_list)
        evidence_true = sum(1 for r in raw_list if bool(r.get("evidence_match")))
        evidence_ratio = (evidence_true / total_candidates) if total_candidates else 1.0

        admin_output_dir = os.path.join(
            "/data", "outputs", session_id, "3.documentos administrativos"
        )
        os.makedirs(admin_output_dir, exist_ok=True)

        # Contador de formas numeradas por prefijo para detectar alucinaciones secuenciales
        _numbered_form_counts: Dict[str, int] = {}

        from app.config.settings import settings
        from app.services.document_deliverable_filter import (
            has_admin_format_template_evidence,
            is_company_credential_present_only,
            is_economic_writer_domain,
            is_generable_tipo_accion,
            normalize_deliverable_key,
            should_show_deliverable_in_ui,
        )

        seen_sigs: Set[str] = set()

        for req in raw_list:
            rid = str(req.get("id", "")).strip().replace(".", "_")
            raw_name = req.get("nombre", "Documento")
            if not rid or rid in seen_ids:
                continue

            raw_name_u = str(raw_name or "")
            desc_u = str(req.get("descripcion", "") or "")
            if not should_show_deliverable_in_ui(
                raw_name_u,
                desc_u,
                str(req.get("snippet") or ""),
                str(req.get("tipo_accion") or ""),
            ):
                seen_ids.add(rid)
                logger.info(
                    "formats_causal_item_skipped",
                    session_id=session_id,
                    rid=rid,
                    nombre=raw_name_u[:80],
                )
                continue

            # Si el compliance clasificó tipo_accion, usarlo directamente
            tipo_accion = str(req.get("tipo_accion", "")).lower()
            if tipo_accion not in action_counts:
                tipo_accion = "unknown"
            action_counts[tipo_accion] = action_counts.get(tipo_accion, 0) + 1
            if not is_generable_tipo_accion(tipo_accion):
                seen_ids.add(rid)
                continue
            if is_company_credential_present_only(
                raw_name_u, desc_u, str(req.get("snippet") or "")
            ):
                seen_ids.add(rid)
                logger.info(
                    "formats_company_credential_skipped",
                    session_id=session_id,
                    rid=rid,
                    nombre=raw_name_u[:80],
                )
                continue
            if is_economic_writer_domain(raw_name_u, desc_u, str(req.get("snippet") or "")):
                seen_ids.add(rid)
                continue
            if re.match(r"^AD[-_]?\d+", rid, re.I):
                integridad_sig = normalize_deliverable_key(
                    "declaracion integridad", "administrativo"
                )
                if integridad_sig in seen_sigs or any(
                    "integridad" in s for s in seen_sigs
                ):
                    seen_ids.add(rid)
                    logger.info(
                        "formats_compliance_integridad_deduped",
                        session_id=session_id,
                        rid=rid,
                    )
                    continue

            blob_ids = f"{rid} {raw_name_u} {desc_u}".upper()

            # ── FILTRO ANTI-ALUCINACIÓN ──────────────────────────────────────────
            # Ítems con tipo_accion == "unknown" y patrón de forma numerada (DD-NN,
            # AT-NN, AE-NN) sin evidencia real son candidatos a alucinación del LLM.
            # El LLM ve DD-01..DD-10 en las bases y "completa" la secuencia hasta
            # DD-29 con descripciones genéricas. Filtramos estos ítems si:
            #   1. tipo_accion es "unknown" (el LLM no los clasificó como "generar")
            #   2. Tienen patrón de forma numerada (DD-NN, AT-NN, AE-NN)
            #   3. NO tienen evidence_match = True (no hay snippet literal en las bases)
            #   4. El snippet es vacío o muy corto (< 20 chars)
            # Excepción: si tipo_accion == "generar", el LLM lo clasificó explícitamente
            # y lo dejamos pasar aunque no tenga evidencia perfecta.
            if tipo_accion == "unknown":
                _numbered_form_match = re.search(
                    r"\b(DD|AT|AE|FO|DC)[-_]?\d{1,2}\b", blob_ids
                )
                if _numbered_form_match:
                    has_evidence = bool(req.get("evidence_match"))
                    snippet = str(req.get("snippet") or "").strip()
                    has_real_snippet = len(snippet) >= 20
                    if not has_evidence and not has_real_snippet:
                        # Alucinación probable: forma numerada sin evidencia literal
                        logger.info(
                            "formats_hallucinated_form_discarded",
                            session_id=session_id,
                            rid=rid,
                            nombre=raw_name_u[:80],
                            reason="numbered_form_no_evidence",
                        )
                        seen_ids.add(rid)  # Marcar como visto para no re-procesar
                        continue

            if not has_admin_format_template_evidence(req):
                seen_ids.add(rid)
                logger.info(
                    "formats_no_template_evidence_skipped",
                    session_id=session_id,
                    rid=rid,
                    nombre=raw_name_u[:80],
                )
                continue

            sig = normalize_deliverable_key(raw_name_u, "administrativo")
            if sig in seen_sigs:
                seen_ids.add(rid)
                continue
            seen_sigs.add(sig)
            reqs_to_process.append(req)
            seen_ids.add(rid)

        # Panel consolidado (misma fuente que UI): anexos «generar» con ancla en bases.
        panel_payload: Dict[str, Any] = {}
        panel_expected = 0
        try:
            from app.services.document_candidate_list_service import (
                build_formats_panel_consolidated,
            )

            panel_payload = await build_formats_panel_consolidated(
                self.context_manager.memory, session_id, session_state
            )
            from app.services.formats_coverage_gate import count_panel_admin_generar

            panel_expected = count_panel_admin_generar(panel_payload)
            panel_authoritative = _build_panel_authoritative_reqs(panel_payload)
            if panel_expected >= 3 and panel_authoritative:
                reqs_to_process = panel_authoritative
                action_counts["generar"] = len(panel_authoritative)
                logger.info(
                    "formats_panel_authoritative_queue",
                    session_id=session_id,
                    count=len(panel_authoritative),
                    panel_expected=panel_expected,
                )
            else:
                panel_added = 0
                for req_panel in panel_authoritative:
                    sig_panel = normalize_deliverable_key(
                        str(req_panel.get("nombre") or ""), "administrativo"
                    )
                    if sig_panel in seen_sigs:
                        continue
                    seen_sigs.add(sig_panel)
                    rid_panel = str(req_panel.get("id") or "")
                    if rid_panel in seen_ids:
                        continue
                    seen_ids.add(rid_panel)
                    reqs_to_process.append(req_panel)
                    panel_added += 1
                    action_counts["generar"] = action_counts.get("generar", 0) + 1
                if panel_added:
                    logger.info(
                        "formats_panel_generables_merged",
                        session_id=session_id,
                        added=panel_added,
                        total=len(reqs_to_process),
                    )
        except Exception as panel_exc:
            logger.warning(
                "formats_panel_merge_failed",
                session_id=session_id,
                error=str(panel_exc)[:200],
            )

        gate = _should_block_by_quality_gate(
            total_items=total_candidates,
            generar_count=action_counts.get("generar", 0),
            unknown_count=action_counts.get("unknown", 0),
            evidence_match_ratio=evidence_ratio,
            presentar_fisico_count=action_counts.get("presentar_fisico", 0),
            triage_context=agent_input.triage_context,
        )
        if gate.get("block"):
            # En lugar de bloquear al usuario con un pendiente técnico,
            # logueamos la situación y continuamos con lo que tenemos.
            # El gate solo interrumpe si no hay absolutamente nada que generar
            # Y no hay documentos para presentar físicamente.
            has_anything_to_do = (
                action_counts.get("generar", 0) > 0
                or action_counts.get("presentar_fisico", 0) > 0
                or len(reqs_to_process) > 0
            )
            logger.warning(
                "formats_quality_gate_triggered",
                session_id=session_id,
                reason=str(gate.get("reason")),
                metrics=gate.get("metrics"),
                continuing=has_anything_to_do,
            )
            if not has_anything_to_do:
                # Solo bloquear si realmente no hay nada que hacer
                missing = [
                    {
                        "field": "document_quality_gate",
                        "label": "Confirmar clasificación documental administrativa",
                        "question": (
                            "La lista administrativa/formatos no es lo suficientemente confiable "
                            "para generar documentos automáticamente. Requiere reclasificación por acción "
                            "o evidencia más sólida."
                        ),
                        "document_hint": f"Motivo gate: {gate.get('reason')}. Métricas: {gate.get('metrics')}",
                        "type": "document_quality_gate_blocking",
                        "blocking_items": [],
                    }
                ]
                await self._save_pending_questions(session_id, missing)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "Pausa por calidad documental: la clasificación de formatos/administrativos "
                        "no es suficientemente confiable para generar archivos sin riesgo."
                    ),
                    data={"missing": missing, "document_quality_gate": gate},
                    correlation_id=correlation_id,
                )

        from app.services.ingested_file_resolver import (
            build_ingested_file_index,
            resolve_ingested_file,
        )
        from app.services.session_template_catalog import build_catalog_mirror_reqs
        from app.services.template_mirror_service import mirror_template_to_output

        session_documents: List[Dict[str, Any]] = []
        try:
            session_documents = await self.context_manager.memory.get_documents(session_id)
        except Exception as doc_exc:
            logger.warning("formats_get_documents_failed", session_id=session_id, error=str(doc_exc))

        file_index = build_ingested_file_index(session_documents)
        catalog_reqs = build_catalog_mirror_reqs(
            session_state,
            seen_ids,
            exclude_sobre=("economico",),
        )
        reqs_to_process = catalog_reqs + reqs_to_process

        mirror_enabled = bool(getattr(settings, "TEMPLATE_MIRROR_ENABLED", True))
        mirror_max = int(getattr(settings, "TEMPLATE_MIRROR_MAX_ADMIN", 40) or 40)
        mirror_queue: List[tuple] = []
        llm_queue: List[Dict[str, Any]] = []
        unresolved_catalog_mirrors: List[str] = []
        session_hint = f"{session_id} {session_state.get('name', '')}"
        for req in reqs_to_process:
            q = str(req.get("archivo_fuente") or req.get("nombre") or "")
            ref = resolve_ingested_file(
                q,
                file_index,
                doc_id=req.get("source_doc_id"),
                source_path=req.get("source_path"),
            )
            ext = (ref.file_path.rsplit(".", 1)[-1].lower() if ref and ref.file_path else "")
            if bool(req.get("from_session_catalog")) and not ref:
                unresolved_catalog_mirrors.append(str(req.get("nombre") or q or "sin_nombre"))
                logger.warning(
                    "formats_catalog_source_unresolved_skip_llm",
                    session_id=session_id,
                    requested=q,
                    source_doc_id=str(req.get("source_doc_id") or ""),
                    source_path=str(req.get("source_path") or ""),
                )
                continue
            if mirror_enabled and ref and ext in ("doc", "docx", "xls", "xlsx"):
                if _mirror_source_has_cross_tender_marker(ref, session_hint):
                    logger.warning(
                        "formats_cross_tender_mirror_proceeding",
                        session_id=session_id,
                        requested=q,
                        source_filename=ref.filename,
                    )
                mirror_queue.append((req, ref))
            else:
                llm_queue.append(req)

        if unresolved_catalog_mirrors:
            logger.warning(
                "formats_catalog_mirror_sources_unresolved",
                session_id=session_id,
                count=len(unresolved_catalog_mirrors),
                names=unresolved_catalog_mirrors[:10],
            )

        if mirror_max > 0 and len(mirror_queue) > mirror_max:
            mirror_queue = mirror_queue[:mirror_max]

        max_formats = int(getattr(settings, "FORMATS_MAX_GENERABLE_DOCS", 18) or 18)
        if max_formats > 0 and panel_expected > max_formats:
            max_formats = min(max(panel_expected + 5, max_formats), 60)
        if max_formats > 0 and len(llm_queue) > max_formats:
            llm_queue = llm_queue[:max_formats]

        _merge_document_inventory_legal(company_data, admin_output_dir, llm_queue, seen_ids)

        logger.info(
            "formats_generation_started",
            agent=self.agent_id,
            session_id=session_id,
            mirror_count=len(mirror_queue),
            llm_count=len(llm_queue),
        )

        generated_files = []
        generation_skipped: List[Dict[str, Any]] = []
        initial_mirror_count = len(mirror_queue)
        initial_llm_count = len(llm_queue)

        def _record_generation_skip(
            nombre: str,
            reason: str,
            **extra: Any,
        ) -> None:
            generation_skipped.append(
                {
                    "nombre": str(nombre or "")[:200],
                    "reason": reason,
                    **extra,
                }
            )
        zonas_detectadas: list[str] = []
        zone_sources: List[str] = []
        for req_item in list(mirror_queue) + [(item, None) for item in llm_queue]:
            req_data = req_item[0] if isinstance(req_item, tuple) else req_item
            zone_sources.append(
                " ".join(str(req_data.get(k) or "") for k in ("nombre", "descripcion", "titulo", "label"))
            )
        for doc in session_documents:
            if not isinstance(doc, dict):
                continue
            content = doc.get("content") if isinstance(doc.get("content"), dict) else {}
            metadata = doc.get("metadata") if isinstance(doc.get("metadata"), dict) else {}
            zone_sources.append(
                " ".join(
                    str(v or "")
                    for v in (
                        content.get("filename"),
                        content.get("doc_type"),
                        metadata.get("filename"),
                        metadata.get("doc_type"),
                        metadata.get("title"),
                    )
                )
            )
        for blob in zone_sources:
            for zone in re.findall(r"\bzona\s+([A-Z])\b", blob, flags=re.IGNORECASE):
                z = str(zone).upper()
                if z not in zonas_detectadas:
                    zonas_detectadas.append(z)
        for z in _zones_from_economic_user_inputs(user_inputs):
            if z not in zonas_detectadas:
                zonas_detectadas.append(z)
        zonas_ofertadas = ", ".join(zonas_detectadas[:-1]) + (" y " + zonas_detectadas[-1] if len(zonas_detectadas) > 1 else (zonas_detectadas[0] if zonas_detectadas else ""))
        tarifa_ref = _pick_reference_service_price(user_inputs)
        if tarifa_ref is None:
            tarifa_ref = await _pick_reference_service_price_from_line_items(
                self.context_manager.memory,
                session_id,
            )
        profile_fill = {
            "rfc": rfc,
            "razon_social": razon_social,
            "representante_legal": representante,
            "domicilio": master_profile.get("domicilio_fiscal") or master_profile.get("domicilio"),
            "fecha": doc_metadata.get("fecha"),
            "licitacion": doc_metadata.get("tender_name") or session_id,
            "zonas_ofertadas": zonas_ofertadas,
            "numero_referencia": doc_metadata.get("tender_name") or session_id,
            "tarifa_mensual_referencia": tarifa_ref,
        }
        contenido_nacional_pct = _infer_formats_contenido_nacional_pct(self.vector_db, session_id)

        def _mirror_output_dir(req_item: Dict[str, Any]) -> str:
            base = os.path.join("/data", "outputs", session_id)
            sobre = str(req_item.get("sobre_inferido") or "")
            tipo = str(req_item.get("tipo") or "").lower()
            if sobre == "tecnico" or tipo == "tecnico":
                sub = "1.propuesta tecnica"
            elif sobre == "economico":
                sub = "2.propuesta_economica"
            else:
                sub = "3.documentos administrativos"
            path = os.path.join(base, sub)
            os.makedirs(path, exist_ok=True)
            return path

        for req, ref in mirror_queue:
            rid = str(req.get("id", "")).strip().replace(".", "_")
            raw_name = req.get("nombre", "Documento")
            safe_name = re.sub(r"[^\w\s-]", "", str(raw_name).replace(" ", "_"))[:60].strip("_")
            filename = f"{rid}_{safe_name}" if rid else safe_name
            filename = re.sub(r"_+", "_", filename).strip("_")
            src_ext = ref.file_path.rsplit(".", 1)[-1].lower()
            out_ext = ".docx" if src_ext == "doc" else f".{src_ext}"
            if out_ext not in (".docx", ".xlsx", ".xls"):
                out_ext = ".docx"
            output_dir = _mirror_output_dir(req)
            filepath = os.path.join(output_dir, f"{filename}{out_ext}")
            existing_path = _formats_inventory_doc_path(output_dir, rid)
            if existing_path:
                generated_files.append(
                    attach_traceability(
                        {
                            "nombre": raw_name,
                            "ruta": existing_path,
                            "status": "REUSED",
                            "tipo": str(req.get("tipo") or "administrativo"),
                            "template_id": None,
                        },
                        source_doc_id=str(ref.doc_id or req.get("source_doc_id") or "") or None,
                        source_filename=ref.filename,
                        source_path=ref.file_path,
                        materialization_route="mirror_reused",
                        output_hash=safe_file_sha256(existing_path),
                    )
                )
                continue
            try:
                normalized_name = _normalize_formats_blob(raw_name or ref.filename)
                if "contenido nacional" in normalized_name:
                    contenido_text = _build_formats_contenido_nacional_text(
                        razon_social=razon_social,
                        rfc=rfc,
                        representante=representante,
                        session_id=session_id,
                        tender_name=doc_metadata.get("tender_name") or session_id,
                        fecha_es=doc_metadata.get("fecha") or "",
                        zonas=zonas_detectadas,
                        porcentaje=contenido_nacional_pct,
                        destinatario=doc_metadata.get("destinatario") or "",
                    )
                    _save_docx(raw_name, contenido_text, filepath, doc_metadata)
                    meta = {
                        "ruta": filepath,
                        "mirror_mode": "deterministic_contenido_nacional",
                        "materialization_route": "mirror",
                        "source_filename": ref.filename,
                        "source_path": ref.file_path,
                        "source_hash": safe_file_sha256(ref.file_path),
                        "output_hash": safe_file_sha256(filepath),
                    }
                else:
                    meta = mirror_template_to_output(
                        ref,
                        filepath,
                        profile_fill,
                        fill_profile=True,
                    )
                canon_fecha = str(doc_metadata.get("fecha") or "").strip()
                if canon_fecha and meta.get("ruta"):
                    normalize_docx_spanish_dates(str(meta["ruta"]), canon_fecha)
                generated_files.append(
                    attach_traceability(
                        {
                        "nombre": raw_name,
                        "ruta": meta["ruta"],
                        "status": "FINAL",
                        "tipo": str(req.get("tipo") or "administrativo"),
                        "template_id": None,
                        },
                        source_doc_id=str(ref.doc_id or req.get("source_doc_id") or "") or None,
                        source_filename=ref.filename,
                        source_path=meta.get("source_path") or ref.file_path,
                        source_hash=meta.get("source_hash"),
                        template_id=None,
                        mirror_mode=meta.get("mirror_mode"),
                        materialization_route=meta.get("materialization_route") or "mirror",
                        output_hash=meta.get("output_hash") or safe_file_sha256(meta.get("ruta")),
                        provenance_ui=req.get("provenance_ui") if isinstance(req.get("provenance_ui"), dict) else None,
                    )
                )
                logger.info(
                    "formats_mirror_ok",
                    session_id=session_id,
                    source=ref.filename,
                    mode=meta.get("mirror_mode"),
                )
            except Exception as mir_exc:
                logger.warning(
                    "formats_mirror_failed",
                    session_id=session_id,
                    source=ref.filename,
                    error=str(mir_exc),
                )
                llm_queue.append(req)

        for req in llm_queue:
            rid = str(req.get("id", "")).strip().replace(".", "_")
            raw_name = req.get('nombre', 'Documento')
            if is_apu_document(raw_name, str(req.get("descripcion") or ""), rid):
                logger.info(
                    "formats_apu_deferred_to_economic_writer",
                    session_id=session_id,
                    req_id=rid,
                    req_name=raw_name[:80],
                )
                continue
            safe_name = re.sub(r"[^\w\s-]", "", raw_name.replace(" ", "_"))[:60].strip("_")
            if rid and not str(rid).startswith("panel_"):
                filename = f"{rid}_{safe_name}"
            else:
                filename = safe_name or "documento"
            filename = re.sub(r"_+", "_", filename).strip("_")
            output_dir_early = _mirror_output_dir(req)
            existing_path = _formats_inventory_doc_path(output_dir_early, rid)
            if existing_path:
                _display = raw_name if str(raw_name).lower().endswith(".docx") else f"{raw_name}.docx"
                generated_files.append(
                    attach_traceability(
                        {
                            "nombre": raw_name,
                            "source_filename": str(req.get("archivo_fuente") or _display).strip() or _display,
                            "ruta": existing_path,
                            "status": "REUSED",
                            "tipo": str(req.get("tipo") or "administrativo"),
                            "template_id": None,
                        },
                        source_doc_id=str(req.get("source_doc_id") or "") or None,
                        materialization_route="llm_reused",
                        output_hash=safe_file_sha256(existing_path),
                    )
                )
                continue
            display_title = raw_name[:120] if str(rid).startswith("panel_") else (
                f"{rid} - {raw_name}"[:120] if rid else raw_name[:120]
            )
            
            template_id = self._template_id_for_requirement(req)
            req_nombre = raw_name
            req_desc = str(req.get("descripcion") or "")
            req_snippet = str(req.get("snippet") or "")
            req_source = str(req.get("archivo_fuente") or "").strip()
            req_context = ""
            context_type = "FRAGMENTADO (RAG)"
            try:
                if req_source and req_source.lower().endswith(".pdf"):
                    logger.info(
                        "formats_mirror_protocol_activated",
                        session_id=session_id,
                        source=req_source,
                    )
                    req_context = self.vector_db.get_full_document_text(session_id, req_source)
                    context_type = "ESPEJO COMPLETO (TEMPLATE OFICIAL)"
                if not req_context:
                    rag_query = f"{req_nombre} {req_desc} {req_snippet}".strip()
                    rag_n = 8 if req.get("inventory_synthetic") else 5
                    req_context_res = self.vector_db.query_texts(
                        session_id, rag_query, n_results=rag_n
                    )
                    docs = req_context_res.get("documents", []) if req_context_res else []
                    req_context = "\n".join(d for d in docs[:6] if d and d.strip())
            except Exception as _rag_err:
                logger.warning(
                    "formats_context_retrieval_failed",
                    agent=self.agent_id,
                    req_name=req_nombre[:80],
                    error=str(_rag_err),
                )

            hint = req.get("generator_hint")
            hint_line = f"\nGUÍA_DE_PLANTILLA_O_CLAVE: {hint}" if hint else ""
            bases_context_block = (
                f"\n--- INICIO DE {context_type} ---\n{req_context}\n--- FIN DE {context_type} ---\n"
                if req_context
                else ""
            )
            snippet_block = ""
            if req_snippet and len(req_snippet.strip()) >= 20:
                snippet_block = (
                    f"\nFRAGMENTO LITERAL DE BASES (respeta estructura y cláusulas):\n"
                    f"{req_snippet.strip()[:2000]}\n"
                )

            from app.services.administrative_letter_clauses import (
                is_obra_pliego_contract_annex,
                is_obra_tabular_annex,
                is_short_acceptance_annex,
                resolve_letter_asunto,
                resolve_letter_session_metadata,
                strip_redundant_signature_blocks,
                try_build_clause_markdown,
            )
            from app.services.obra_economic_annex_clauses import (
                is_official_obra_e1_mirror_content,
            )

            req_letter_meta = resolve_letter_session_metadata(
                session_state,
                triage_context=agent_input.triage_context if hasattr(agent_input, "triage_context") else None,
                req_snippet=req_snippet or req_desc,
            )
            req_doc_metadata = {
                **doc_metadata,
                **{k: v for k, v in req_letter_meta.items() if v},
                "req_snippet": req_snippet or req_desc,
                "req_desc": req_desc,
                "session_id": session_id,
                "session_state": session_state,
            }
            req_doc_metadata["obra_tabular"] = is_obra_tabular_annex(raw_name)
            req_doc_metadata["obra_pliego_contract"] = is_obra_pliego_contract_annex(
                raw_name
            )
            if req_doc_metadata.get("obra_pliego_contract"):
                req_doc_metadata["document_title"] = resolve_letter_asunto(
                    raw_name, req_snippet or req_desc
                )
            if req_doc_metadata.get("obra_tabular") or req_doc_metadata.get(
                "obra_pliego_contract"
            ):
                req_doc_metadata["ciudad"] = resolve_document_ciudad(
                    master_profile,
                    str(_dom_fiscal),
                    letter_meta=req_letter_meta,
                )
            short_acceptance = is_short_acceptance_annex(raw_name, req_desc, req_snippet)

            acceptance_rule = ""
            if short_acceptance:
                acceptance_rule = (
                    "REGLA ESPECIAL: Este anexo solo requiere una carta breve de aceptación o negativa "
                    "del documento publicado en las bases. NO redactes un aviso de privacidad completo, "
                    "ni políticas extensas, ni listas de finalidades. Máximo dos párrafos sustantivos. "
                    "NO incluyas bloque de firma al final (el sistema lo agrega).\n"
                )
            prompt = (
                f"Genera el contenido legal oficial para el requisito {req.get('id')}: {req_nombre}\n"
                f"Descripción: {req_desc}\nEmpresa: {razon_social}\n"
                f"Representante: {representante}\nRFC: {rfc}\n"
                f"Domicilio: {master_profile.get('domicilio_fiscal', '')}\n"
                f"Destinatario: {req_doc_metadata.get('destinatario', '')}\n"
                f"{snippet_block}{bases_context_block}\n"
                f"{economic_block}\n{hint_line}\n"
                f"{acceptance_rule}"
                "OBLIGATORIO: Redacta como concursante (quien suscribe / mi representada). "
                "Incluye al menos un párrafo con «bajo protesta de decir verdad» o «manifiesto». "
                "No devuelvas solo títulos ni encabezados."
            )

            content = ""
            materialization_route = "generate_controlled"

            clause_body = try_build_clause_markdown(
                req_label=raw_name,
                master_profile=master_profile,
                doc_metadata=req_doc_metadata,
                req_snippet=req_snippet or req_desc,
            )
            if clause_body and not template_id:
                content = clause_body
                materialization_route = "deterministic_clause"
                if is_official_obra_e1_mirror_content(clause_body):
                    req_doc_metadata["official_bases_mirror"] = True
                    req_doc_metadata["formal_closing"] = False
            elif template_id:
                tpl_data = self._template_data(session_id, master_profile, doc_metadata, user_inputs)
                content = self.template_engine.render(template_id, tpl_data)
                if not self.template_engine.verify_integrity(content, template_id):
                    raise TemplateIntegrityError(f"Integridad inválida para template {template_id}")
                materialization_route = "template_locked"
            else:
                resp = await llm.generate(
                    prompt=prompt, system_prompt=system_prompt, correlation_id=correlation_id
                )
                if not resp.success:
                    logger.error(
                        "llm_generation_failed",
                        agent=self.agent_id,
                        req_name=raw_name,
                        error=resp.error,
                    )
                    _record_generation_skip(raw_name, "llm_generation_failed", req_id=rid)
                    continue
                llm_content = (resp.response or "").strip()
                if llm_content:
                    content = strip_llm_meta_leaks(llm_content)
            if not content.strip():
                logger.warning("llm_empty_response", agent=self.agent_id, req_name=raw_name)
                _record_generation_skip(raw_name, "llm_empty_response", req_id=rid)
                continue

            content = _sanitize_legal_content(
                content,
                session_id=session_id,
                metadata=req_doc_metadata,
            )
            content = strip_redundant_signature_blocks(content)

            from app.services.document_body_quality import is_substantive_markdown

            if (
                not template_id
                and materialization_route != "deterministic_clause"
                and not short_acceptance
                and not is_substantive_markdown(content)
            ):
                retry_prompt = (
                    f"{prompt}\n\nREINTENTO OBLIGATORIO: Redacta el cuerpo legal completo "
                    "(declaración o carta bajo protesta de decir verdad) con al menos tres "
                    "párrafos sustantivos. Debe aparecer literalmente «bajo protesta de decir verdad» "
                    "o «manifiesto». No devuelvas solo encabezados ni líneas vacías."
                )
                resp_retry = await llm.generate(
                    prompt=retry_prompt,
                    system_prompt=system_prompt,
                    correlation_id=correlation_id,
                )
                if resp_retry.success and (resp_retry.response or "").strip():
                    content = strip_llm_meta_leaks((resp_retry.response or "").strip())
                    content = _sanitize_legal_content(
                        content,
                        session_id=session_id,
                        metadata=req_doc_metadata,
                    )
                    content = strip_redundant_signature_blocks(content)
            if (
                not template_id
                and materialization_route != "deterministic_clause"
                and not short_acceptance
                and not is_substantive_markdown(content)
            ):
                from app.services.legal_document_fallback import (
                    build_administrative_fallback_markdown,
                )

                content = build_administrative_fallback_markdown(
                    req_nombre=raw_name,
                    req_desc=req_desc,
                    req_snippet=req_snippet,
                    master_profile=master_profile,
                    doc_metadata=req_doc_metadata,
                    session_state=session_state,
                )
                logger.warning(
                    "formats_insufficient_body_using_fallback",
                    session_id=session_id,
                    req_name=raw_name[:80],
                    req_id=rid,
                )

            filepath = os.path.join(output_dir_early, f"{filename}.docx")
            try:
                _save_docx(display_title, content, filepath, req_doc_metadata)
                try:
                    from docx import Document as DocxDocument

                    from app.services.document_body_quality import (
                        scan_materialized_doc_text,
                    )

                    post_paras = [
                        (p.text or "").strip()
                        for p in DocxDocument(filepath).paragraphs
                        if (p.text or "").strip()
                    ]
                    shell_hit = scan_materialized_doc_text("\n".join(post_paras))
                    if shell_hit and not template_id:
                        from app.services.legal_document_fallback import (
                            build_administrative_fallback_markdown,
                        )

                        fallback_body = build_administrative_fallback_markdown(
                            req_nombre=raw_name,
                            req_desc=req_desc,
                            req_snippet=req_snippet,
                            master_profile=master_profile,
                            doc_metadata=doc_metadata,
                            session_state=session_state,
                        )
                        os.remove(filepath)
                        _save_docx(display_title, fallback_body, filepath, doc_metadata)
                        logger.warning(
                            "formats_post_save_shell_replaced_with_fallback",
                            session_id=session_id,
                            req_name=raw_name[:80],
                            detail=shell_hit.get("detected_value"),
                        )
                except Exception as post_chk_exc:
                    logger.warning(
                        "formats_post_save_check_failed",
                        session_id=session_id,
                        error=str(post_chk_exc)[:120],
                    )
                _display = raw_name if str(raw_name).lower().endswith(".docx") else f"{raw_name}.docx"
                generated_files.append(
                    attach_traceability(
                        {
                            "nombre": raw_name,
                            "source_filename": str(req.get("archivo_fuente") or _display).strip() or _display,
                            "ruta": filepath,
                            "status": "FINAL",
                            "tipo": str(req.get("tipo") or "administrativo"),
                            "template_id": template_id,
                            "template_static_hash": self.template_engine.static_hash(template_id) if template_id else None,
                        },
                        source_doc_id=str(req.get("source_doc_id") or "") or None,
                        source_filename=str(req.get("archivo_fuente") or _display).strip() or _display,
                        source_path=str(req.get("source_path") or "") or None,
                        source_hash=safe_file_sha256(str(req.get("source_path") or "")),
                        template_id=template_id,
                        mirror_mode=None,
                        materialization_route=materialization_route,
                        output_hash=safe_file_sha256(filepath),
                        provenance_ui=req.get("provenance_ui") if isinstance(req.get("provenance_ui"), dict) else None,
                    )
                )
                logger.info("docx_generated", agent=self.agent_id, filename=filename)
            except Exception as e:
                logger.error("docx_save_failed", agent=self.agent_id, filename=filename, error=str(e))

        result_data = {
            "documentos": generated_files,
            "count": len(generated_files),
            "folder": admin_output_dir,
            "action_type_stats": action_counts,
            "generation_skipped": generation_skipped,
            "generation_expected": {
                "panel_admin_generar": panel_expected,
                "mirror_queue": initial_mirror_count,
                "llm_queue": initial_llm_count,
            },
            "materialization_metrics": build_materialization_metrics(
                stage="formats",
                documents=generated_files,
                elapsed_ms=int((time.perf_counter() - started_at) * 1000),
            ),
        }
        fill_gate = validate_generated_documents_fill(
            stage="formats",
            generated_documents=generated_files,
            master_profile=master_profile,
            provenance_context={
                "source": "formats_writer",
                "confidence": 0.9,
                "session_hint": session_hint,
                "fecha_es": doc_metadata.get("fecha"),
                "deadline_dt_iso": doc_metadata.get("deadline_dt_iso"),
            },
        )
        result_data["document_fill_quality_gate"] = fill_gate
        result_data["validation_events"] = [
            build_fill_validation_event(it, stage="formats")
            for it in (fill_gate.get("issues") or [])
            if isinstance(it, dict)
        ]
        if not bool(fill_gate.get("validation_passed", True)):
            from app.services.document_fill_ux_messages import (
                build_fill_blocking_question,
                pick_fill_gate_pending_label,
            )

            company_name = str(master_profile.get("razon_social") or "").strip()
            human_question = build_fill_blocking_question(
                "formats",
                fill_gate.get("issues") or [],
                company_name=company_name,
            )
            missing = [
                {
                    "field": "document_fill_quality_gate",
                    "label": pick_fill_gate_pending_label(fill_gate.get("issues") or []),
                    "question": human_question,
                    "document_hint": f"Gate llenado: {fill_gate.get('metrics')}",
                    "type": "document_fill_quality_gate_blocking",
                    "blocking_items": fill_gate.get("issues") or [],
                }
            ]
            await self._save_pending_questions(session_id, missing)
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=human_question,
                data={**result_data, "missing": missing, "stage": "formats"},
                correlation_id=correlation_id,
            )

        from app.services.formats_coverage_gate import evaluate_formats_stage_completeness

        generated_count_for_gate = len(generated_files)
        panel_expected_for_gate = panel_expected
        try:
            from app.services.delivery_coverage_report import build_delivery_coverage_report

            cov = build_delivery_coverage_report(
                session_id, session_state, session_documents
            )
            cov_summary = cov.get("summary") if isinstance(cov.get("summary"), dict) else {}
            cov_generadas = int(cov_summary.get("generadas") or 0)
            cov_esperadas = int(cov_summary.get("esperadas_generar") or 0)
            if cov_esperadas >= 3:
                panel_expected_for_gate = cov_esperadas
                generated_count_for_gate = max(generated_count_for_gate, cov_generadas)
        except Exception as cov_exc:
            logger.warning(
                "formats_coverage_report_for_gate_failed",
                session_id=session_id,
                error=str(cov_exc)[:160],
            )

        completeness_block = evaluate_formats_stage_completeness(
            generated_count=generated_count_for_gate,
            mirror_queue_size=initial_mirror_count,
            llm_queue_size=initial_llm_count,
            generation_skipped=generation_skipped,
            panel_expected=panel_expected_for_gate,
        )
        if completeness_block:
            missing = [
                {
                    "field": "formats_completeness_gate",
                    "label": "Anexos administrativos incompletos",
                    "question": completeness_block.get("message"),
                    "document_hint": (
                        f"Generados: {completeness_block.get('generated_count')}/"
                        f"{completeness_block.get('expected_count')}"
                    ),
                    "type": "formats_completeness_gate_blocking",
                    "blocking_items": completeness_block.get("skipped") or [],
                    "pending_names": completeness_block.get("pending_names") or [],
                }
            ]
            result_data["formats_completeness_gate"] = completeness_block
            await self._save_pending_questions(session_id, missing)
            return AgentOutput(
                status=AgentStatus.WAITING_FOR_DATA,
                agent_id=self.agent_id,
                session_id=session_id,
                message=str(completeness_block.get("message") or ""),
                data={**result_data, "missing": missing, "stage": "formats"},
                correlation_id=correlation_id,
            )

        await self._clear_resolved_formats_pending(session_id)
        await self.context_manager.record_task_completion(session_id, "formats_generation_COMPLETED", result_data)

        return AgentOutput(
            status=AgentStatus.SUCCESS,
            agent_id=self.agent_id,
            session_id=session_id,
            data=result_data,
            correlation_id=correlation_id
        )

    async def _clear_resolved_formats_pending(self, session_id: str) -> None:
        """Quita preguntas HITL de formatos ya resueltas (evita chat con mensajes obsoletos)."""
        try:
            session_state = await self.context_manager.memory.get_session(session_id) or {}
            drop_fields = frozenset(
                {
                    "formats_completeness_gate",
                    "document_fill_quality_gate",
                    "quality.fill.review",
                }
            )
            drop_types = frozenset(
                {
                    "formats_completeness_gate_blocking",
                    "document_fill_quality_gate_blocking",
                    "quality_validation_blocking",
                }
            )
            before = list(session_state.get("pending_questions") or [])
            after = [
                q
                for q in before
                if isinstance(q, dict)
                and str(q.get("field") or "") not in drop_fields
                and str(q.get("type") or "") not in drop_types
            ]
            if len(after) != len(before):
                session_state["pending_questions"] = after
                session_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, session_state)
        except Exception as e:
            logger.warning(
                "formats_clear_resolved_pending_failed",
                session_id=session_id,
                error=str(e)[:120],
            )

    async def _save_pending_questions(self, session_id: str, missing_fields: List[Dict]):
        """Persiste preguntas para el chatbot."""
        try:
            session_state = await self.context_manager.memory.get_session(session_id)
            if session_state:
                session_state["pending_questions"] = missing_fields
                session_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, session_state)
        except Exception as e:
            logger.error("save_questions_failed", agent=self.agent_id, session_id=session_id, error=str(e))

def _save_docx(title: str, content: str, file_path: str, metadata: dict = None):
    from app.services.administrative_letter_clauses import (
        city_from_domicilio,
        format_letter_lugar_ciudad,
        is_invalid_letter_lugar,
    )

    doc = docx.Document()
    section = doc.sections[0]
    
    # Header: Logo y Datos
    header = section.header
    htable = header.add_table(1, 2, Inches(6.5))
    
    # Logo
    if metadata and metadata.get("logo_path"):
        from app.utils.doc_formatting import add_logo_picture_to_run

        if not add_logo_picture_to_run(
            htable.cell(0, 0).paragraphs[0].add_run(),
            str(metadata["logo_path"]),
            width_inches=1.5,
        ):
            logger.warning(
                "logo_insert_failed",
                agent="formats_001",
                path=(metadata or {}).get("logo_path"),
            )
            
    # Datos Licitación
    p_info = htable.cell(0, 1).paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if metadata:
        run = p_info.add_run(f"{metadata.get('tender_name', '').upper()}")
        run.bold = True
        run.font.size = Pt(8)

    # Footer
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if metadata:
        p_foot.add_run(f"{metadata.get('footer_text', '')}").font.size = Pt(7)

    heading = (
        str((metadata or {}).get("document_title") or title).strip()
        if (metadata or {}).get("obra_pliego_contract")
        else title
    )
    obra_tabular = bool((metadata or {}).get("obra_tabular"))
    obra_pliego_contract = bool((metadata or {}).get("obra_pliego_contract"))
    from app.services.obra_economic_annex_clauses import is_official_obra_e1_mirror_content

    official_mirror = bool((metadata or {}).get("official_bases_mirror")) or (
        obra_pliego_contract and is_official_obra_e1_mirror_content(content)
    )

    if not official_mirror:
        doc.add_heading(heading.upper(), 1)

    # LUGAR Y FECHA (omitido en espejo del machote E-1 publicado en bases)
    footer_text = metadata.get("footer_text", "") if metadata else ""
    domicilio_ref = str((metadata or {}).get("domicilio") or "").strip()
    if not domicilio_ref and "Domicilio:" in footer_text:
        domicilio_ref = footer_text.split("Domicilio:")[-1].strip()
    ciudad = str((metadata or {}).get("ciudad") or "").strip()
    if is_invalid_letter_lugar(ciudad):
        ciudad = format_letter_lugar_ciudad(city_from_domicilio(domicilio_ref), domicilio_ref)
    if is_invalid_letter_lugar(ciudad):
        ciudad = "México"
    if not official_mirror:
        doc.add_paragraph(f"LUGAR Y FECHA: {ciudad}, a {metadata.get('fecha', '')}").alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not obra_pliego_contract:
        doc.add_paragraph("Hoja 1 de 1").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if not obra_tabular and not obra_pliego_contract:
        # Destinatario (convocante desde sesión o genérico; sin hardcode por licitación)
        dest = (metadata or {}).get("destinatario") or "A QUIEN CORRESPONDA:"
        p_dest = doc.add_paragraph(f"\n{dest}")
        p_dest.bold = True
        
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    canon_fecha = str((metadata or {}).get("fecha") or "").strip()
    if canon_fecha and content:
        content = normalize_body_spanish_dates(content, canon_fecha)

    # --- PROCESAMIENTO DE CUERPO CON SOPORTE PARA TABLAS NATIVAS ---
    raw_lines = (content or "").split("\n")
    table_buffer: List[str] = []
    
    def flush_table(buffer: List[str]):
        if not buffer: return
        matrix = parse_markdown_table(buffer)
        if not matrix: return
        
        # Crear tabla real en Word
        rows = len(matrix)
        cols = max(len(row) for row in matrix)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid' # Borde estándar oficial
        
        for r_idx, row_data in enumerate(matrix):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < cols:
                    cell_p = row_cells[c_idx].paragraphs[0]
                    # Limpiar markdown del texto de la celda
                    clean_text = strip_markdown_for_docx(cell_text)
                    run = cell_p.add_run(clean_text)
                    if r_idx == 0: # Negrita para cabeceras
                        run.bold = True
                    cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in raw_lines:
        if is_markdown_table_line(line):
            table_buffer.append(line)
            continue
        
        # Si veníamos procesando una tabla y esta línea no lo es, imprimir tabla acumulada
        if table_buffer:
            flush_table(table_buffer)
            table_buffer = []
            
        if line.strip():
            clean_p = strip_markdown_for_docx(line)
            if clean_p.strip():
                p = doc.add_paragraph(clean_p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Imprimir tabla final si quedó en buffer
    if table_buffer:
        flush_table(table_buffer)
            
    # Firma al calce (formal) — el machote E-1 ya trae ATENTAMENTE y firma del participante
    if metadata and metadata.get("formal_closing", True) and not official_mirror:
        doc.add_paragraph("\n\n")
        p_at = doc.add_paragraph("A T E N T A M E N T E\n")
        p_at.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")
        p_line = doc.add_paragraph("________________________________________\n")
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rep = str(metadata.get("representante") or "").upper()
        emp = str(metadata.get("empresa") or "").upper()
        rfc_val = str(metadata.get("rfc") or "")
        run_firma = p_line.add_run(f"{rep}\n")
        run_firma.bold = True
        p_line.add_run("REPRESENTANTE LEGAL\n").bold = True
        if emp:
            p_line.add_run(f"{emp}\n").bold = True
        if rfc_val:
            p_line.add_run(f"R.F.C. {rfc_val}\n")
        if metadata.get("require_rubrica_hint"):
            doc.add_paragraph(
                "[Rúbrica al margen derecho de cada hoja — según bases del procedimiento]"
            ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save(file_path)
    if canon_fecha:
        normalize_docx_spanish_dates(file_path, canon_fecha)
