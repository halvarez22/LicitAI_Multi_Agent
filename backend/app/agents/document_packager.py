"""
Empacador de documentos: organiza archivos generados en SOBRE_1/2/3.

Por defecto usa mapeo **determinístico** (carpeta de origen + clasificación legal/económica),
sin depender del LLM — evita duplicados (p. ej. 13× el mismo Anexo M en SOBRE_2).
"""
from __future__ import annotations

import json
import os
import re
import shutil
import time
from datetime import datetime
from typing import Any, Dict, List, Optional, Set, Tuple

import docx.shared as docx_shared
from docx import Document

from app.agents.base_agent import BaseAgent
from app.agents.mcp_context import MCPContextManager
from app.contracts.agent_contracts import AgentInput, AgentOutput, AgentStatus
from app.services.resilient_llm import ResilientLLMClient
from app.services.vector_service import VectorDbServiceClient
from app.services.document_traceability import build_materialization_metrics

# Carpetas de generación del orquestador → categoría de origen
_OUTPUT_FOLDER_TO_CATEGORY: Tuple[Tuple[str, str], ...] = (
    ("1.propuesta tecnica", "tecnica"),
    ("2.propuesta_economica", "economica"),
    ("2.propuesta economica", "economica"),
    ("economic_proposal", "economica"),
    ("3.documentos administrativos", "administrativa"),
)

_ALLOWED_EXTENSIONS = (".doc", ".docx", ".pdf", ".xlsx", ".xls", ".jpg", ".jpeg", ".png")

_SOBR_E_SHELLS = {
    "sobre_1": {
        "titulo": "SOBRE 1 - DOCUMENTACIÓN ADMINISTRATIVA",
        "nombre_carpeta": "SOBRE_1_ADMINISTRATIVO",
    },
    "sobre_2": {
        "titulo": "SOBRE 2 - PROPUESTA TÉCNICA",
        "nombre_carpeta": "SOBRE_2_TECNICO",
    },
    "sobre_3": {
        "titulo": "SOBRE 3 - PROPUESTA ECONÓMICA",
        "nombre_carpeta": "SOBRE_3_ECONOMICO",
    },
}


def _packager_use_llm_mapping() -> bool:
    return os.getenv("PACKAGER_USE_LLM_MAPPING", "").strip().lower() in (
        "1",
        "true",
        "yes",
    )


def _doc_path_key(doc: Dict[str, Any]) -> str:
    ruta = doc.get("ruta")
    if ruta and os.path.isfile(ruta):
        return os.path.normcase(os.path.abspath(ruta))
    return f"nombre:{str(doc.get('nombre') or '').strip().lower()}"


def _doc_deliverable_key(doc: Dict[str, Any]) -> str:
    """Clave de negocio para fusionar variantes del mismo anexo (p. ej. dos Anexo M)."""
    from app.services.pliego_formats_enrichment_service import pliego_format_dedupe_key

    label = str(doc.get("source_filename") or doc.get("nombre") or "").strip()
    if label:
        fk = pliego_format_dedupe_key(label)
        if fk:
            return f"deliverable:{fk}"
    return _doc_path_key(doc)


def _deliverable_label_priority(doc: Dict[str, Any]) -> int:
    """Mayor = preferido al fusionar duplicados semánticos (Anexo legible > código AD)."""
    label = str(doc.get("source_filename") or doc.get("nombre") or "")
    base = os.path.basename(label).lower()
    if re.search(r"(?i)\banexo[\s_.-]+[ivx]+", base) and not re.search(
        r"(?i)^ad[-_]", base
    ):
        return 80
    if re.search(r"(?i)^ad[-_]", base):
        return 20
    if re.match(r"(?i)^carta[_\s-]*compromiso", base):
        return 15
    return 40


def _materialization_priority(doc: Dict[str, Any]) -> int:
    """Mayor = preferido al fusionar duplicados del mismo anexo."""
    route = str(doc.get("materialization_route") or "").lower()
    if route.startswith("deterministic"):
        return 100
    if route == "mirror":
        return 80
    if route == "template_locked":
        return 60
    if route == "generate_controlled":
        return 30
    return 10


def _pick_preferred_doc(
    current: Dict[str, Any], candidate: Dict[str, Any]
) -> Dict[str, Any]:
    """Prefiere ruta determinística/mirror sobre cascarones LLM más pesados."""
    cur_label_pri = _deliverable_label_priority(current)
    cand_label_pri = _deliverable_label_priority(candidate)
    if cand_label_pri != cur_label_pri:
        return candidate if cand_label_pri > cur_label_pri else current
    cur_pri = _materialization_priority(current)
    cand_pri = _materialization_priority(candidate)
    if cand_pri != cur_pri:
        return candidate if cand_pri > cur_pri else current
    cur_path = str(current.get("ruta") or "")
    cand_path = str(candidate.get("ruta") or "")
    cur_size = os.path.getsize(cur_path) if cur_path and os.path.isfile(cur_path) else 0
    cand_size = os.path.getsize(cand_path) if cand_path and os.path.isfile(cand_path) else 0
    if cand_size != cur_size:
        return candidate if cand_size > cur_size else current
    cur_tech = "propuesta tecnica" in cur_path.replace("\\", "/").lower()
    cand_tech = "propuesta tecnica" in cand_path.replace("\\", "/").lower()
    if cand_tech and not cur_tech:
        return candidate
    return current


def _dedupe_doc_list(docs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    by_key: Dict[str, Dict[str, Any]] = {}
    seen_paths: Set[str] = set()
    for d in docs or []:
        if not isinstance(d, dict):
            continue
        pkey = _doc_path_key(d)
        if pkey in seen_paths:
            continue
        dkey = _doc_deliverable_key(d)
        if dkey in by_key:
            prev = by_key[dkey]
            seen_paths.discard(_doc_path_key(prev))
            merged = _pick_preferred_doc(prev, d)
            by_key[dkey] = merged
            seen_paths.add(_doc_path_key(merged))
        else:
            by_key[dkey] = d
            seen_paths.add(pkey)
    return list(by_key.values())


def _infer_category_from_path(ruta: str) -> Optional[str]:
    norm = ruta.replace("\\", "/").lower()
    for folder, cat in _OUTPUT_FOLDER_TO_CATEGORY:
        if f"/{folder}/" in norm or norm.endswith(f"/{folder}"):
            return cat
    return None


def _classify_doc_to_sobre_key(doc: Dict[str, Any]) -> str:
    """
    Sobre 1 = administrativo/legal, 2 = técnico, 3 = económico (convención CompraNet piloto).
    """
    ruta = str(doc.get("ruta") or "")
    nombre = str(doc.get("nombre") or "")
    cat = str(doc.get("categoria") or "")
    blob = f"{nombre} {ruta}".lower()
    label = str(
        doc.get("source_filename") or doc.get("archivo_fuente") or nombre or ""
    ).strip()

    from app.services.session_template_catalog import (
        infer_plantilla_sobre,
        is_anexo_tecnico_propuesta_entregable,
    )

    base = os.path.basename(ruta).upper() if ruta else nombre.upper()
    if is_anexo_tecnico_propuesta_entregable(label or nombre):
        return "sobre_2"
    if re.search(
        r"(?i)modelo.*presentaci[oó]n.*propuesta\s+t[eé]cnica|"
        r"presentaci[oó]n.*de\s+la\s+propuesta\s+t[eé]cnica|"
        r"fo[-_]35.*modelo|modelo.*considerar.*presentaci",
        label or nombre,
    ):
        return "sobre_2"
    # Administrativo/legal: prioridad sobre inferencia por carpeta de generación.
    if re.search(r"\bAD[-_]", base) or re.search(
        r"(?i)anexo\s+m|declaraci[oó]n\s+de\s+integridad", label or nombre
    ):
        return "sobre_1"
    if re.search(r"\b(FO|DD)[-_]", base) or "MANIFEST" in base:
        return "sobre_1"

    inferred = infer_plantilla_sobre(label)
    if inferred == "tecnico":
        return "sobre_2"
    if inferred == "economico":
        return "sobre_3"

    path_cat = _infer_category_from_path(ruta) if ruta else None
    if path_cat == "administrativa":
        return "sobre_1"
    if path_cat == "economica":
        return "sobre_3"
    if path_cat == "tecnica":
        return "sobre_2"

    if cat == "administrativa" or "administrativ" in blob:
        return "sobre_1"
    if cat == "economica" or "propuesta_econom" in blob or "propuesta econom" in blob:
        return "sobre_3"
    if cat == "tecnica" or "propuesta tecnica" in blob or "propuesta técnica" in blob:
        return "sobre_2"

    if re.search(r"\b(TE|AT)[-_]", base) or "CARTA_PRESENTACION" in base:
        return "sobre_2"
    if re.search(r"\b(AE|ANEXO_AE|TABLA_PRECIOS|CARTA_COMPROMISO_PRECIOS)\b", base):
        return "sobre_3"

    from app.services.compliance_consolidation_service import classify_deliverable_sobre

    zone = classify_deliverable_sobre(nombre, "")
    if zone == "requisitos_legales":
        return "sobre_1"
    if zone == "sobre_2_economico":
        return "sobre_3"
    return "sobre_2"


def _sort_docs_for_sobre(docs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    def sort_key(d: Dict[str, Any]) -> Tuple[int, str]:
        n = (d.get("nombre") or os.path.basename(d.get("ruta") or "")).lower()
        if "carta" in n and ("present" in n or "presentación" in n):
            return (0, n)
        if re.match(r"^0?1[_\s]", n):
            return (1, n)
        return (2, n)

    return sorted(docs, key=sort_key)


def _attach_generated_metadata(
    item: Dict[str, Any], gen_docs: Optional[Dict[str, Any]]
) -> None:
    """Propaga metadata de lineage desde documentos_generados si el basename coincide."""
    base = os.path.basename(str(item.get("ruta") or ""))
    if not base:
        return
    copy_fields = (
        "source_filename",
        "archivo_fuente",
        "source_doc_id",
        "source_path",
        "source_hash",
        "output_hash",
        "template_id",
        "mirror_mode",
        "materialization_route",
        "provenance_ui",
    )
    for docs in (gen_docs or {}).values():
        if not isinstance(docs, list):
            continue
        for d in docs:
            if not isinstance(d, dict):
                continue
            dr = str(d.get("ruta") or "")
            if not dr:
                continue
            if os.path.basename(dr) == base or base in dr:
                for field in copy_fields:
                    if d.get(field) is not None:
                        item[field] = d.get(field)
                return


def _is_canonical_pipeline_name(name: str) -> bool:
    return bool(
        re.search(
            r"(?i)^[a-z0-9]{8,}_[\w]+_Sobre(?:Complementaria|Tecnica|Economica)_\d+\.\w+$",
            (name or "").strip(),
        )
    )


def _should_skip_reference_source(doc: Dict[str, Any]) -> bool:
    """
    Excluye espejos de PDFs/pliegos de referencia que no deben materializarse como entregable.
    """
    from app.services.session_template_catalog import classify_ingested_filename

    source_label = str(
        doc.get("source_filename") or doc.get("archivo_fuente") or ""
    ).strip()
    if not source_label:
        return False
    doc_class, accion, _sobre = classify_ingested_filename(source_label)
    return doc_class in ("pliego_referencia", "informativo") or accion != "generar"


def _should_exclude_from_delivery_pack(doc: Dict[str, Any]) -> bool:
    """Excluye espejos corporativos o plantillas ajenas al pliego de la licitación."""
    label = str(doc.get("source_filename") or doc.get("nombre") or "")
    blob = label.lower().replace("_", " ")
    if re.search(r"cat[\s_-]*formato", blob) and "cmyt" in blob:
        return True
    if re.search(r"cmyt.*zen.*ofertas|formato[\s_-]*hoja[\s_-]*membretada[\s_-]*cmyt", blob):
        return True
    return False


def collect_documentos_para_empaque(
    session_id: str,
    gen_docs: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """
    Une listas del orquestador con archivos existentes en disco (fuente de verdad).
    """
    merged: Dict[str, List[Dict[str, Any]]] = {
        "administrativa": [],
        "tecnica": [],
        "economica": [],
    }
    for cat, docs in (gen_docs or {}).items():
        if cat in merged and isinstance(docs, list):
            for d in docs:
                if isinstance(d, dict):
                    item = dict(d)
                    item.setdefault("categoria", cat)
                    if _should_skip_reference_source(item):
                        continue
                    if _should_exclude_from_delivery_pack(item):
                        continue
                    merged[cat].append(item)

    output_base = os.path.join("/data", "outputs", session_id)
    for folder, cat in _OUTPUT_FOLDER_TO_CATEGORY:
        dir_path = os.path.join(output_base, folder)
        if not os.path.isdir(dir_path):
            continue
        for fn in sorted(os.listdir(dir_path)):
            if fn.startswith(".") or fn.upper().startswith("00_CARATULA"):
                continue
            ext = os.path.splitext(fn)[1].lower()
            if ext not in _ALLOWED_EXTENSIONS:
                continue
            full = os.path.join(dir_path, fn)
            from app.services.session_template_catalog import infer_plantilla_sobre

            item: Dict[str, Any] = {
                "nombre": fn,
                "ruta": full,
                "categoria": cat,
                "status": "OK",
            }
            # Ajustar categoría de origen si el nombre en disco oculta «Anexo …» (guiones bajos).
            inferred = infer_plantilla_sobre(fn)
            if inferred == "tecnico" and cat != "tecnica":
                item["categoria"] = "tecnica"
            elif inferred == "economico" and cat != "economica":
                item["categoria"] = "economica"
            _attach_generated_metadata(item, gen_docs)
            if _should_skip_reference_source(item):
                continue
            if _should_exclude_from_delivery_pack(item):
                continue
            merged[item["categoria"]].append(item)

    unified: List[Dict[str, Any]] = []
    for cat in ("administrativa", "tecnica", "economica"):
        unified.extend(_dedupe_doc_list(merged.get(cat) or []))
    # Quitar copias con nombre canónico RFC si ya existe el mismo anexo con etiqueta convocante.
    filtered: List[Dict[str, Any]] = []
    seen_conv_keys: Set[str] = set()
    for d in unified:
        label = str(d.get("source_filename") or d.get("nombre") or "")
        from app.services.pliego_formats_enrichment_service import pliego_format_dedupe_key

        dkey = pliego_format_dedupe_key(label)
        if dkey and not _is_canonical_pipeline_name(label):
            seen_conv_keys.add(dkey)
    for d in unified:
        label = str(d.get("source_filename") or d.get("nombre") or "")
        from app.services.pliego_formats_enrichment_service import pliego_format_dedupe_key

        dkey = pliego_format_dedupe_key(label)
        if _is_canonical_pipeline_name(label) and dkey in seen_conv_keys:
            continue
        filtered.append(d)
    return filtered


def mapear_sobres_deterministico(
    session_id: str,
    gen_docs: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Distribuye documentos en tres sobres sin LLM."""
    unified = collect_documentos_para_empaque(session_id, gen_docs)
    buckets: Dict[str, List[Dict[str, Any]]] = {
        "sobre_1": [],
        "sobre_2": [],
        "sobre_3": [],
    }
    seen_paths: Set[str] = set()
    for doc in unified:
        ruta = doc.get("ruta")
        if not ruta or not os.path.isfile(ruta):
            continue
        pkey = os.path.normcase(os.path.abspath(ruta))
        if pkey in seen_paths:
            continue
        seen_paths.add(pkey)
        sk = _classify_doc_to_sobre_key(doc)
        buckets[sk].append(doc)

    out: Dict[str, Any] = {}
    for sk, shell in _SOBR_E_SHELLS.items():
        docs = _sort_docs_for_sobre(_dedupe_doc_list(buckets[sk]))
        out[sk] = {
            **shell,
            "documentos": docs,
        }
    return out


def _sanitize_llm_estructura(
    parsed: Dict[str, Any],
    session_id: str,
    gen_docs: Dict[str, Any],
) -> Dict[str, Any]:
    """Dedup y valida rutas del JSON del LLM; rellena huecos con mapeo determinístico."""
    det = mapear_sobres_deterministico(session_id, gen_docs)
    if not isinstance(parsed, dict):
        return det

    known_paths = {
        os.path.normcase(os.path.abspath(d["ruta"]))
        for d in collect_documentos_para_empaque(session_id, gen_docs)
        if d.get("ruta") and os.path.isfile(d["ruta"])
    }

    out: Dict[str, Any] = {}
    assigned: Set[str] = set()
    for sk in ("sobre_1", "sobre_2", "sobre_3"):
        raw = parsed.get(sk) if isinstance(parsed.get(sk), dict) else {}
        shell = _SOBR_E_SHELLS[sk]
        docs_in: List[Dict[str, Any]] = []
        seen_local: Set[str] = set()
        for doc in raw.get("documentos") or []:
            if not isinstance(doc, dict):
                continue
            ruta = doc.get("ruta")
            if not ruta or not os.path.isfile(ruta):
                continue
            pkey = os.path.normcase(os.path.abspath(ruta))
            if pkey in seen_local or pkey not in known_paths:
                continue
            seen_local.add(pkey)
            assigned.add(pkey)
            docs_in.append(
                {
                    "nombre": doc.get("nombre") or os.path.basename(ruta),
                    "ruta": ruta,
                    "source_filename": doc.get("source_filename"),
                    "archivo_fuente": doc.get("archivo_fuente"),
                    "source_doc_id": doc.get("source_doc_id"),
                    "source_path": doc.get("source_path"),
                    "source_hash": doc.get("source_hash"),
                    "output_hash": doc.get("output_hash"),
                    "template_id": doc.get("template_id"),
                    "mirror_mode": doc.get("mirror_mode"),
                    "materialization_route": doc.get("materialization_route"),
                    "provenance_ui": doc.get("provenance_ui"),
                }
            )
        out[sk] = {
            "titulo": raw.get("titulo") or shell["titulo"],
            "nombre_carpeta": raw.get("nombre_carpeta") or shell["nombre_carpeta"],
            "documentos": _sort_docs_for_sobre(docs_in),
        }

    # Documentos no asignados por el LLM → clasificación determinística
    for doc in collect_documentos_para_empaque(session_id, gen_docs):
        ruta = doc.get("ruta")
        if not ruta or not os.path.isfile(ruta):
            continue
        pkey = os.path.normcase(os.path.abspath(ruta))
        if pkey in assigned:
            continue
        sk = _classify_doc_to_sobre_key(doc)
        out[sk]["documentos"].append(
            {
                "nombre": doc.get("nombre") or os.path.basename(ruta),
                "ruta": ruta,
                "source_filename": doc.get("source_filename"),
                "archivo_fuente": doc.get("archivo_fuente"),
                "source_doc_id": doc.get("source_doc_id"),
                "source_path": doc.get("source_path"),
                "source_hash": doc.get("source_hash"),
                "output_hash": doc.get("output_hash"),
                "template_id": doc.get("template_id"),
                "mirror_mode": doc.get("mirror_mode"),
                "materialization_route": doc.get("materialization_route"),
                "provenance_ui": doc.get("provenance_ui"),
            }
        )
        assigned.add(pkey)

    for sk in out:
        out[sk]["documentos"] = _sort_docs_for_sobre(
            _dedupe_doc_list(out[sk]["documentos"])
        )
    return out


def packager_sobres_stale(
    session_id: str,
    gen_docs: Optional[Dict[str, Any]] = None,
) -> bool:
    """
    True si las carpetas SOBRE_* en disco tienen menos archivos que el mapeo determinístico actual.

    Ocurre cuando packager quedó ``done`` tras una corrida parcial (p. ej. económico bloqueado).
    """
    mapped = mapear_sobres_deterministico(session_id, gen_docs)
    output_base = os.path.join("/data", "outputs", session_id)
    folder_by_key = {sk: shell["nombre_carpeta"] for sk, shell in _SOBR_E_SHELLS.items()}
    for sk, folder in folder_by_key.items():
        expected = len((mapped.get(sk) or {}).get("documentos") or [])
        sobre_dir = os.path.join(output_base, folder)
        if expected <= 0:
            continue
        if not os.path.isdir(sobre_dir):
            return True
        actual = sum(
            1
            for fn in os.listdir(sobre_dir)
            if not fn.upper().startswith("00_CARATULA")
            and os.path.isfile(os.path.join(sobre_dir, fn))
        )
        if actual < expected:
            return True
    return False


def packager_pdata_incomplete(
    session_id: str,
    packager_data: Dict[str, Any],
    gen_docs: Optional[Dict[str, Any]] = None,
) -> tuple[bool, Dict[str, int], Dict[str, int]]:
    """
    True si la estructura materializada por DocumentPackager no cubre el mapeo determinístico.

    Usar tras ``DocumentPackagerAgent`` (no comparar disco vs mapeo: puede haber carrera).
    """
    expected_map = mapear_sobres_deterministico(session_id, gen_docs)
    est = (packager_data or {}).get("estructura_sobres") or {}
    expected: Dict[str, int] = {}
    actual: Dict[str, int] = {}
    for sk in ("sobre_1", "sobre_2", "sobre_3"):
        expected[sk] = len((expected_map.get(sk) or {}).get("documentos") or [])
        actual[sk] = len((est.get(sk) or {}).get("documentos") or [])
    incomplete = any(
        actual.get(sk, 0) < expected.get(sk, 0)
        for sk in expected
        if expected.get(sk, 0) > 0
    )
    econ_dir = os.path.join("/data", "outputs", session_id, "2.propuesta_economica")
    if os.path.isdir(econ_dir):
        econ_files = sum(
            1
            for fn in os.listdir(econ_dir)
            if not fn.startswith(".") and os.path.isfile(os.path.join(econ_dir, fn))
        )
        if econ_files > 0 and actual.get("sobre_3", 0) < 1:
            incomplete = True
    return incomplete, expected, actual


def validated_pack_complete(session_id: str, *, min_ratio: float = 0.85) -> bool:
    """
    True si ``_compranet_validated`` ya refleja un empaquetado utilizable (incluye económico).

    Evita re-empaquetar CompraNet sobre un manifiesto bueno cuando ``SOBRE_*`` fue podado en disco.
    """
    import json
    from pathlib import Path

    index_path = Path("/data/outputs") / session_id / "_compranet_validated" / "INDICE_ENTREGA.json"
    if not index_path.is_file():
        return False
    try:
        payload = json.loads(index_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return False
    files = payload.get("files") or []
    if not isinstance(files, list) or not files:
        return False
    by_sobre: Dict[str, int] = {}
    for row in files:
        if isinstance(row, dict):
            label = str(row.get("sobre") or "")
            by_sobre[label] = by_sobre.get(label, 0) + 1
    if by_sobre.get("SobreEconomica", 0) < 1:
        return False
    expected_map = mapear_sobres_deterministico(session_id, None)
    exp_total = sum(len((expected_map.get(sk) or {}).get("documentos") or []) for sk in expected_map)
    if exp_total <= 0:
        return len(files) >= 3
    return len(files) >= max(3, int(exp_total * min_ratio))


class DocumentPackagerAgent(BaseAgent):
    """
    Organiza archivos generados en carpetas de sobres oficiales y carátulas.

    Contrato de rutas: ``/data/outputs/{session_id}/``.
    """

    def __init__(self, context_manager: MCPContextManager):
        super().__init__(
            agent_id="document_packager",
            name="Document Packager Agent",
            description="Organizador de expedientes de licitación en estructura de sobres oficiales.",
            context_manager=context_manager,
        )
        self.llm = ResilientLLMClient()
        self.vector_db = VectorDbServiceClient()

    async def process(self, agent_input: AgentInput) -> AgentOutput:
        session_id = agent_input.session_id
        correlation_id = agent_input.correlation_id or "no-id"
        started_at = time.perf_counter()
        print(f"[{self.name}] 📦 Iniciando empaquetado de expediente para {session_id}...", flush=True)

        gen_docs = agent_input.company_data.get("documentos_generados", {}) or {}
        master_profile = agent_input.company_data.get("master_profile", {})
        try:
            from app.services.mini_dictamen_anexos_service import (
                build_and_persist_mini_dictamen,
                build_stage_blocking_questions,
                get_blocking_annex_rows_for_stage,
            )

            await build_and_persist_mini_dictamen(self.context_manager.memory, session_id)
            session_state = await self.context_manager.memory.get_session(session_id) or {}
            blocking_rows = get_blocking_annex_rows_for_stage(session_state, "packager")
            if blocking_rows:
                session_state["pending_questions"] = build_stage_blocking_questions(
                    "packager", blocking_rows
                ) + list(session_state.get("pending_questions") or [])
                session_state["current_question_index"] = 0
                await self.context_manager.memory.save_session(session_id, session_state)
                return AgentOutput(
                    status=AgentStatus.WAITING_FOR_DATA,
                    agent_id=self.agent_id,
                    session_id=session_id,
                    message=(
                        "El empaquetado fue bloqueado por anexos obligatorios con aclaraciones "
                        "abiertas o sin cobertura válida."
                    ),
                    data={"missing": blocking_rows},
                    correlation_id=correlation_id,
                )
        except Exception:
            pass

        estructura = await self._mapear_sobres(gen_docs, session_id, correlation_id)
        mapping_mode = "deterministic" if not _packager_use_llm_mapping() else "llm_sanitized"

        output_base = os.path.join("/data", "outputs", session_id)
        reporte_sobres: Dict[str, Any] = {}
        caratulas: List[str] = []

        for key, info in estructura.items():
            if not isinstance(info, dict):
                continue
            sobre_dir = os.path.join(output_base, info["nombre_carpeta"])
            os.makedirs(sobre_dir, exist_ok=True)
            for stale_name in os.listdir(sobre_dir):
                stale_path = os.path.join(sobre_dir, stale_name)
                if os.path.isfile(stale_path):
                    os.remove(stale_path)

            print(f"[{self.name}] 📨 Organizando {info['titulo']}...", flush=True)

            docs_finales: List[Dict[str, Any]] = []
            seen_paths: Set[str] = set()
            orden = 0
            for doc in info.get("documentos") or []:
                raw_path = doc.get("ruta")
                if not raw_path or not os.path.exists(raw_path):
                    continue
                norm_path = os.path.normcase(os.path.abspath(raw_path))
                if norm_path in seen_paths:
                    continue
                seen_paths.add(norm_path)
                orden += 1
                nuevo_nombre = f"{orden:02d}_{os.path.basename(raw_path)}"
                destino = os.path.join(sobre_dir, nuevo_nombre)
                shutil.copy2(raw_path, destino)
                docs_finales.append(
                    {
                        "orden": orden,
                        "nombre": doc.get("nombre") or os.path.basename(raw_path),
                        "archivo": nuevo_nombre,
                        "source_filename": doc.get("source_filename"),
                        "archivo_fuente": doc.get("archivo_fuente"),
                        "source_doc_id": doc.get("source_doc_id"),
                        "source_path": doc.get("source_path"),
                        "source_hash": doc.get("source_hash"),
                        "template_id": doc.get("template_id"),
                        "mirror_mode": doc.get("mirror_mode"),
                        "materialization_route": doc.get("materialization_route"),
                        "provenance_ui": doc.get("provenance_ui"),
                    }
                )

            caratula_path = os.path.join(sobre_dir, "00_CARATULA_SOBRE.docx")
            self._generate_caratula(
                caratula_path, info["titulo"], docs_finales, master_profile, session_id
            )
            caratulas.append(caratula_path)

            reporte_sobres[key] = {
                "nombre": info["titulo"],
                "carpeta": sobre_dir,
                "documentos": docs_finales,
                "total_documentos": len(docs_finales),
            }

        print(
            f"[{self.name}] ✅ Expediente organizado en {len(reporte_sobres)} sobres "
            f"({mapping_mode}).",
            flush=True,
        )
        all_docs: List[Dict[str, Any]] = []
        for info in reporte_sobres.values():
            if isinstance(info, dict):
                all_docs.extend([d for d in (info.get("documentos") or []) if isinstance(d, dict)])

        return AgentOutput(
            status=AgentStatus.SUCCESS,
            agent_id=self.agent_id,
            session_id=session_id,
            data={
                "estructura_sobres": reporte_sobres,
                "caratulas_generadas": caratulas,
                "folder_raiz": output_base,
                "mapping_mode": mapping_mode,
                "materialization_metrics": build_materialization_metrics(
                    stage="document_packager",
                    documents=all_docs,
                    elapsed_ms=int((time.perf_counter() - started_at) * 1000),
                ),
            },
            correlation_id=correlation_id,
        )

    async def _mapear_sobres(
        self,
        gen_docs: Dict[str, Any],
        session_id: str,
        correlation_id: str = "",
    ) -> Dict[str, Any]:
        if not _packager_use_llm_mapping():
            print(
                f"[{self.name}] 📋 Mapeo determinístico (carpeta origen + tipo documental).",
                flush=True,
            )
            return mapear_sobres_deterministico(session_id, gen_docs)

        context_rag = await self.smart_search(
            session_id,
            "orden presentación documentos sobre técnica administrativa económica foliado",
            n_results=5,
            vector_db=self.vector_db,
        )
        parsed = await self._mapear_sobres_llm_raw(
            context_rag, gen_docs, session_id, correlation_id
        )
        return _sanitize_llm_estructura(parsed, session_id, gen_docs)

    async def _mapear_sobres_llm_raw(
        self,
        context: str,
        gen_docs: Dict[str, Any],
        session_id: str,
        correlation_id: str = "",
    ) -> Dict[str, Any]:
        documentos_disponibles = []
        for doc in collect_documentos_para_empaque(session_id, gen_docs):
            documentos_disponibles.append(
                {
                    "nombre": doc.get("nombre"),
                    "ruta": doc.get("ruta"),
                    "categoria": doc.get("categoria"),
                }
            )

        prompt = f"""
        Clasifica cada archivo en sobre_1 (administrativo), sobre_2 (técnico) o sobre_3 (económico).
        NO repitas el mismo archivo en más de un sobre. NO dupliques rutas.

        BASES (extracto):
        {context[:5000]}

        ARCHIVOS:
        {json.dumps(documentos_disponibles, indent=2)}

        JSON: {{ "sobre_1": {{ "titulo": "...", "nombre_carpeta": "SOBRE_1_ADMINISTRATIVO", "documentos": [...] }}, ... }}
        """

        resp = await self.llm.generate(prompt=prompt, format="json", correlation_id=correlation_id)
        if not resp.success:
            print(f"[{self.name}] ⚠️ LLM error: {resp.error}. Fallback determinístico.", flush=True)
            return {}

        raw_text = resp.response or ""
        try:
            if "```json" in raw_text:
                raw_text = raw_text.split("```json")[1].split("```")[0].strip()
            elif "```" in raw_text:
                raw_text = raw_text.split("```")[1].split("```")[0].strip()
            start = raw_text.find("{")
            end = raw_text.rfind("}")
            if start != -1 and end != -1:
                return json.loads(raw_text[start : end + 1])
        except Exception as e:
            print(f"[{self.name}] ⚠️ JSON inválido ({e}).", flush=True)
        return {}

    def _fallback_estructura_por_claves(self, gen_docs: Dict[str, Any]) -> Dict[str, Any]:
        """Compatibilidad tests: alias del mapeo por categoría de generación."""
        return mapear_sobres_deterministico("", gen_docs)

    def _generate_caratula(
        self,
        path: str,
        titulo: str,
        docs: List[Dict[str, Any]],
        profile: Dict[str, Any],
        session_id: str,
    ) -> None:
        doc = Document()
        for section in doc.sections:
            section.top_margin = docx_shared.Inches(1)

        doc.add_paragraph("\n" * 2)
        p_titulo = doc.add_paragraph()
        run_titulo = p_titulo.add_run(titulo.upper())
        run_titulo.bold = True
        run_titulo.font.size = docx_shared.Pt(24)
        p_titulo.alignment = 1

        doc.add_paragraph("-" * 40).alignment = 1

        p_licit = doc.add_paragraph()
        p_licit.add_run(f"LICITACIÓN: {session_id.upper()}").bold = True
        p_licit.alignment = 1

        doc.add_paragraph("\n")

        p_empresa = doc.add_paragraph()
        p_empresa.add_run(f"EMPRESA: {profile.get('razon_social', '...')}\n").bold = True
        p_empresa.add_run(f"RFC: {profile.get('rfc', '...')}\n")
        p_empresa.add_run(f"REPRESENTANTE: {profile.get('representante_legal', '...')}")
        p_empresa.alignment = 1

        doc.add_paragraph("\n" * 2)
        doc.add_heading("ÍNDICE DE CONTENIDO", 2)

        for doc_item in docs:
            doc.add_paragraph(
                f"{doc_item['orden']}. {doc_item['nombre']}", style="List Bullet"
            )

        doc.add_paragraph("\n" * 3)
        doc.add_paragraph(
            f"FECHA DE GENERACIÓN: {datetime.now().strftime('%d/%m/%Y')}"
        ).alignment = 1

        doc.save(path)
