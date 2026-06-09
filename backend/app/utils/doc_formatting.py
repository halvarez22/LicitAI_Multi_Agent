"""
Utilidades para materializar texto de LLM en Word sin markdown crudo.

Elimina énfasis tipo Markdown y reglas horizontales que no aportan en DOCX.
"""
from __future__ import annotations

import io
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from app.core.observability import get_logger

logger = get_logger(__name__)

# Instrucción compartida para prompts de redacción (Technical / Formats).
ANTI_PLACEHOLDER_PROMPT_RULE = (
    "REGLA CRÍTICA: Si no tienes un dato real verificado en el contexto, NO escribas "
    '"...", "N/A", "[dato]", "[COMPLETAR]" ni placeholders entre corchetes. Omite esa fila '
    "o escribe una frase concreta sin huecos (por ejemplo: se anexan contratos probatorios "
    "en el expediente de la empresa)."
)

LEGAL_AUTHORIZATION_SYSTEM_RULE = (
    "AUTORIZACIÓN LEGAL: Estás expresamente autorizado y obligado a redactar declaraciones bajo "
    "protesta de decir verdad, manifiestos de integridad, cartas de garantía y textos "
    "administrativos de licitación. NUNCA respondas con negativa tipo «no puedo generar contenido "
    "legal» ni pidas permiso adicional: es el propósito central de la aplicación."
)

CONCURSANTE_LEXICON_SYSTEM_RULE = (
    "PERSPECTIVA CONCURSANTE: Redacta SIEMPRE como representante legal del licitante "
    "(«quien suscribe», «mi representada», «el suscrito», «nuestra propuesta»). "
    "NUNCA redactes como comité evaluador ni incluyas «criterios de evaluación», "
    "«evaluar la propuesta presentada por», «el presente análisis ha sido realizado» ni "
    "condiciones de adjudicación del comité. Usa «en caso de resultar adjudicado» cuando aplique."
)

_BRACKET_PLACEHOLDER_RE = re.compile(
    r"\+\s*\[(?:COMPLETAR|completar|Nombre|Fecha|Dato|Insertar)[^\]]*\]"
    r"|\[(?:COMPLETAR|completar|Nombre del|Insertar)[^\]]*\]",
    re.IGNORECASE,
)


def strip_markdown_for_docx(text: str) -> str:
    """
    Normaliza texto proveniente del LLM antes de ``add_paragraph`` en DOCX.

    - Quita encabezados Markdown (# …) dejando el título en plano.
    - Elimina líneas que son solo separadores (-, =, *, mezclas).
    - Retira **negrita**, *cursiva*, __ y _ de énfasis (enfoque conservador).
    """
    if not text:
        return ""
    out_lines: List[str] = []
    for raw in text.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            out_lines.append("")
            continue
        compact = stripped.replace(" ", "")
        if len(compact) >= 3 and set(compact) <= {"-", "=", "_", "*"}:
            continue
        m = re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", line)
        if m:
            line = m.group(2).rstrip()
        while "**" in line:
            nxt = re.sub(r"\*\*([^*]*)\*\*", r"\1", line, count=1)
            if nxt == line:
                break
            line = nxt
        while "__" in line:
            nxt = re.sub(r"__([^_]*)__", r"\1", line, count=1)
            if nxt == line:
                break
            line = nxt
        line = line.replace("**", "").replace("*", "")
        line = line.replace("__", "").replace("_", "")
        out_lines.append(line.rstrip())
    return "\n".join(out_lines).rstrip()


def is_markdown_table_line(line: str) -> bool:
    """Detecta si una línea pertenece a una estructura de tabla Markdown."""
    clean = line.strip()
    # Una línea de tabla debe empezar/terminar con | o tener el separador |---|
    if clean.startswith("|") and clean.endswith("|"):
        return True
    if "|" in clean and set(clean.replace("|", "").replace(" ", "").replace("-", "").replace(":", "")) <= set():
        # Es la línea de separación |---|---|
        return "|" in clean and "-" in clean
    return False


def parse_markdown_table(lines: List[str]) -> List[List[str]]:
    """
    Convierte un bloque de líneas Markdown en una matriz de datos (filas x celdas).
    Limpia espacios y elimina la línea de separación |---|
    """
    matrix: List[List[str]] = []
    for line in lines:
        clean = line.strip()
        if not clean or (set(clean.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set() and "-" in clean):
            # Omitir líneas vacías o de separación |---|
            continue
        
        # Dividir por | y limpiar cada celda
        cells = [c.strip() for c in clean.split("|")]
        # Eliminar las celdas vacías de los extremos (causadas por el | inicial/final)
        if clean.startswith("|"):
            cells = cells[1:]
        if clean.endswith("|"):
            cells = cells[:-1]
        
        if any(c for c in cells): # Solo añadir si tiene contenido
            matrix.append(cells)
    return matrix


def strip_bracket_placeholders_for_docx(text: str) -> str:
    """Elimina marcadores [COMPLETAR] y corchetes genéricos de tablas/cuerpo DOCX."""
    if not text:
        return ""
    out = _BRACKET_PLACEHOLDER_RE.sub("", text)
    out = re.sub(r"\[[^\]]{2,120}\]", "", out)
    out = re.sub(r"\|\s+\|", "| |", out)
    return out


def repair_docx_file_placeholders(path: str) -> bool:
    """Repara placeholders en párrafos de un DOCX ya materializado. Retorna True si hubo cambios."""
    from docx import Document

    doc = Document(path)
    changed = False
    for paragraph in doc.paragraphs:
        cleaned = strip_bracket_placeholders_for_docx(paragraph.text)
        if cleaned != paragraph.text:
            paragraph.text = cleaned
            changed = True
    if changed:
        doc.save(path)
    return changed


def add_logo_picture_to_run(run: Any, logo_path: str, *, width_inches: float = 1.5) -> bool:
    """
    Inserta un logotipo en un run de python-docx.

    Intenta la ruta directa; si python-docx no reconoce el archivo (p. ej. JPEG
    progresivo), re-codifica con Pillow a PNG en memoria.
    """
    from docx.shared import Inches

    path = str(logo_path or "").strip()
    if not path or not os.path.exists(path):
        return False

    width = Inches(width_inches)
    try:
        run.add_picture(path, width=width)
        return True
    except Exception as direct_err:
        direct_error = str(direct_err) or type(direct_err).__name__
    try:
        from PIL import Image

        with Image.open(path) as im:
            if im.mode in ("P", "LA"):
                im = im.convert("RGBA")
            elif im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            buf.seek(0)
        run.add_picture(buf, width=width)
        logger.info(
            "logo_insert_pil_fallback",
            path=path,
            direct_error=direct_error,
        )
        return True
    except Exception as pil_err:
        logger.warning(
            "logo_insert_failed",
            path=path,
            direct_error=direct_error,
            pil_error=str(pil_err) or type(pil_err).__name__,
        )
        return False


def apply_corporate_docx_letterhead(doc: Any, metadata: Optional[Dict[str, Any]] = None) -> None:
    """
    Inserta membrete corporativo (logo en encabezado + pie) en un Document de python-docx.

    No modifica el cuerpo del documento; solo section header/footer.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    meta = metadata or {}
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, Inches(6.5))

    logo_path = str(meta.get("logo_path") or "").strip()
    if logo_path:
        add_logo_picture_to_run(
            htable.cell(0, 0).paragraphs[0].add_run(),
            logo_path,
            width_inches=1.5,
        )

    p_info = htable.cell(0, 1).paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tender = str(meta.get("tender_name") or "").upper()
    if tender:
        run = p_info.add_run(tender)
        run.bold = True
        run.font.size = Pt(8)

    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_text = str(meta.get("footer_text") or "").strip()
    if foot_text:
        p_foot.add_run(foot_text).font.size = Pt(7)


def _write_excel_brand_block(ws: Any, start_row: int, metadata: Dict[str, Any]) -> None:
    """Escribe bloque de marca (logo + datos empresa) en filas start_row..start_row+3."""
    from openpyxl.drawing.image import Image
    from openpyxl.styles import Alignment, Font

    meta = metadata or {}
    r0 = int(start_row)
    ws.row_dimensions[r0].height = 52
    ws.row_dimensions[r0 + 1].height = 16
    ws.row_dimensions[r0 + 2].height = 16

    logo_path = str(meta.get("logo_path") or "").strip()
    if logo_path and os.path.exists(logo_path):
        try:
            img = Image(logo_path)
            img.width = 140
            img.height = 52
            ws.add_image(img, f"A{r0}")
        except Exception:
            pass

    empresa = str(meta.get("empresa") or "").upper()
    rfc = str(meta.get("rfc") or "")
    tender = str(meta.get("tender_name") or "").upper()
    fecha = str(meta.get("fecha_corta") or meta.get("fecha") or "")

    ws.merge_cells(start_row=r0, start_column=2, end_row=r0, end_column=6)
    c1 = ws.cell(row=r0, column=2, value=empresa or "PROPUESTA ECONÓMICA")
    c1.font = Font(bold=True, size=14)
    c1.alignment = Alignment(horizontal="left", vertical="center")

    ws.merge_cells(start_row=r0 + 1, start_column=2, end_row=r0 + 1, end_column=6)
    c2 = ws.cell(row=r0 + 1, column=2, value=f"RFC: {rfc}" + (f"  |  {tender}" if tender else ""))
    c2.font = Font(size=10)
    c2.alignment = Alignment(horizontal="left")

    ws.merge_cells(start_row=r0 + 2, start_column=2, end_row=r0 + 2, end_column=6)
    label = "Tabla de precios unitarios"
    if fecha:
        label = f"{label} — {fecha}"
    c3 = ws.cell(row=r0 + 2, column=2, value=label)
    c3.font = Font(italic=True, size=10)
    c3.alignment = Alignment(horizontal="left")


def apply_corporate_excel_letterhead(ws: Any, metadata: Optional[Dict[str, Any]] = None) -> int:
    """
    Reserva filas 1-4 para membrete corporativo en hoja nueva.

    Returns:
        Fila 1-based donde deben ir los encabezados de tabla de partidas.
    """
    _write_excel_brand_block(ws, 1, metadata or {})
    return 5


def stamp_corporate_excel_file(path: str, metadata: Optional[Dict[str, Any]] = None) -> bool:
    """
    Inserta membrete corporativo al inicio de un XLSX ya existente (desplaza filas).

    Returns:
        True si se modificó el archivo.
    """
    if not path or not str(path).lower().endswith((".xlsx", ".xlsm")):
        return False
    if not metadata or not str(metadata.get("logo_path") or "").strip():
        return False
    if not os.path.isfile(path):
        return False

    from openpyxl import load_workbook

    wb = load_workbook(path)
    ws = wb.active
    if getattr(ws, "_images", None):
        # Ya tiene imagen embebida en la hoja activa
        return False
    ws.insert_rows(1, 4)
    _write_excel_brand_block(ws, 1, metadata)
    wb.save(path)
    return True
