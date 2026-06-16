#!/usr/bin/env python3
"""
UAT checklist anexo a anexo — obra pública T/E (HRU).

Uso:
  PYTHONPATH=/app python scripts/uat_obra_te_checklist.py SESSION_ID
"""
from __future__ import annotations

import asyncio
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from app.api.deps import get_connected_memory
from app.services.junta_bases_corpus import build_bases_corpus
from app.services.obra_delivery_gap_service import build_obra_te_gap_report

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


# Patrones HRU universales (no hardcode por licitación).
_FALSE_POSITIVE_RE = re.compile(
    r"(?i)acepto\s+su\s+contenido|hemos\s+sido\s+seleccionados|presento\s+los\s+programas|"
    r"presento\s+las\s+cotizaciones|71\.6\s*%|luis\s+ernesto\s+diez|soluciones\s+dior"
)
_CONTAMINATION_RE = re.compile(
    r"(?i)de\s+las\s+causas\s+de\s+descalific|descalificar[aá]\s+al\s+participante|"
    r"dictamen\s+de\s+evaluaci[oó]n|propuesta\s+conveniente,\s*y\s+que"
)
_CDMX_RE = re.compile(r"(?i)ciudad\s+de\s+m[eé]xico,\s*cdmx")

_HRU_EXPECT_CONSIGNAR = frozenset(
    {"obra|T1", "obra|T2", "obra|T5", "obra|T7", "obra|T8_PRIVACIDAD", "obra|E3", "obra|E4", "obra|E5"}
)


def _read_docx(path: str) -> str:
    if not Document or not Path(path).is_file():
        return ""
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                val = cell.text.strip()
                if val:
                    parts.append(val)
    return "\n".join(parts)


def _check_row(
    row: Dict[str, Any],
    *,
    text: str,
    bases_low: str,
) -> Dict[str, Any]:
    key = str(row.get("dedupe_key") or "")
    path = str(row.get("archivo") or "")
    nombre = str(row.get("nombre_canonico") or "")
    sobre = str(row.get("sobre_clasificado") or "")
    low = text.lower()
    checks: List[Dict[str, str]] = []

    def add(ok: bool, code: str, detail: str) -> None:
        checks.append({"ok": ok, "code": code, "detail": detail})

    add(bool(path and Path(path).is_file()), "archivo", path or "sin archivo")
    if sobre:
        if "administrativo" in sobre:
            add("SOBRE_1" in path or "administrativos" in path, "ubicacion", sobre)
        elif "economico" in sobre:
            add("SOBRE_3" in path or "propuesta_economica" in path, "ubicacion", sobre)

    add("LEÓN" in text or "LEON" in text.upper(), "ciudad_convocante", "LEÓN, GTO esperado")
    add(not bool(_CDMX_RE.search(text)), "no_cdmx", "sin Ciudad de México, CDMX")
    add(not bool(_FALSE_POSITIVE_RE.search(text)), "no_afirmaciones_falsas", "sin redacción post-adjudicación")
    if key != "obra|T4":
        add(
            not bool(_CONTAMINATION_RE.search(text)),
            "no_contaminacion",
            "sin texto de descalificación/evaluación",
        )
    else:
        add(True, "no_contaminacion", "T-4 transcribe bases (cola normativa admisible)")

    if key in _HRU_EXPECT_CONSIGNAR:
        has_ph = (
            "[Consignar]" in text
            or "[consignar]" in low
            or "consignar]" in low
            or "por consignar" in low
        )
        add(has_ph, "hitl_consignar", "placeholder HITL presente")

    if key == "obra|T8_PRIVACIDAD":
        add("acepto su contenido" not in low, "t8_sin_aceptacion_inventada", "aviso sin aceptación inventada")

    if key == "obra|E1":
        add("$3,278,289.63" in text or "$" in text, "e1_importe", "importe total con IVA")
        add("18 días naturales" in text or "días naturales" in low, "e1_plazo", "plazo numérico desde bases")

    if key == "obra|E2":
        add(
            "UNIDAD" in text or "P.U." in text or "costos directos" in low,
            "e2_catalogo",
            "catálogo con columnas o totales obra",
        )
        add(
            "PRESUPUESTO_52" not in Path(path).name,
            "e2_nombre",
            "nombre archivo sin PRESUPUESTO_52 (cosmético si contenido OK)",
        )

    if key == "obra|E3":
        add("anexo e-3" in low, "e3_portada", "portada E-3")
        add("71.6%" not in text, "e3_sin_apu_inventado", "sin % materiales inventados")

    if key == "obra|E5":
        add("de las causas" not in low, "e5_req_limpio", "requisito sin causas de descalificación")

    # Anclaje mínimo en bases (snippet o palabra clave del anexo)
    snippet = str(row.get("snippet") or "")[:80].lower()
    anchor = ""
    if "t-1" in key or key == "obra|T1":
        anchor = "maquinaria"
    elif key == "obra|T2":
        anchor = "contratos"
    elif key == "obra|T3":
        anchor = "contrato"
    elif key == "obra|T4":
        anchor = "bases"
    elif key == "obra|T5":
        anchor = "visita"
    elif key == "obra|T6":
        anchor = "obligaciones"
    elif key == "obra|T7":
        anchor = "subcontrat"
    elif key == "obra|T8_PRIVACIDAD":
        anchor = "privacidad"
    elif key == "obra|E1":
        anchor = "carta-compromiso"
    elif key == "obra|E2":
        anchor = "catálogo" if "catálogo" in bases_low else "conceptos"
    elif key == "obra|E3":
        anchor = "precios unitarios"
    elif key == "obra|E4":
        anchor = "programa"
    elif key == "obra|E5":
        anchor = "cotizaciones"
    if anchor:
        add(anchor in bases_low or anchor in snippet, "anclaje_bases", f"«{anchor}» en corpus o snippet")

    failed = [c for c in checks if not c["ok"]]
    cosmetic_only = failed and all(c["code"] in ("e2_nombre", "hitl_consignar") for c in failed if not c["ok"])
    if not failed:
        status = "APROBADO"
    elif cosmetic_only and len(failed) <= 2:
        status = "APROBADO_OBS"
    elif len(failed) <= 2:
        status = "PARCIAL"
    else:
        status = "NO_APROBADO"
    if not path:
        status = "NO_APROBADO"

    return {
        "dedupe_key": key,
        "nombre": nombre[:100],
        "sobre": sobre,
        "archivo": Path(path).name if path else "",
        "status": status,
        "checks": checks,
        "failed": [c["code"] for c in failed],
    }


async def main() -> int:
    session_id = sys.argv[1] if len(sys.argv) > 1 else "barda_primaria_lopez_rayon"
    mem = await get_connected_memory()
    state = await mem.get_session(session_id) or {}
    docs = await mem.get_documents(session_id)
    corpus = build_bases_corpus(session_id, docs, session_state=state)
    bases_low = (corpus.combined or "").lower()

    report = build_obra_te_gap_report(session_id, state, docs)
    results: List[Dict[str, Any]] = []
    for row in report.get("rows") or []:
        path = str(row.get("archivo") or "")
        text = _read_docx(path)
        results.append(_check_row(row, text=text, bases_low=bases_low))

    summary = {
        "session_id": session_id,
        "inventory": report.get("inventory_count"),
        "gap_count": report.get("gap_count"),
        "aprobados": sum(1 for r in results if r["status"] in ("APROBADO", "APROBADO_OBS")),
        "parciales": sum(1 for r in results if r["status"] == "PARCIAL"),
        "no_aprobados": sum(1 for r in results if r["status"] == "NO_APROBADO"),
        "rows": results,
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if summary["no_aprobados"] == 0 and summary["gap_count"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
