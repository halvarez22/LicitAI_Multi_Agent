import sys
from docx import Document

path = "/app/../../../costos de servicios/Formato hoja membretada CMyT ZEN ofertas.docx"
try:
    doc = Document(path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{i}: {para.text}")
    for i, table in enumerate(doc.tables):
        print(f"\n=== TABLA {i} ===")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            print(" | ".join(cells))
except Exception as e:
    print("Error:", e)
    # Try alternate path
    import os
    base = "/app"
    for root, dirs, files in os.walk("/"):
        for f in files:
            if "membretada" in f.lower() or "CMyT" in f:
                print("FOUND:", os.path.join(root, f))
        if root.count(os.sep) > 6:
            break
