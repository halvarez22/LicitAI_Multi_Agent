"""Resolver de archivos ingestados y espejo de plantillas."""

import os
import tempfile
from pathlib import Path

from app.services.ingested_file_resolver import (
    build_ingested_file_index,
    resolve_ingested_file,
)
from app.services.template_mirror_service import fill_docx_with_profile, mirror_template_to_output


def test_resolve_ingested_by_filename():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "16. Anexo III P 1 Zona A.xlsx"
        path.write_bytes(b"xlsx")
        docs = [
            {
                "id": "d1",
                "content": {
                    "filename": "16. Anexo III P 1 Zona A.xlsx",
                    "file_path": str(path),
                },
                "metadata": {},
            }
        ]
        idx = build_ingested_file_index(docs)
        ref = resolve_ingested_file("16. Anexo III P 1 Zona A.xlsx", idx)
        assert ref is not None
        assert ref.exists


def test_mirror_docx_copy():
    with tempfile.TemporaryDirectory() as tmp:
        try:
            from docx import Document
        except ImportError:
            return
        src = Path(tmp) / "formato.docx"
        out = Path(tmp) / "out.docx"
        doc = Document()
        doc.add_paragraph("RFC: {{RFC}}")
        doc.save(str(src))
        docs = [
            {
                "id": "1",
                "content": {"filename": "formato.docx", "file_path": str(src)},
                "metadata": {},
            }
        ]
        idx = build_ingested_file_index(docs)
        ref = resolve_ingested_file("formato.docx", idx)
        assert ref is not None
        meta = mirror_template_to_output(
            ref,
            str(out),
            {"rfc": "ABC123456789", "razon_social": "Empresa Test"},
        )
        assert os.path.isfile(meta["ruta"])
        assert meta["mirror_mode"].startswith("copy_docx")


def test_resolve_ingested_by_doc_id_and_source_path():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "14. Anexo S Modelo Contrato Federal.doc"
        path.write_bytes(b"doc")
        docs = [
            {
                "id": "doc-anexo-s",
                "content": {
                    "filename": "14. Anexo S Modelo Contrato Federal.doc",
                    "file_path": str(path),
                },
                "metadata": {},
            }
        ]
        idx = build_ingested_file_index(docs)

        ref_by_id = resolve_ingested_file(
            "nombre-incorrecto.doc",
            idx,
            doc_id="doc-anexo-s",
        )
        ref_by_path = resolve_ingested_file(
            "otro-nombre.doc",
            idx,
            source_path=str(path),
        )

        assert ref_by_id is not None
        assert ref_by_id.filename == "14. Anexo S Modelo Contrato Federal.doc"
        assert ref_by_path is not None
        assert ref_by_path.file_path == str(path)


def test_fill_docx_with_profile_rellena_fecha_zona_y_denominacion():
    with tempfile.TemporaryDirectory() as tmp:
        try:
            from docx import Document
        except ImportError:
            return
        path = Path(tmp) / "modelo.docx"
        doc = Document()
        doc.add_paragraph("__________de __________ de ______________ (1)")
        doc.add_paragraph("Zona: _____")
        doc.add_paragraph("DENOMINACIÓN SOCIAL: __________")
        doc.add_paragraph("DOMICILIO: __________________.")
        doc.add_paragraph("Me refiero al procedimiento ISAPEG(3)___________, No. ____(4)_______ en el que mi representada, la empresa________(5)___________ participa a través del presente sobre.")
        doc.save(str(path))

        fill_docx_with_profile(
            str(path),
            {
                "fecha": "26 de mayo de 2026",
                "zonas_ofertadas": "A, C y D",
                "razon_social": "Comercializadora Mayo",
                "domicilio": "León, Guanajuato",
                "licitacion": "ISAPEG-2026",
            },
        )

        out = Document(str(path))
        text = "\n".join(p.text for p in out.paragraphs)
        assert "26 de mayo de 2026" in text
        assert "Zona: A, C y D" in text
        assert "DENOMINACIÓN SOCIAL: Comercializadora Mayo" in text
        assert "DOMICILIO: León, Guanajuato" in text
        assert "No. ISAPEG-2026" in text
        assert "empresa Comercializadora Mayo participa" in text


def test_fill_docx_with_profile_rellena_tarifa_mensual_y_numero_referencia():
    with tempfile.TemporaryDirectory() as tmp:
        try:
            from docx import Document
        except ImportError:
            return
        path = Path(tmp) / "anexo_d_iii.docx"
        doc = Document()
        doc.add_paragraph("Tarifa mensual para horario: _______________")
        doc.add_paragraph("Número (s): ____________________")
        doc.save(str(path))

        fill_docx_with_profile(
            str(path),
            {
                "tarifa_mensual_referencia": 13326.63,
                "numero_referencia": "ISAPEG-2026",
            },
        )

        out = Document(str(path))
        text = "\n".join(p.text for p in out.paragraphs)
        assert "Tarifa mensual para horario: $13,326.63 MXN" in text
        assert "Número (s): ISAPEG-2026" in text
