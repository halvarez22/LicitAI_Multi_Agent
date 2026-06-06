"""Tests de reparación de placeholders en DOCX materializados."""
from pathlib import Path

from docx import Document

from app.services.document_fill_quality_gate import validate_generated_documents_fill, _is_placeholder, _scan_docx
from app.utils.doc_formatting import repair_docx_file_placeholders, strip_bracket_placeholders_for_docx


def test_strip_completar_marker():
    raw = "| Cliente | Monto |\n| + [COMPLETAR] | + [COMPLETAR] |"
    out = strip_bracket_placeholders_for_docx(raw)
    assert "[COMPLETAR]" not in out
    assert "Cliente" in out


def test_strip_client_table_placeholders():
    raw = (
        "| Número | Nombre del Cliente | Domicilio | Teléfono |\n"
        "| 1 |  | [Domicilio del cliente 1] | [Teléfono del cliente 1] |\n"
        "Se anexan contratos probatorios en el expediente de la empresa."
    )
    out = strip_bracket_placeholders_for_docx(raw)
    assert "[Domicilio" not in out
    assert "Se anexan contratos" in out


def test_repair_docx_removes_client_table_placeholders(tmp_path, monkeypatch):
    from app.config.settings import settings as app_settings

    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    f = tmp_path / "02_TE-03_Propuesta_tecnica.docx"
    doc = Document()
    doc.add_paragraph("SEGURIDAD PRIVADA INTEGRAL MANAVIL SA DE CV")
    doc.add_paragraph("RFC: SPI060200AG5")
    doc.add_paragraph("Representante: YUNUEN IVON ACEVES SANCHEZ")
    doc.add_paragraph("| 1 |  | [Domicilio del cliente 1] | [Teléfono del cliente 1] |")
    doc.add_paragraph("Se anexan contratos probatorios en el expediente de la empresa.")
    doc.save(f)

    assert repair_docx_file_placeholders(str(f)) is True
    chunks = _scan_docx(str(f))
    assert not any(_is_placeholder(c) for c in chunks)
    gate = validate_generated_documents_fill(
        stage="technical",
        generated_documents=[{"ruta": str(f), "tipo": "tecnico"}],
        master_profile={
            "razon_social": "SEGURIDAD PRIVADA INTEGRAL MANAVIL SA DE CV",
            "rfc": "SPI060200AG5",
            "representante_legal": "YUNUEN IVON ACEVES SANCHEZ",
        },
        provenance_context={"source": "technical_writer", "confidence": 0.9},
    )
    assert gate.get("validation_passed") is True
