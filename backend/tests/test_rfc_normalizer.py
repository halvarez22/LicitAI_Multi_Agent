"""Normalización universal de RFC (SAT)."""

from __future__ import annotations

from app.api.v1.routes.companies import _canonicalize_profile_rfc, _normalize_profile_scalar_fields
from app.config.settings import settings as app_settings
from app.services.document_fill_quality_gate import validate_generated_documents_fill
from app.utils.rfc_normalizer import (
    find_rfcs_in_text,
    is_valid_rfc_sat,
    normalize_rfc_sat,
    rfc_present_in_text,
    strip_rfc_formatting,
)
from docx import Document


def test_strip_rfc_formatting_removes_separators_and_prefix():
    assert strip_rfc_formatting("RFC: CIN-250608-9A3") == "CIN2506089A3"
    assert strip_rfc_formatting("cmt 160107 s83") == "CMT160107S83"
    assert strip_rfc_formatting("ACM.010101.AAA") == "ACM010101AAA"


def test_normalize_rfc_sat_persona_moral_y_fisica():
    assert normalize_rfc_sat("CIN-250608-9A3") == "CIN2506089A3"
    assert normalize_rfc_sat("CMT160107S83") == "CMT160107S83"
    assert normalize_rfc_sat("TODE820602-FR4") == "TODE820602FR4"


def test_normalize_rfc_sat_rechaza_invalidos():
    assert normalize_rfc_sat("RFC_MALO") is None
    assert normalize_rfc_sat("") is None
    assert normalize_rfc_sat("NO ENCONTRADO") is None
    assert not is_valid_rfc_sat("ABC-12")


def test_find_rfcs_in_text_free_form():
    blob = "El contribuyente RFC CIN-250608-9A3 y también CMT160107S83 aparecen."
    found = find_rfcs_in_text(blob)
    assert "CIN2506089A3" in found
    assert "CMT160107S83" in found


def test_rfc_present_in_text_with_separators():
    assert rfc_present_in_text("RFC: CIN-250608-9A3", "CIN2506089A3")
    assert not rfc_present_in_text("sin datos", "CIN2506089A3")


def test_companies_normalize_profile_scalar_fields_rfc():
    profile = {"rfc": "CIN-250608-9A3"}
    _normalize_profile_scalar_fields(profile)
    assert profile["rfc"] == "CIN2506089A3"


def test_companies_canonicalize_invalid_keeps_upper():
    assert _canonicalize_profile_rfc("RFC_MALO") == "RFC_MALO"


def test_fill_gate_accepts_hyphenated_rfc_in_profile_and_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    f = tmp_path / "manifestacion.docx"
    doc = Document()
    doc.add_paragraph("CONSTRUCTORA SA DE CV")
    doc.add_paragraph("RFC: CIN-250608-9A3")
    doc.add_paragraph("Representante legal")
    doc.save(f)

    out = validate_generated_documents_fill(
        stage="formats",
        generated_documents=[
            {
                "ruta": str(f),
                "tipo": "administrativo",
                "template_id": "manifestacion_subcontratacion",
            }
        ],
        master_profile={
            "razon_social": "CONSTRUCTORA SA DE CV",
            "rfc": "CIN-250608-9A3",
            "representante_legal": "Juan Perez",
        },
        provenance_context={"source": "formats_writer", "confidence": 0.9},
    )
    rfc_issues = [
        i
        for i in out["issues"]
        if i.get("field_key") == "rfc" and i.get("error_type") == "cross_field_inconsistency"
    ]
    assert not rfc_issues
