"""Extracción heurística de partidas desde Excel (sin Postgres)."""

import pandas as pd
from docx import Document

from app.services.tabular_line_item_extract import (
    extract_line_items_from_docx_path,
    extract_line_items_from_excel_path,
)


def test_extract_line_items_detecta_concepto_y_precio(tmp_path):
    path = tmp_path / "costos.xlsx"
    df = pd.DataFrame(
        {
            "Concepto": ["Servicio de limpieza", "Vigilancia"],
            "Unidad": ["m2", "hora"],
            "Cantidad": [100, 720],
            "Precio unitario": [12.5, 45.0],
        }
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="Partidas", index=False)

    rows = extract_line_items_from_excel_path(str(path), "costos.xlsx")
    assert len(rows) >= 2
    by_concept = {r["concepto_norm"]: r for r in rows}
    assert "servicio de limpieza" in by_concept
    assert by_concept["servicio de limpieza"]["precio_unitario"] == 12.5
    assert by_concept["servicio de limpieza"]["unidad"] == "m2"
    assert by_concept["servicio de limpieza"]["extra"]["price_column_index"] == 3
    assert "vigilancia" in by_concept
    assert by_concept["vigilancia"]["precio_unitario"] == 45.0


def test_extract_line_items_desde_docx_con_tabla_economica(tmp_path):
    path = tmp_path / "oferta.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Partida"
    table.rows[0].cells[1].text = "Descripcion"
    table.rows[0].cells[2].text = "Unidad"
    table.rows[0].cells[3].text = "Precio Unitario"
    row1 = table.add_row().cells
    row1[0].text = "1"
    row1[1].text = "Suministro e instalacion de paneles solares"
    row1[2].text = "Lote"
    row1[3].text = "$2,586,233.00"
    row2 = table.add_row().cells
    row2[0].text = "2"
    row2[1].text = "Servicio UVIE"
    row2[2].text = "Servicio"
    row2[3].text = "150000"
    doc.save(path)

    rows = extract_line_items_from_docx_path(str(path), "oferta.docx")
    assert len(rows) >= 2
    by_concept = {r["concepto_norm"]: r for r in rows}
    assert "suministro e instalacion de paneles solares" in by_concept
    assert by_concept["suministro e instalacion de paneles solares"]["precio_unitario"] == 2586233.0
    assert by_concept["suministro e instalacion de paneles solares"]["unidad"] == "Lote"
    assert "servicio uvie" in by_concept
    assert by_concept["servicio uvie"]["precio_unitario"] == 150000.0


def test_extract_line_items_docx_dedupes_repeated_tables(tmp_path):
    """Dos tablas idénticas en el DOCX no deben duplicar la misma partida en cuadratura."""
    from docx import Document

    path = tmp_path / "oferta_dup.docx"
    doc = Document()
    for _ in range(2):
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = "Partida"
        table.rows[0].cells[1].text = "Descripcion"
        table.rows[0].cells[2].text = "Unidad"
        table.rows[0].cells[3].text = "Precio Unitario"
        r1 = table.rows[1].cells
        r1[0].text = "1"
        r1[1].text = "Suministro e instalacion de paneles solares"
        r1[2].text = "Lote"
        r1[3].text = "2586233"
    doc.save(path)

    rows = extract_line_items_from_docx_path(str(path), "oferta_dup.docx")
    assert len(rows) == 1
    assert rows[0]["precio_unitario"] == 2586233.0


def test_extract_line_items_blank_price_templates_preserve_quantity_and_structure(tmp_path):
    path = tmp_path / "anexo_blank.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        service = pd.DataFrame(
            [
                ["SERVICIO DE LIMPIEZA", "", "", "", "", "", "", "", ""],
                ["ZONA", "NÚM.", "UNIDAD", "DOMICILIO", "CIUDAD", "NÚM. DE ELEMENTOS", "HORARIO", "", ""],
                ["A", "LIAI-001", "CAISES GUANAJUATO", "PARDO 5", "GUANAJUATO", 6, "LUNES A VIERNES (8 HORAS)", "", ""],
                ["A", "LIAI-002", "UMAPS ARPEROS", "MORELOS 47", "GUANAJUATO", 1, "LUNES A VIERNES (8 HORAS)", "", ""],
            ]
        )
        materials = pd.DataFrame(
            [
                ["", "ANEXO III PARTIDA 2 ZONA A PROPUESTA ECONÓMICA", "", "", "", "", "", ""],
                ["", "No.", "Descripción del material", "Presentación", "Cantidad mensual", "Costo Unitario (I.V.A. incluido)", "COSTO MENSUAL I.V.A INCLUIDO", "COSTO TOTAL I.V.A INCLUIDO"],
                ["", 1, "BOLSA DE PLÁSTICO CHICA 55X60", "KILO", 1528, "", "", ""],
                ["", 2, "CEPILLO DE PLÁSTICO", "PIEZA", 197, "", "", ""],
            ]
        )
        service.to_excel(w, sheet_name="PARTIDA 1 ZONA A", index=False, header=False)
        materials.to_excel(w, sheet_name="PARTIDA 2 ZONA A", index=False, header=False)

    rows = extract_line_items_from_excel_path(str(path), "anexo_blank.xlsx")
    assert len(rows) == 4

    by_concept = {r["concepto_norm"]: r for r in rows}
    site = by_concept["caises guanajuato"]
    assert site["cantidad"] == 6
    assert site["precio_unitario"] == 0.0
    assert site["unidad"] == "ELEMENTO"
    assert site["extra"]["layout"] == "structured_template"
    assert site["extra"]["template_kind"] == "service_zone_elements"
    assert site["extra"]["zone"] == "A"
    assert "lunes a viernes" in str(site["extra"]["schedule"]).lower()
    assert site["extra"]["price_input_kind"] == "cost_per_element"

    material = by_concept["bolsa de plástico chica 55x60"]
    assert material["cantidad"] == 1528
    assert material["precio_unitario"] == 0.0
    assert material["unidad"] == "KILO"
    assert material["extra"]["template_kind"] == "monthly_material_requirement"
    assert material["extra"]["zone"] == "A"
    assert material["extra"]["price_input_kind"] == "unit_cost"


def test_extract_line_items_marks_material_support_matrix_role(tmp_path):
    path = tmp_path / "soporte_materiales.xlsx"
    df = pd.DataFrame(
        {
            "Descripción del material": [
                "BOLSA DE PLÁSTICO CHICA 55X60",
                "ALCOHOL EN GEL",
                "ATOMIZADOR",
                "DETERGENTE EN POLVO",
                "FIBRA VERDE",
                "GERMICIDA",
                "GUANTES DE HULE",
                "JABÓN LÍQUIDO",
            ],
            "Código": [101, 102, 103, 104, 105, 106, 107, 108],
            "CAISES A": [1, 2, 3, 4, 5, 6, 7, 8],
            "CAISES B": [2, 3, 4, 5, 6, 7, 8, 9],
            "ZONA C": [3, 4, 5, 6, 7, 8, 9, 10],
            "TOTAL": [6, 9, 12, 15, 18, 21, 24, 27],
            "Entrega mensual": [1, 1, 1, 1, 1, 1, 1, 1],
        }
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="MATERIALES", index=False)

    rows = extract_line_items_from_excel_path(str(path), "soporte_materiales.xlsx")
    assert len(rows) >= 8
    by_concept = {r["concepto_norm"]: r for r in rows}
    alcohol = by_concept["alcohol en gel"]
    assert alcohol["extra"]["document_role"] == "material_support_matrix"
    assert alcohol["precio_unitario"] == 0.0
    assert alcohol["extra"]["price_values_suppressed"] is True


def test_extract_line_items_transposed_material_support_matrix(tmp_path):
    path = tmp_path / "matriz_transpuesta.xlsx"
    df = pd.DataFrame(
        [
            ["", "", "ANEXO SOPORTE DE MATERIALES", "", ""],
            ["ZONA", "DESCRIPCIÓN DEL MATERIAL", "", "ALCOHOL EN GEL", "ATOMIZADOR"],
            ["", "PRESENTACIÓN", "", "LITRO", "PIEZA"],
            ["", "UNIDAD MÉDICA", "ELEMENTOS", "CANTIDAD", "CANTIDAD"],
            ["A", "CAISES A", 2, 5, 1],
            ["A", "CAISES B", 3, 4, 2],
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="ZONA A", index=False, header=False)

    rows = extract_line_items_from_excel_path(str(path), "matriz_transpuesta.xlsx")
    by_concept = {r["concepto_norm"]: r for r in rows}
    alcohol = by_concept["alcohol en gel"]
    assert alcohol["extra"]["document_role"] == "material_support_matrix"
    assert alcohol["cantidad"] == 9.0
    assert alcohol["unidad"] == "LITRO"
    assert alcohol["precio_unitario"] == 0.0


def test_extract_line_items_location_price_grid_zb_style(tmp_path):
    """Regresión layout localidades × precio IVA incluido (estilo ZB, Ítem D.20)."""
    path = tmp_path / "anexo_zb.xlsx"
    df = pd.DataFrame(
        [
            ["", "ANEXO III PARTIDA 2 ZB — PROPUESTA ECONÓMICA", "", ""],
            ["", "LOCALIDAD", "NÚM. DE ELEMENTOS", "COSTO POR ELEMENTO I.V.A INCLUIDO"],
            ["", "Acambaro", 1, ""],
            ["", "Celaya", 1, ""],
            ["", "Salamanca", 1, ""],
            ["", "Leon", 1, ""],
            ["", "Irapuato", 1, ""],
            ["", "Silao", 1, ""],
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="ZB", index=False, header=False)

    rows = extract_line_items_from_excel_path(str(path), "33. Anexo III P1-2 ZB.xlsx")
    assert len(rows) >= 6
    by_concept = {r["concepto_norm"]: r for r in rows}
    acambaro = by_concept["acambaro"]
    assert acambaro["extra"]["template_kind"] == "location_price_grid"
    assert acambaro["extra"]["layout"] == "structured_template"
    header = str(acambaro["extra"].get("price_column_header") or "").lower()
    assert "incl" in header or "i.v.a" in header
    assert acambaro["precio_unitario"] in (None, 0.0)
    assert acambaro["extra"].get("price_input_pending") is True


def test_extract_line_items_material_support_list(tmp_path):
    path = tmp_path / "lista_soporte.xlsx"
    df = pd.DataFrame(
        [
            ["Anexo soporte", "", "", ""],
            ["Número", "Descripción del material", "Presentación", "Observaciones"],
            [1, "ALCOHOL EN GEL", "LITRO", ""],
            [2, "ATOMIZADOR", "PIEZA", ""],
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="MATERIALES", index=False, header=False)

    rows = extract_line_items_from_excel_path(str(path), "lista_soporte.xlsx")
    by_concept = {r["concepto_norm"]: r for r in rows}
    alcohol = by_concept["alcohol en gel"]
    assert alcohol["extra"]["document_role"] == "material_support_list"
    assert alcohol["cantidad"] is None
    assert alcohol["unidad"] == "LITRO"
    assert alcohol["precio_unitario"] == 0.0
