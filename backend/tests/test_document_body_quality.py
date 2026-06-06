"""Tests de cuerpo sustantivo universal en documentos."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.config.settings import settings as app_settings
from app.services.document_body_quality import (
    is_substantive_markdown,
    scan_materialized_doc_text,
    substantive_body_metrics,
)
from app.services.document_fill_quality_gate import validate_generated_documents_fill


def test_substantive_markdown_detects_protesta():
    text = (
        "Por medio de la presente, quien suscribe manifiesta bajo protesta de decir verdad "
        "que la empresa cumple con las bases y anexos del procedimiento de contratación."
    )
    assert is_substantive_markdown(text) is True


def test_shell_markdown_rejected():
    assert is_substantive_markdown("A QUIEN CORRESPONDA:\n________________") is False


def test_markdown_requires_protesta_by_default():
    filler = " ".join(["palabra"] * 50)
    assert is_substantive_markdown(filler) is False
    assert is_substantive_markdown(
        f"{filler} bajo protesta de decir verdad que cumplimos las bases."
    ) is True


def test_fill_gate_blocks_shell_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    monkeypatch.setattr(app_settings, "DOCUMENT_MIN_SUBSTANTIVE_WORDS", 40)
    path = tmp_path / "anexo_v.docx"
    doc = Document()
    doc.add_heading("PANEL_ADMINISTRATIVO|ANEXO V", 1)
    doc.add_paragraph("LUGAR Y FECHA: Ciudad a 3 de junio de 2026")
    doc.add_paragraph("A QUIEN CORRESPONDA:")
    doc.add_paragraph("_" * 40)
    doc.save(path)

    out = validate_generated_documents_fill(
        stage="formats",
        generated_documents=[{"ruta": str(path), "tipo": "administrativo"}],
        master_profile={
            "razon_social": "ACME",
            "rfc": "ACM010101AAA",
            "representante_legal": "Ana",
        },
    )
    assert out["validation_passed"] is False
    types = {i["error_type"] for i in out["issues"]}
    assert "document_shell_detected" in types or "document_metadata_leak" in types


def test_scan_materialized_returns_none_for_legal_body():
    hit = scan_materialized_doc_text(
        "El suscrito declara bajo protesta de decir verdad que cumple con las bases."
    )
    assert hit is None


def test_operational_technical_form_skips_shell_gate(tmp_path, monkeypatch):
    """Anexos III-B / partidas: espejo operativo, no manifiesto legal."""
    from app.services.document_body_quality import is_operational_technical_form_document

    assert is_operational_technical_form_document(
        "mirror_01_21_Anexo_III-B_Actividades_del_supervisor_de_limp.docx"
    )
    assert is_operational_technical_form_document(
        "mirror_04_24_Anexo_III-D_Partida_2_Entrega_de_materiales.docx"
    )
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    path = tmp_path / "mirror_01_21_Anexo_III-B_Actividades.docx"
    doc = Document()
    doc.add_paragraph("ACTIVIDADES DEL SUPERVISOR DE LIMPIEZA")
    doc.add_paragraph("HORARIO | ACTIVIDAD | FRECUENCIA")
    doc.add_paragraph("Lunes a viernes | Supervisión de turno | Diario")
    doc.add_paragraph("Nombre del supervisor: _________________________")
    doc.save(path)
    out = validate_generated_documents_fill(
        stage="technical",
        generated_documents=[
            {
                "ruta": str(path),
                "tipo": "tecnico_mirror",
                "materialization_route": "mirror",
            }
        ],
        master_profile={
            "razon_social": "MANAVIL SA",
            "rfc": "SPI060200AG5",
            "representante_legal": "Representante Legal",
        },
    )
    assert out["validation_passed"] is True


def test_operational_mirror_skips_adjudication_template_lexicon(tmp_path, monkeypatch):
    """Texto «proveedor adjudicado» en plantilla espejo III-B no debe bloquear técnica."""
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    path = tmp_path / "mirror_01_21_Anexo_III-B_Actividades.docx"
    doc = Document()
    doc.add_paragraph(
        "El proveedor adjudicado deberá cumplir las actividades del supervisor de limpieza."
    )
    doc.add_paragraph("ACTIVIDAD | FRECUENCIA | RESPONSABLE")
    doc.add_paragraph("Supervisión de turno | Diario | _________________________")
    doc.save(path)
    out = validate_generated_documents_fill(
        stage="technical",
        generated_documents=[
            {
                "ruta": str(path),
                "tipo": "tecnico_mirror",
                "materialization_route": "mirror",
                "mirror_mode": "copy",
            }
        ],
        master_profile={
            "razon_social": "MANAVIL SA",
            "rfc": "SPI060200AG5",
            "representante_legal": "Representante Legal",
        },
        provenance_context={"session_hint": "isapeg_servicios_de_limpieza"},
    )
    assert out["validation_passed"] is True
    assert not any(
        i.get("error_type") == "adjudication_language_in_proposal_stage"
        and i.get("severity") == "block"
        for i in (out.get("issues") or [])
    )


def test_catalog_template_skips_cross_tender_and_contract_dates(tmp_path, monkeypatch):
    """Plantillas catálogo ISAPEG: siglas hospital (CINCO) y fechas de vigencia no bloquean."""
    from app.services.document_fill_quality_gate import detect_cross_tender_marker

    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    path = tmp_path / "cat_anexo_s_modelo_contrato_federal.docx"
    doc = Document()
    doc.add_paragraph("Servicio en el área (CINCO) y también (CINCO) según bases.")
    doc.add_paragraph("Atención en (CEYE) y (CEYE) hospitalaria.")
    doc.add_paragraph("Criterios de evaluación del contrato federal.")
    doc.add_paragraph("Vigencia del 18 de diciembre de 2025 al 31 de diciembre de 2026.")
    doc.save(path)
    chunks = [p.text for p in Document(path).paragraphs if p.text.strip()]
    assert detect_cross_tender_marker(chunks, "isapeg_servicios_de_limpieza") in ("CINCO", "CEYE")
    out = validate_generated_documents_fill(
        stage="formats",
        generated_documents=[{"ruta": str(path), "materialization_route": "mirror"}],
        master_profile={
            "razon_social": "MANAVIL SA",
            "rfc": "SPI060200AG5",
            "representante_legal": "Representante Legal",
        },
        provenance_context={
            "session_hint": "isapeg_servicios_de_limpieza",
            "deadline_dt_iso": "2026-04-15T23:59:00",
            "fecha_es": "3 de junio de 2026",
        },
    )
    assert out["validation_passed"] is True


def test_pliego_xlsx_licitante_blank_not_blocking(tmp_path, monkeypatch):
    """Constancia de visitas (panel pliego): línea NOMBRE DEL LICITANTE: ___ es válida."""
    from app.services.document_fill_quality_gate import _is_placeholder_llm_residual_only

    line = "NOMBRE DEL LICITANTE:_____________________________________________________________"
    assert _is_placeholder_llm_residual_only(line) is False
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    path = tmp_path / "panel_pliego_constancia_visitas_Constancia_de_Visitas.xlsx"
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = line
    wb.save(path)
    out = validate_generated_documents_fill(
        stage="formats",
        generated_documents=[{"ruta": str(path)}],
        master_profile={
            "razon_social": "MANAVIL SA",
            "rfc": "SPI060200AG5",
            "representante_legal": "Representante Legal",
        },
        provenance_context={"session_hint": "isapeg_servicios_de_limpieza"},
    )
    assert out["validation_passed"] is True
