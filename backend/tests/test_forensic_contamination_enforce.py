"""Tests del gate forense P0 en empaquetado y escaneo CI."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.agents.packager import CompraNetPackager, build_pack_session_data_from_outputs
from app.config.settings import settings as app_settings
from app.services.delivery_content_audit import (
    forensic_contamination_blocking,
    run_forensic_contamination_audit,
)
from app.services.document_contamination_gate import (
    contamination_enforce_at_pack,
    scan_conflicting_document_dates,
)


def test_scan_conflicting_document_dates_detects_extra():
    text = (
        "LUGAR Y FECHA: Queretaro, a 23 de abril de 2026\n"
        "En Querétaro, a 3 de junio de 2026, manifiesto..."
    )
    hit = scan_conflicting_document_dates(
        text,
        canonical_fecha_es="23 de abril de 2026",
        dedupe_key="pliego|ANEXO_II",
        basename="Anexo_II.docx",
    )
    assert hit is not None
    assert hit.error_type == "document_multiple_dates_in_body"


def test_contamination_enforce_at_pack_respects_settings(monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    monkeypatch.setattr(app_settings, "DELIVERY_CONTAMINATION_ENFORCE_AT_PACK", True)
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    assert contamination_enforce_at_pack() is True
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "audit")
    assert contamination_enforce_at_pack() is False


def test_packager_blocks_contaminated_delivery(tmp_path, monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    monkeypatch.setattr(app_settings, "DELIVERY_CONTAMINATION_ENFORCE_AT_PACK", True)
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")

    root = tmp_path / "outputs" / "sess_test"
    sobre = root / "SOBRE_1_ADMINISTRATIVO"
    sobre.mkdir(parents=True)
    bad = sobre / "01_Anexo_X_malo.docx"
    doc = Document()
    doc.add_paragraph("Lo siento, no puedo generar contenido legal.")
    doc.save(bad)

    pack_session = build_pack_session_data_from_outputs(
        "sess_test",
        {"folder_raiz": str(root), "estructura_sobres": {}},
        {"master_profile": {"rfc": "ACM010101AAA", "razon_social": "ACME"}},
        session_state={},
    )
    result = CompraNetPackager().pack(pack_session)
    assert result.success is False
    assert result.validation_passed is False
    assert result.contamination_report is not None
    assert forensic_contamination_blocking(result.contamination_report)


def test_forensic_audit_clean_admin_letter(tmp_path, monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    staged = tmp_path / "validated"
    folder = staged / "SobreComplementaria"
    folder.mkdir(parents=True)
    fpath = folder / "Anexo V.docx"
    doc = Document()
    doc.add_paragraph("LUGAR Y FECHA: Ciudad, a 22 de abril de 2026")
    doc.add_paragraph("Bajo protesta de decir verdad, manifiesto cumplimiento.")
    doc.save(fpath)

    indice = [
        {
            "sobre": "SobreComplementaria",
            "nombre_entrega": "Anexo V.docx",
            "path": "SobreComplementaria/Anexo V.docx",
        }
    ]
    report = run_forensic_contamination_audit(
        "sess_clean",
        session_state={
            "last_analysis": {
                "cronograma": {"presentacion_proposiciones": "27 de abril de 2026"},
            }
        },
        validated_root=staged,
        indice_files=indice,
    )
    assert report.get("gate_passed") is True
    assert not forensic_contamination_blocking(report)


def test_forensic_audit_mirror_template_skips_adjudication_lexicon(tmp_path, monkeypatch):
    """Plantilla espejo (mirror_/cat_): léxico del pliego no bloquea entrega."""
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    monkeypatch.setattr(app_settings, "DELIVERY_CONTAMINATION_ENFORCE_AT_PACK", True)

    staged = tmp_path / "validated"
    folder = staged / "SobreTecnica"
    folder.mkdir(parents=True)
    fpath = folder / "mirror 01 21 Anexo III-B Actividades del supervisor de limpi.docx"
    doc = Document()
    doc.add_paragraph(
        "El proveedor adjudicado deberá entregar materiales conforme a criterios de evaluación."
    )
    doc.save(fpath)

    indice = [
        {
            "sobre": "SobreTecnica",
            "nombre_entrega": fpath.name,
            "path": f"SobreTecnica/{fpath.name}",
        }
    ]
    report = run_forensic_contamination_audit(
        "sess_mirror",
        session_state={},
        validated_root=staged,
        indice_files=indice,
    )
    assert report.get("gate_passed") is True
    assert not forensic_contamination_blocking(report)


def test_forensic_audit_unaq_style_letter_still_blocks_llm_refusal(tmp_path, monkeypatch):
    """Carta redactada (FO-35 / propuesta técnica UNAQ): rechazo LLM sigue bloqueando."""
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    monkeypatch.setattr(app_settings, "DELIVERY_CONTAMINATION_ENFORCE_AT_PACK", True)

    staged = tmp_path / "validated"
    folder = staged / "SobreTecnica"
    folder.mkdir(parents=True)
    fpath = folder / "FO-35_Modelo_presentacion_Propuesta_Tecnica.docx"
    doc = Document()
    doc.add_paragraph("Lo siento, no puedo generar contenido legal.")
    doc.save(fpath)

    indice = [
        {
            "sobre": "SobreTecnica",
            "nombre_entrega": fpath.name,
            "path": f"SobreTecnica/{fpath.name}",
        }
    ]
    report = run_forensic_contamination_audit(
        "sess_unaq_refusal",
        session_state={},
        validated_root=staged,
        indice_files=indice,
    )
    assert report.get("gate_passed") is False
    assert forensic_contamination_blocking(report)


def test_forensic_audit_unaq_style_letter_still_blocks_adjudication(tmp_path, monkeypatch):
    """Carta del licitante (no espejo): lenguaje de adjudicado sigue bloqueando."""
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    monkeypatch.setattr(app_settings, "DELIVERY_CONTAMINATION_ENFORCE_AT_PACK", True)

    staged = tmp_path / "validated"
    folder = staged / "SobreComplementaria"
    folder.mkdir(parents=True)
    fpath = folder / "02_CARTA_COMPROMISO.docx"
    doc = Document()
    doc.add_paragraph("Bajo protesta de decir verdad declaro que fuimos seleccionados como proveedor adjudicado.")
    doc.save(fpath)

    indice = [
        {
            "sobre": "SobreComplementaria",
            "nombre_entrega": fpath.name,
            "path": f"SobreComplementaria/{fpath.name}",
        }
    ]
    report = run_forensic_contamination_audit(
        "sess_unaq_adj",
        session_state={},
        validated_root=staged,
        indice_files=indice,
    )
    assert report.get("gate_passed") is False
    assert forensic_contamination_blocking(report)
