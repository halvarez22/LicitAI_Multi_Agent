"""Tests universales: experiencia desde Fuentes sin hardcode de licitación."""
from __future__ import annotations

import docx

from app.services.company_experience_context import (
    build_experience_sources_ux_summary,
    extract_client_references_from_documents,
    fill_te03_client_placeholders,
)

_GENERIC_CONSTANCIA = """
Coordinación General de Recursos Materiales
Dependencia de Servicios Integrales del Estado de Ejemplo
Asunto: Constancia de contrato

Por medio de la presente hago constar que se tiene signado el contrato número ABC-2024-001,
correspondiente al Servicio de Limpieza para las Unidades de este Organismo Estatal de Ejemplo,
con vigencia del 1 de enero al 31 de diciembre de 2024.

Av. Reforma #100, Col. Centro, Ciudad Ejemplo, Edo. C.P. 01000
Tel. (55) 1234 5678
"""


def test_extract_client_references_generic_constancia():
    docs = [
        {
            "content": {
                "filename": "curriculum_empresarial.pdf",
                "status": "ANALYZED",
                "extracted_text": _GENERIC_CONSTANCIA,
            }
        }
    ]
    refs = extract_client_references_from_documents(docs)
    assert refs
    assert refs[0].get("contrato") == "ABC-2024-001"
    assert refs[0].get("telefono")
    assert refs[0].get("domicilio")
    assert "Ejemplo" in refs[0]["nombre"] or "Organismo" in refs[0]["nombre"]


def test_build_experience_sources_ux_summary_uses_source_filename():
    docs = [
        {
            "content": {
                "filename": "curriculum_empresarial.pdf",
                "status": "ANALYZED",
                "extracted_text": _GENERIC_CONSTANCIA,
            }
        }
    ]
    summary = build_experience_sources_ux_summary(docs)
    assert summary
    assert "curriculum_empresarial.pdf" in summary
    assert "experiencia previa.pdf" not in summary


def test_fill_te03_client_placeholders_generic():
    import tempfile
    import os

    with tempfile.TemporaryDirectory() as tmp:
        f = os.path.join(tmp, "te03.docx")
        doc = docx.Document()
        doc.add_paragraph("| 1 |  | [Domicilio del cliente 1] | [Teléfono del cliente 1] |")
        doc.save(f)
        clients = [
            {
                "nombre": "Organismo Estatal de Ejemplo",
                "domicilio": "Av. Reforma #100, C.P. 01000",
                "telefono": "5512345678",
            }
        ]
        assert fill_te03_client_placeholders(f, clients) is True
        saved = docx.Document(f)
        text = saved.paragraphs[0].text
        assert "01000" in text
        assert "[Domicilio" not in text
