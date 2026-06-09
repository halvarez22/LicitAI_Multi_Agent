"""Inserción de logo en DOCX: JPEG progresivo y formatos habituales."""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from app.utils.doc_formatting import add_logo_picture_to_run, apply_corporate_docx_letterhead


def _header_has_image(doc: Document) -> bool:
    header = doc.sections[0].header
    xml = header._element.xml
    return "blip" in xml or bool(header._element.findall(".//" + qn("wp:inline")))


def test_progressive_jpeg_logo_inserts_via_pil_fallback(tmp_path):
    logo = tmp_path / "logo_progressive.jpg"
    im = Image.new("RGB", (120, 80), color=(20, 90, 160))
    im.save(logo, format="JPEG", progressive=True)

    doc = Document()
    ok = add_logo_picture_to_run(doc.add_paragraph().add_run(), str(logo))
    assert ok is True
    assert _header_has_image(doc) is False  # run en cuerpo, no header

    doc2 = Document()
    apply_corporate_docx_letterhead(
        doc2,
        {
            "logo_path": str(logo),
            "tender_name": "LICITACION TEST",
            "footer_text": "Empresa Test | RFC: TST010101TST",
        },
    )
    assert _header_has_image(doc2) is True


def test_png_logo_inserts_directly(tmp_path):
    logo = tmp_path / "logo.png"
    Image.new("RGBA", (64, 64), color=(255, 0, 0, 200)).save(logo, format="PNG")

    doc = Document()
    apply_corporate_docx_letterhead(doc, {"logo_path": str(logo), "tender_name": "X"})
    assert _header_has_image(doc) is True
