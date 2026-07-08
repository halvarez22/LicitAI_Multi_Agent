"""Tests universales de fecha documental (HRU)."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from app.config.settings import settings as app_settings
from app.services.document_date_resolver import (
    apply_document_date_override_from_chat,
    extract_document_date_user_override,
    normalize_body_spanish_dates,
    normalize_docx_spanish_dates,
    parse_document_date_override_from_chat,
    resolve_document_date,
    resolve_generation_header_date,
)
from app.services.document_fill_quality_gate import validate_generated_documents_fill


def test_resolve_document_date_defaults_to_generation_timestamp():
    state = {
        "submission_checklist": {
            "hitos": [
                {
                    "id": "presentacion_proposiciones",
                    "fecha_texto_raw": "19 de diciembre de 2025, 09:30",
                    "fecha_hora": "2025-12-19T09:30:00",
                }
            ]
        }
    }
    gen_at = datetime(2025, 12, 18, 9, 0, 0)
    out = resolve_document_date(state, at=gen_at)
    assert out["source"] == "generation_timestamp"
    assert out["fecha_es"] == "18 de diciembre de 2025"
    assert out["deadline_dt"] is not None
    assert out["is_after_deadline"] is False


def test_user_override_from_session():
    state = {
        "document_date_user_override": "17 de diciembre de 2025",
        "submission_checklist": {
            "hitos": [
                {
                    "id": "presentacion_proposiciones",
                    "fecha_hora": "2025-12-19T09:30:00",
                }
            ]
        },
    }
    assert extract_document_date_user_override(state) == "17 de diciembre de 2025"
    out = resolve_document_date(state)
    assert out["source"] == "user_override"
    assert out["fecha_es"] == "17 de diciembre de 2025"


def test_parse_document_date_from_chat():
    msg = "Fecha de los formatos administrativos: 17 de diciembre de 2025"
    assert parse_document_date_override_from_chat(msg) == "17 de diciembre de 2025"


def test_apply_chat_override_clamps_after_deadline():
    state = {
        "submission_checklist": {
            "hitos": [
                {
                    "id": "presentacion_proposiciones",
                    "fecha_hora": "2025-12-19T09:30:00",
                }
            ]
        }
    }
    hitl = apply_document_date_override_from_chat(
        state,
        "fecha documentos: 8 de junio de 2026",
    )
    assert hitl["applied"] is True
    assert hitl["fecha_es"] == "19 de diciembre de 2025"
    assert hitl["session_patch"]["document_date_override_provenance"]["clamped_to_deadline"] is True


def test_normalize_body_spanish_dates_replaces_wrong_dates():
    canon = "17 de diciembre de 2025"
    raw = "Carta-Compromiso en León, a 8 de junio de 2026."
    out = normalize_body_spanish_dates(raw, canon)
    assert "8 de junio" not in out
    assert canon in out


def test_normalize_docx_spanish_dates(tmp_path):
    canon = "17 de diciembre de 2025"
    path = tmp_path / "carta.docx"
    doc = Document()
    doc.add_paragraph("En León, Guanajuato, a 20 de diciembre de 2025.")
    doc.save(path)
    changed = normalize_docx_spanish_dates(str(path), canon)
    assert changed is True
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "20 de diciembre" not in text
    assert canon in text


def test_fill_gate_no_date_block_when_canonical_applied(tmp_path, monkeypatch):
    monkeypatch.setattr(app_settings, "DOCUMENT_FILL_QUALITY_GATE_MODE", "enforce")
    monkeypatch.setattr(app_settings, "DOCUMENT_CONTAMINATION_GATE_ENABLED", True)
    f = tmp_path / "carta.docx"
    doc = Document()
    doc.add_paragraph("CONSTRUCTORA SA DE CV")
    doc.add_paragraph("RFC: CIN2506089A3")
    doc.add_paragraph("En León, a 17 de diciembre de 2025.")
    doc.save(f)
    out = validate_generated_documents_fill(
        stage="formats",
        generated_documents=[{"ruta": str(f), "tipo": "administrativo", "template_id": "carta_compromiso"}],
        master_profile={
            "razon_social": "CONSTRUCTORA SA DE CV",
            "rfc": "CIN2506089A3",
            "representante_legal": "Juan Perez",
        },
        provenance_context={
            "source": "formats_writer",
            "confidence": 0.9,
            "fecha_es": "17 de diciembre de 2025",
            "deadline_dt_iso": "2025-12-19T09:30:00",
        },
    )
    date_blocks = [
        i
        for i in out["issues"]
        if i.get("error_type") == "document_date_after_submission_deadline"
    ]
    assert not date_blocks


def test_resolve_generation_and_documental_match_without_override():
    state = {
        "submission_checklist": {
            "hitos": [
                {
                    "id": "presentacion_proposiciones",
                    "fecha_hora": "2025-12-19T09:30:00",
                }
            ]
        }
    }
    gen_at = datetime(2025, 12, 18, 10, 30, 0)
    doc = resolve_document_date(state, at=gen_at)
    gen = resolve_generation_header_date(at=gen_at)
    assert doc["source"] == "generation_timestamp"
    assert doc["fecha_es"] == gen["fecha_es"]


def test_normalize_docx_preserves_generation_header_stamp(tmp_path):
    canon = "17 de diciembre de 2025"
    gen = "29 de junio de 2026"
    path = tmp_path / "carta.docx"
    doc = Document()
    doc.add_paragraph(f"LUGAR Y FECHA: León, a {gen}")
    doc.add_paragraph("En León, a 8 de junio de 2026.")
    doc.save(path)
    changed = normalize_docx_spanish_dates(str(path), canon)
    assert changed is True
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert gen in text
    assert canon in text
    assert "8 de junio" not in text


def test_save_docx_official_mirror_prints_generation_header(tmp_path):
    from app.agents.formats import _save_docx

    gen = resolve_generation_header_date(at=datetime(2026, 6, 29, 14, 0, 0))
    path = tmp_path / "e1.docx"
    _save_docx(
        "ANEXO E-1",
        "CARTA-COMPROMISO DE LA PROPOSICIÓN\n\nContenido espejo bases.",
        str(path),
        {
            "official_bases_mirror": True,
            "obra_pliego_contract": True,
            "fecha_documental": gen["fecha_es"],
            "fecha_encabezado": gen["fecha_es"],
            "ciudad": "León",
            "footer_text": "EMPRESA | RFC: X | Domicilio: León, Gto.",
            "formal_closing": False,
        },
    )
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "LUGAR Y FECHA: León, a 29 de junio de 2026" in text
    assert "17 de diciembre de 2025" not in text.split("LUGAR Y FECHA")[0]


def test_generation_header_stamp_exempt_from_deadline_gate():
    from app.services.document_contamination_gate import scan_all_document_dates

    text = (
        "LUGAR Y FECHA: León, a 29 de junio de 2026\n"
        "En León, a 17 de diciembre de 2025.\n"
    )
    hits = scan_all_document_dates(
        text,
        deadline_dt_iso="2025-12-19T09:30:00",
        canonical_fecha_es="17 de diciembre de 2025",
    )
    assert not hits
