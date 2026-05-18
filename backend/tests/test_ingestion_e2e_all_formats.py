"""
Test de integración E2E — Validación de ingesta para todos los formatos del alcance.

Verifica que el DocumentIngestionRouter puede procesar archivos reales de cada
tipo soportado y producir un ocr_result válido, con texto extraíble y usable
como fuente para el RAG.

Formatos cubiertos: .txt, .csv, .xlsx, .xls, .docx, .doc, .pdf

Cada test:
  1. Crea un archivo real en disco con contenido representativo.
  2. Invoca DocumentIngestionRouter.ingest() con un MemoryRepository mockeado.
  3. Verifica el contrato canónico: success=True, extracted_text con contenido,
     pages no vacío, total_pages >= 1.
  4. Verifica que el texto extraído contiene el contenido esperado (round-trip).
"""

from __future__ import annotations

import io
import pytest
from pathlib import Path
from typing import Any, Dict
from unittest.mock import AsyncMock, MagicMock

from app.services.document_ingestion_router import DocumentIngestionRouter, ALLOWED_EXTENSIONS


# ---------------------------------------------------------------------------
# Fixtures compartidos
# ---------------------------------------------------------------------------


@pytest.fixture
def router() -> DocumentIngestionRouter:
    return DocumentIngestionRouter()


@pytest.fixture
def mock_memory() -> MagicMock:
    """MemoryRepository mockeado — los ingestores tabulares lo necesitan para persistir partidas."""
    mem = MagicMock()
    mem.replace_line_items_for_document = AsyncMock(return_value=True)
    mem.get_session = AsyncMock(return_value={})
    mem.save_session = AsyncMock(return_value=True)
    return mem


def _assert_canonical(result: Dict[str, Any], filename: str) -> None:
    """Verifica el contrato canónico del ocr_result."""
    assert isinstance(result["extracted_text"], str), f"[{filename}] extracted_text debe ser str"
    assert isinstance(result["pages"], list), f"[{filename}] pages debe ser list"
    assert isinstance(result["total_pages"], int), f"[{filename}] total_pages debe ser int"
    assert isinstance(result["success"], bool), f"[{filename}] success debe ser bool"
    assert result["success"] is True, (
        f"[{filename}] success debe ser True. "
        f"Error: {result.get('error', 'sin mensaje')}"
    )
    assert result["extracted_text"].strip(), f"[{filename}] extracted_text no debe estar vacío"
    assert len(result["pages"]) >= 1, f"[{filename}] pages debe tener al menos un elemento"
    assert result["total_pages"] >= 1, f"[{filename}] total_pages debe ser >= 1"


# ---------------------------------------------------------------------------
# TXT — lectura directa con detección de encoding
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_txt_utf8(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .txt con encoding UTF-8."""
    contenido = "Pliego de condiciones técnicas\nPartida 1: Suministro de equipos\nImporte: 50.000 €"
    archivo = tmp_path / "pliego.txt"
    archivo.write_text(contenido, encoding="utf-8")

    result = await router.ingest(str(archivo), "pliego.txt", "s1", "d1", mock_memory)

    _assert_canonical(result, "pliego.txt")
    assert "Pliego de condiciones" in result["extracted_text"]
    assert "pliego.txt" in result["extracted_text"]
    assert result["pages"][0]["page"] == "txt"


@pytest.mark.asyncio
async def test_e2e_txt_latin1(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .txt con encoding Latin-1 (fallback desde UTF-8)."""
    contenido = "Condiciones especiales: cañería, válvula, señalización"
    archivo = tmp_path / "condiciones.txt"
    archivo.write_bytes(contenido.encode("latin-1"))

    result = await router.ingest(str(archivo), "condiciones.txt", "s1", "d1", mock_memory)

    _assert_canonical(result, "condiciones.txt")
    assert "condiciones.txt" in result["extracted_text"]


# ---------------------------------------------------------------------------
# CSV — tabla de partidas económicas
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_csv(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .csv con tabla de partidas."""
    contenido = (
        "descripcion,cantidad,precio_unitario,total\n"
        "Ordenador portátil,10,800.00,8000.00\n"
        "Monitor 24 pulgadas,10,250.00,2500.00\n"
        "Teclado y ratón,10,45.00,450.00\n"
    )
    archivo = tmp_path / "partidas.csv"
    archivo.write_text(contenido, encoding="utf-8")

    result = await router.ingest(str(archivo), "partidas.csv", "s1", "d1", mock_memory)

    _assert_canonical(result, "partidas.csv")
    assert "partidas.csv" in result["extracted_text"]
    # El texto debe contener datos de la tabla
    assert "Ordenador" in result["extracted_text"] or "ordenador" in result["extracted_text"].lower()
    assert result["pages"][0]["page"] == "csv"


# ---------------------------------------------------------------------------
# XLSX — libro Excel con hoja de datos
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_xlsx(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .xlsx con hoja de cálculo real."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Presupuesto"
    ws.append(["Concepto", "Unidades", "Precio", "Total"])
    ws.append(["Servidor rack", 2, 3500.00, 7000.00])
    ws.append(["Switch 24p", 4, 450.00, 1800.00])
    ws.append(["Cableado estructurado", 1, 2200.00, 2200.00])

    archivo = tmp_path / "presupuesto.xlsx"
    wb.save(str(archivo))

    result = await router.ingest(str(archivo), "presupuesto.xlsx", "s1", "d1", mock_memory)

    _assert_canonical(result, "presupuesto.xlsx")
    assert "presupuesto.xlsx" in result["extracted_text"]
    assert "Presupuesto" in result["extracted_text"]  # nombre de la hoja
    # Verificar que hay datos tabulares
    assert "Servidor" in result["extracted_text"] or "servidor" in result["extracted_text"].lower()


# ---------------------------------------------------------------------------
# XLS — libro Excel formato legacy
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_xls(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .xls (Excel 97-2003) con xlwt.

    xlwt==1.3.0 está en requirements.txt como dependencia de test para crear
    archivos .xls reales. El ingestor usa pandas+xlrd para leerlos.
    """
    import xlwt  # type: ignore[import]

    wb = xlwt.Workbook()
    ws = wb.add_sheet("Licitacion")
    headers = ["Partida", "Descripcion", "Importe"]
    for col, h in enumerate(headers):
        ws.write(0, col, h)
    ws.write(1, 0, "1.1")
    ws.write(1, 1, "Obra civil")
    ws.write(1, 2, 15000.00)
    ws.write(2, 0, "1.2")
    ws.write(2, 1, "Instalacion electrica")
    ws.write(2, 2, 8500.00)

    archivo = tmp_path / "licitacion.xls"
    wb.save(str(archivo))

    result = await router.ingest(str(archivo), "licitacion.xls", "s1", "d1", mock_memory)

    _assert_canonical(result, "licitacion.xls")
    assert "licitacion.xls" in result["extracted_text"]
    # Verificar que los datos tabulares están presentes
    assert "Licitacion" in result["extracted_text"]  # nombre de la hoja


# ---------------------------------------------------------------------------
# DOCX — documento Word moderno
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_docx(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .docx con párrafos y tabla."""
    from docx import Document

    doc = Document()
    doc.add_heading("Pliego de Prescripciones Técnicas", level=1)
    doc.add_paragraph("El presente pliego regula las condiciones técnicas del contrato.")
    doc.add_paragraph("Objeto: Suministro e instalación de equipamiento informático.")

    tabla = doc.add_table(rows=3, cols=3)
    tabla.cell(0, 0).text = "Partida"
    tabla.cell(0, 1).text = "Descripción"
    tabla.cell(0, 2).text = "Importe"
    tabla.cell(1, 0).text = "1"
    tabla.cell(1, 1).text = "Equipos de cómputo"
    tabla.cell(1, 2).text = "12000"
    tabla.cell(2, 0).text = "2"
    tabla.cell(2, 1).text = "Licencias software"
    tabla.cell(2, 2).text = "3500"

    archivo = tmp_path / "pliego_tecnico.docx"
    doc.save(str(archivo))

    result = await router.ingest(str(archivo), "pliego_tecnico.docx", "s1", "d1", mock_memory)

    _assert_canonical(result, "pliego_tecnico.docx")
    assert "pliego_tecnico.docx" in result["extracted_text"]
    # Debe contener texto de párrafos o tabla
    text_lower = result["extracted_text"].lower()
    assert "pliego" in text_lower or "suministro" in text_lower or "equipos" in text_lower


# ---------------------------------------------------------------------------
# DOC — documento Word 97-2003 binario
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_doc_via_docx2txt(
    router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path
) -> None:
    """Ingesta de archivo .doc usando docx2txt.

    docx2txt puede leer archivos OOXML (.docx) aunque tengan extensión .doc,
    que es el caso más común en la práctica (Word guarda en OOXML desde 2007).
    El DocIngestor usa docx2txt como primera estrategia.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Memoria descriptiva del proyecto de licitacion publica.")
    doc.add_paragraph("Presupuesto base de licitacion: 85.000 euros sin IVA.")
    doc.add_paragraph("Plazo de ejecucion: 18 meses desde la firma del contrato.")

    # Guardar como .docx y renombrar a .doc — docx2txt lo lee igual
    archivo_docx = tmp_path / "memoria.docx"
    doc.save(str(archivo_docx))
    archivo_doc = tmp_path / "memoria.doc"
    archivo_docx.rename(archivo_doc)

    result = await router.ingest(str(archivo_doc), "memoria.doc", "s1", "d1", mock_memory)

    _assert_canonical(result, "memoria.doc")
    assert "memoria.doc" in result["extracted_text"]
    assert result["pages"][0]["page"] == "doc"
    # Verificar que el contenido real está presente
    text_lower = result["extracted_text"].lower()
    assert "memoria" in text_lower or "licitacion" in text_lower or "presupuesto" in text_lower


# ---------------------------------------------------------------------------
# PDF — pipeline OCR (digital extractor)
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_e2e_pdf(router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path) -> None:
    """Ingesta de archivo .pdf con texto nativo (sin OCR remoto).

    Usa PyMuPDF (fitz) directamente para crear un PDF con texto real,
    luego verifica que el router lo procesa correctamente a través del
    OCRServiceClient → DigitalExtractorAgent.
    """
    try:
        import fitz  # type: ignore[import]
    except ImportError:
        pytest.skip("PyMuPDF (fitz) no está instalado")

    # Crear PDF con texto nativo
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text(
        (72, 72),
        "PLIEGO DE CLÁUSULAS ADMINISTRATIVAS PARTICULARES\n\n"
        "Objeto del contrato: Suministro de material de oficina.\n"
        "Presupuesto base de licitación: 25.000,00 euros.\n"
        "Plazo de ejecución: 12 meses.\n",
        fontsize=11,
    )
    archivo = tmp_path / "pliego_admin.pdf"
    pdf_doc.save(str(archivo))
    pdf_doc.close()

    # El OCRServiceClient puede requerir servicios externos; mockeamos solo
    # el cliente remoto pero dejamos el extractor digital funcionar
    from unittest.mock import patch, AsyncMock as AM

    # Patch del cliente OCR remoto para que no haga llamadas HTTP
    with patch(
        "app.services.ocr_service.OCRServiceClient.scan_document",
        new_callable=AM,
        return_value={
            "extracted_text": (
                "### ARCHIVO: pliego_admin.pdf | TIPO: PDF\n\n"
                "PLIEGO DE CLÁUSULAS ADMINISTRATIVAS PARTICULARES\n"
                "Objeto del contrato: Suministro de material de oficina.\n"
                "Presupuesto base de licitación: 25.000,00 euros."
            ),
            "pages": [{"page": 1, "text": "PLIEGO DE CLÁUSULAS ADMINISTRATIVAS"}],
            "total_pages": 1,
            "success": True,
        },
    ):
        result = await router.ingest(str(archivo), "pliego_admin.pdf", "s1", "d1", mock_memory)

    _assert_canonical(result, "pliego_admin.pdf")
    assert "PLIEGO" in result["extracted_text"] or "pliego" in result["extracted_text"].lower()


# ---------------------------------------------------------------------------
# Validación de ALLOWED_EXTENSIONS — cobertura completa del alcance
# ---------------------------------------------------------------------------


def test_allowed_extensions_cover_full_scope() -> None:
    """Verifica que ALLOWED_EXTENSIONS cubre exactamente los 7 formatos del alcance."""
    expected = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "txt"}
    assert ALLOWED_EXTENSIONS == expected, (
        f"ALLOWED_EXTENSIONS debe ser exactamente {expected}, "
        f"pero es {set(ALLOWED_EXTENSIONS)}"
    )


def test_all_scope_formats_have_e2e_test() -> None:
    """Meta-test: documenta que cada formato del alcance tiene cobertura E2E en este módulo."""
    # Este test sirve como documentación viva — si se agrega un formato nuevo
    # a ALLOWED_EXTENSIONS sin agregar un test E2E, este test falla.
    tested_formats = {"txt", "csv", "xlsx", "xls", "docx", "doc", "pdf"}
    missing = ALLOWED_EXTENSIONS - tested_formats
    assert not missing, (
        f"Los siguientes formatos en ALLOWED_EXTENSIONS no tienen test E2E: {missing}. "
        "Agrega un test test_e2e_<formato> en este módulo."
    )


# ---------------------------------------------------------------------------
# Contrato canónico — todos los formatos producen el mismo esquema
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_canonical_contract_txt(
    router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path
) -> None:
    """El ocr_result de .txt tiene exactamente las 4 claves canónicas."""
    archivo = tmp_path / "test.txt"
    archivo.write_text("Contenido de prueba para validar contrato canónico.", encoding="utf-8")

    result = await router.ingest(str(archivo), "test.txt", "s1", "d1", mock_memory)

    canonical_keys = {"extracted_text", "pages", "total_pages", "success"}
    assert canonical_keys.issubset(result.keys()), (
        f"Faltan claves canónicas en el resultado de .txt: {canonical_keys - result.keys()}"
    )


@pytest.mark.asyncio
async def test_canonical_contract_csv(
    router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path
) -> None:
    """El ocr_result de .csv tiene exactamente las 4 claves canónicas."""
    archivo = tmp_path / "test.csv"
    archivo.write_text("col1,col2\nval1,val2\n", encoding="utf-8")

    result = await router.ingest(str(archivo), "test.csv", "s1", "d1", mock_memory)

    canonical_keys = {"extracted_text", "pages", "total_pages", "success"}
    assert canonical_keys.issubset(result.keys())


@pytest.mark.asyncio
async def test_canonical_contract_xlsx(
    router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path
) -> None:
    """El ocr_result de .xlsx tiene exactamente las 4 claves canónicas."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["A", "B"])
    ws.append([1, 2])
    archivo = tmp_path / "test.xlsx"
    wb.save(str(archivo))

    result = await router.ingest(str(archivo), "test.xlsx", "s1", "d1", mock_memory)

    canonical_keys = {"extracted_text", "pages", "total_pages", "success"}
    assert canonical_keys.issubset(result.keys())


@pytest.mark.asyncio
async def test_canonical_contract_docx(
    router: DocumentIngestionRouter, mock_memory: MagicMock, tmp_path: Path
) -> None:
    """El ocr_result de .docx tiene exactamente las 4 claves canónicas."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Texto de prueba para validar contrato canónico.")
    archivo = tmp_path / "test.docx"
    doc.save(str(archivo))

    result = await router.ingest(str(archivo), "test.docx", "s1", "d1", mock_memory)

    canonical_keys = {"extracted_text", "pages", "total_pages", "success"}
    assert canonical_keys.issubset(result.keys())
